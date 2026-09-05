"""PDF/Word import engine + import endpoints.

KEPT VERBATIM from the original backend (14.8.0) - the PDF pipeline is
intentionally untouched. Only cross-module names are imported.
"""
import base64
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import time
import uuid

import pymupdf
from flask import Response, jsonify, request, send_from_directory
from PIL import Image

from .cloud import _publish_live
from .common import BASE_DIR, DOCS_DIR, IMG_DIR, JOBS_DIR, UPLOAD_DIR
from .core import app
from .notify import _notify_add


# ============ 문서 가져오기 (PDF / Word → 편집 가능한 요소) ============
# PDF 와 Word 를 읽어서 프런트의 문서 형식(pages[].els[])으로 변환한다.
#  · 글자 → type:'text'   (띄어쓰기 단위 단어별 상자, 정확한 위치)
#  · 그림 → type:'image'  (배경: 원본 그대로)
# 프런트 A4 용지 방향에 맞춰 세로 800 x 1100 또는 가로 1100 x 800
# 좌표계로 환산한다. 한 문서 안에 방향이 섞이면 더 많은 쪽의 방향을 따른다.

PAGE_W, PAGE_H = 800, 1100
PAGE_PRESETS = {
    "a4_portrait": (800, 1100),
    "a4_landscape": (1100, 800),
}


def _preset_dims(size_preset):
    return PAGE_PRESETS.get(size_preset, PAGE_PRESETS["a4_portrait"])


def _pdf_size_preset(src):
    """PDF 쪽 비율의 다수결로 노트의 세로/가로 용지를 고른다."""
    doc = None
    try:
        doc = pymupdf.open(src)
        portrait = landscape = 0
        for pno in range(min(doc.page_count, IMPORT_MAX_PAGES)):
            rect = doc[pno].rect
            if rect.width > rect.height:
                landscape += 1
            else:
                portrait += 1
        return "a4_landscape" if landscape > portrait else "a4_portrait"
    except Exception as e:
        print(f"[import] PDF 방향 판별 실패, 세로로 처리: {e}")
        return "a4_portrait"
    finally:
        if doc is not None:
            try:
                doc.close()
            except Exception:
                pass


def _docx_size_preset(data):
    """Word 구역(section) 용지 비율의 다수결로 세로/가로를 고른다."""
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        word_doc = Document(io.BytesIO(data))
        portrait = landscape = 0
        for section in word_doc.sections:
            is_landscape = section.orientation == WD_ORIENT.LANDSCAPE
            if is_landscape or int(section.page_width or 0) > int(section.page_height or 0):
                landscape += 1
            else:
                portrait += 1
        return "a4_landscape" if landscape > portrait else "a4_portrait"
    except Exception as e:
        print(f"[import] Word 방향 판별 실패, 세로로 처리: {e}")
        return "a4_portrait"


IMPORT_MAX_MB = 120
IMPORT_MAX_PAGES = 1200    # 500쪽+ 원서 대응
IMPORT_BG_DPI = 300       # 배경 래스터 해상도 (SVG 실패/사진 쪽 예비값)
IMPORT_BG_MAX_PX = 9_000_000   # 배경 래스터 픽셀 상한 (고화질 사진 대응)
# 쪽 수가 많으면 디스크/변환 시간을 위해 화질을 자동 조절 (자식 프로세스별 적용)
BG_DPI_CUR = IMPORT_BG_DPI
BG_PX_CUR = IMPORT_BG_MAX_PX


BG_JPEG_Q = 62      # 배경 JPEG 품질 (글자는 텍스트 상자로 분리돼 있어 넉넉함)


def _bg_tier(total):
    if total >= 400:
        return 150, 2_500_000
    if total >= 200:
        return 170, 3_000_000
    return 200, 4_000_000


def _save_import_svg(svg_text):
    """배경 SVG 를 파일로 저장하고 주소를 돌려준다."""
    try:
        # 공백을 줄여 크기를 아낀다 (내용은 그대로)
        s = re.sub(r">\s+<", "><", svg_text)
        s = re.sub(r"[ \t]{2,}", " ", s)
        name = f"{uuid.uuid4().hex[:16]}.svg"
        with open(os.path.join(IMG_DIR, name), "w", encoding="utf-8") as fp:
            fp.write(s)
        return f"/api/import/img/{name}"
    except Exception:
        return None


class _NoBackground(Exception):
    """이 쪽은 배경(그림·선)이 없어 굽지 않는다는 내부 신호."""


def _save_pixmap_bg(pm):
    """픽스맵을 곧바로 파일로 굽는다 (base64/Pillow 왕복 없음).

    옛 경로: pixmap → PNG bytes → Pillow 열기 → 재인코딩 → base64 →
             다시 디코드 → 파일. 쪽당 0.5초 이상을 여기서 태웠다.
    새 경로: pixmap → (JPEG|PNG) → 파일. 같은 화질에 수십 배 빠르다.
    글자는 이미 텍스트 상자로 분리돼 배경엔 그림·선만 남으므로
    JPEG 로 구워도 눈에 보이는 손해가 없다.
    """
    try:
        # 배경은 두 부류다.
        #  · 선화(그래프·표·도형): 색이 몇 가지뿐 → PNG 가 몇 배 작고 선명
        #  · 사진: 색이 많음 → JPEG 가 몇 배 작다
        # 픽셀을 성기게 훑어 색 가짓수로 판별한다 (수 ms).
        photo = False
        try:
            s = pm.samples
            nch = pm.n
            npx = pm.width * pm.height
            if nch >= 3 and npx > 0:
                step = max(1, npx // 20000)
                cols = set()
                for i in range(0, npx, step):
                    o = i * nch
                    cols.add(s[o:o + 3])
                    if len(cols) > 700:
                        photo = True
                        break
        except Exception:
            photo = pm.n >= 3

        if photo:
            name = f"{uuid.uuid4().hex[:16]}.jpg"
            pm.save(os.path.join(IMG_DIR, name), jpg_quality=BG_JPEG_Q)
        else:
            name = f"{uuid.uuid4().hex[:16]}.png"
            pm.save(os.path.join(IMG_DIR, name))
        return f"/api/import/img/{name}"
    except Exception as e:
        print(f"[import] 배경 저장 실패: {e}")
        return None


def _save_import_img(data_url):
    """data URL 을 파일로 저장하고 짧은 주소를 돌려준다.

    문서에 그림을 통째로 넣으면 브라우저 저장 한도를 넘겨
    노트가 통째로 사라진다. 주소만 넣으면 수십 배 가볍다.
    """
    try:
        head, b64data = data_url.split(",", 1)
        ext = "png" if "png" in head else "jpg"
        raw = base64.b64decode(b64data)
        name = f"{uuid.uuid4().hex[:16]}.{ext}"
        with open(os.path.join(IMG_DIR, name), "wb") as fp:
            fp.write(raw)
        return f"/api/import/img/{name}"
    except Exception:
        return data_url          # 실패하면 원래대로


@app.route("/api/import/img/<path:name>", methods=["GET"])
def import_img(name):
    """가져오기로 만들어진 배경(벡터 SVG 또는 그림)"""
    if not re.fullmatch(r"[0-9a-f]{8,32}\.(png|jpg|svg)", name or ""):
        return jsonify({"error": "잘못된 이름"}), 400
    path = os.path.join(IMG_DIR, name)
    if not os.path.exists(path):
        return jsonify({"error": "없는 파일"}), 404

    if name.endswith(".svg"):
        with open(path, "rb") as fp:
            raw = fp.read()
        # 벡터는 글자 정보가 많아 압축 효과가 크다 (약 1/8)
        if "gzip" in (request.headers.get("Accept-Encoding") or "").lower():
            import gzip as _gz
            body = _gz.compress(raw, 6)
            resp = Response(body, mimetype="image/svg+xml")
            resp.headers["Content-Encoding"] = "gzip"
        else:
            resp = Response(raw, mimetype="image/svg+xml")
        resp.headers["Content-Length"] = str(len(resp.get_data()))
        resp.headers["Cache-Control"] = "public, max-age=31536000"
        return resp

    resp = send_from_directory(IMG_DIR, name)
    resp.headers["Cache-Control"] = "public, max-age=31536000"
    return resp


# ── 슬라이스 메모리 캐시 ────────────────────────────────────────────────
# 이미 만들어 둔 노트를 여는 경로(GET .s{n}.gz)는 읽기 전용 gzip 파일을 그대로
# 흘려보낸다. 같은 파일을 열 때마다 디스크를 때리지 않도록 최근 것만 담아 둔다.
# 키에 ETag(수정시각+크기)를 넣으므로, 저장으로 파일이 바뀌면 캐시는 자동으로
# 빗나가고 새 내용을 읽는다 — 오래된 본문이 나갈 수 없다.
try:                       # (_imp_env_int 는 이 파일 아래쪽에서 정의된다)
    _SLICE_CACHE_MAX = max(16, int(os.environ.get("SDY_SLICE_CACHE_MAX", "512")))
except (TypeError, ValueError):
    _SLICE_CACHE_MAX = 512                                    # 슬라이스 개수
_slice_cache = {}
_slice_cache_lock = threading.Lock()


def _slice_cache_get(path, etag):
    key = (path, etag)
    with _slice_cache_lock:
        hit = _slice_cache.get(key)
        if hit is not None:
            _slice_cache.pop(key, None)     # LRU: 최근 쓴 것을 뒤로
            _slice_cache[key] = hit
            return hit
    with open(path, "rb") as fp:
        body = fp.read()
    with _slice_cache_lock:
        _slice_cache[key] = body
        while len(_slice_cache) > _SLICE_CACHE_MAX:
            _slice_cache.pop(next(iter(_slice_cache)), None)
    return body


@app.route("/api/import/docfile/<jid>", methods=["GET", "POST"])
def import_docfile(jid):
    """대용량 문서 본문 서버 보관소.

    브라우저 localStorage(약 5MB) 를 넘기는 문서는 기기 대신 여기에 두고,
    프런트는 참조 번호만 기억한다. 편집분도 여기로 저장된다.
    """
    import gzip as _gz
    if not re.fullmatch(r"[0-9a-zA-Z_\-]{4,40}", jid or ""):
        return jsonify({"error": "잘못된 이름"}), 400
    path = os.path.join(DOCS_DIR, f"{jid}.json.gz")
    if request.method == "POST":
        d = request.get_json(silent=True) or {}
        pages = d.get("pages")
        if not isinstance(pages, list):
            return jsonify({"ok": False, "error": "pages 없음"}), 400
        # 14.4 · 슬라이스만 갱신 (번역/편집분을 원본 문서에 저장).
        # 전체 pages 를 다시 올리면 대용량 문서에서 타임아웃·유실이 난다.
        s0_in = d.get("from")
        if s0_in is not None:
            try:
                s0 = int(s0_in)
            except Exception:
                return jsonify({"ok": False, "error": "from 오류"}), 400
            if s0 < 0 or (s0 % IMP_SLICE) != 0:
                return jsonify({"ok": False, "error": "from 오류"}), 400
            try:
                total_n = int(d.get("total") or (s0 + len(pages)))
                chunk = pages[:IMP_SLICE]
                sp = os.path.join(DOCS_DIR, f"{jid}.s{s0}.gz.tmp")
                with _gz.open(sp, "wt", encoding="utf-8") as fp:
                    json.dump({"ok": True, "pages": chunk, "total": total_n},
                              fp, ensure_ascii=False)
                os.replace(sp, os.path.join(DOCS_DIR, f"{jid}.s{s0}.gz"))
                if os.path.exists(path):
                    try:
                        with _gz.open(path, "rt", encoding="utf-8") as fp:
                            data_ = json.load(fp)
                        allp = data_.get("pages") if isinstance(data_, dict) else None
                        if isinstance(allp, list):
                            for i, pg in enumerate(chunk):
                                idx = s0 + i
                                if idx < len(allp):
                                    allp[idx] = pg
                                elif idx == len(allp):
                                    allp.append(pg)
                            tmpf = "%s.tmp.%s" % (path, uuid.uuid4().hex[:8])
                            with _gz.open(tmpf, "wt", encoding="utf-8") as fp:
                                json.dump({"pages": allp,
                                           "sizePreset": d.get("sizePreset")
                                           or data_.get("sizePreset", "a4_portrait")},
                                          fp, ensure_ascii=False)
                            os.replace(tmpf, path)
                    except Exception:
                        pass
                mp = os.path.join(DOCS_DIR, f"{jid}.meta.json.tmp")
                with open(mp, "w", encoding="utf-8") as fp:
                    json.dump({"total": total_n,
                               "sizePreset": d.get("sizePreset", "a4_portrait"),
                               "version": time.time()}, fp)
                os.replace(mp, os.path.join(DOCS_DIR, f"{jid}.meta.json"))
                try:
                    _publish_live("notes", "")
                except Exception:
                    pass
                return jsonify({"ok": True, "slice": s0, "version": time.time()})
            except Exception as e:
                return jsonify({"ok": False, "error": str(e)}), 500
        try:
            total_n = len(pages)
            for s0 in range(0, total_n, IMP_SLICE):
                sp = os.path.join(DOCS_DIR, f"{jid}.s{s0}.gz.tmp")
                with _gz.open(sp, "wt", encoding="utf-8") as fp:
                    json.dump({"ok": True, "pages": pages[s0:s0 + IMP_SLICE],
                               "total": total_n}, fp, ensure_ascii=False)
                os.replace(sp, os.path.join(DOCS_DIR, f"{jid}.s{s0}.gz"))
            for s0 in range(((total_n // IMP_SLICE) + 1) * IMP_SLICE,
                            ((total_n // IMP_SLICE) + 5) * IMP_SLICE, IMP_SLICE):
                oldp = os.path.join(DOCS_DIR, f"{jid}.s{s0}.gz")
                if os.path.exists(oldp):
                    os.remove(oldp)
            mp = os.path.join(DOCS_DIR, f"{jid}.meta.json.tmp")
            with open(mp, "w", encoding="utf-8") as fp:
                json.dump({"total": total_n,
                           "sizePreset": d.get("sizePreset", "a4_portrait"),
                           "version": time.time()}, fp)
            os.replace(mp, os.path.join(DOCS_DIR, f"{jid}.meta.json"))
            tmp = "%s.tmp.%s" % (path, uuid.uuid4().hex[:8])
            with _gz.open(tmp, "wt", encoding="utf-8") as fp:
                json.dump({"pages": pages,
                           "sizePreset": d.get("sizePreset", "a4_portrait")},
                          fp, ensure_ascii=False)
            os.replace(tmp, path)
            return jsonify({"ok": True})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500
    meta_path = os.path.join(DOCS_DIR, f"{jid}.meta.json")
    slice0 = os.path.join(DOCS_DIR, f"{jid}.s0.gz")
    if not (os.path.exists(path) or os.path.exists(meta_path)
            or os.path.exists(slice0)):
        return jsonify({"ok": False, "error": "없는 문서"}), 404
    if request.args.get("meta") == "1":
        try:
            with open(meta_path, encoding="utf-8") as fp:
                m = json.load(fp)
            return jsonify({"ok": True, "version": m.get("version", 0),
                            "total": m.get("total", 0)})
        except Exception:
            pass
        try:   # 옛 단일 파일 폴백
            st = os.stat(path)
            with _gz.open(path, "rt", encoding="utf-8") as fp:
                total = len(json.load(fp).get("pages", []))
            return jsonify({"ok": True, "version": st.st_mtime, "total": total})
        except Exception:
            return jsonify({"ok": False, "error": "읽기 실패"}), 500
    fr = request.args.get("from", type=int)
    to = request.args.get("to", type=int)
    if fr is not None:
        fr = max(0, fr)
        to = fr + IMP_SLICE if to is None else max(fr, min(fr + IMP_SLICE, to))
        # 정렬된 슬라이스면 파싱 없이 압축 그대로 스트리밍 (메모리 최소화)
        if fr % IMP_SLICE == 0 and (to - fr) <= IMP_SLICE:
            sp = os.path.join(DOCS_DIR, f"{jid}.s{fr}.gz")
            if os.path.exists(sp):
                # 14.29.4 · 노트 여는 속도.
                #   같은 슬라이스를 열 때마다 디스크에서 다시 읽었다. 이제
                #   ① 조건부 요청(ETag) → 안 바뀌었으면 304 만 보내고 본문 0바이트
                #   ② 최근 슬라이스는 메모리에 캐시 → 디스크 I/O 자체를 건너뜀
                #   ETag 는 (수정시각, 크기) 라 저장(POST)하면 즉시 달라진다.
                st = os.stat(sp)
                etag = '"%x-%x"' % (int(st.st_mtime_ns), st.st_size)
                inm = request.headers.get("If-None-Match") or ""
                if etag in [t.strip() for t in inm.split(",")]:
                    resp = Response(status=304)
                    resp.headers["ETag"] = etag
                    resp.headers["Cache-Control"] = "private, max-age=0, must-revalidate"
                    return resp
                body = _slice_cache_get(sp, etag)
                resp = Response(body, mimetype="application/json")
                resp.headers["Content-Encoding"] = "gzip"
                resp.headers["ETag"] = etag
                resp.headers["Cache-Control"] = "private, max-age=0, must-revalidate"
                return resp
        # 폴백: 옛 단일 파일 또는 비정렬 범위
        if os.path.exists(path):
            with _gz.open(path, "rt", encoding="utf-8") as fp:
                data_ = json.load(fp)
            pages_all = data_.get("pages", [])
            total = len(pages_all)
            fr2 = max(0, min(fr, total))
            to2 = max(fr2, min(total, to))
            resp = jsonify({"ok": True, "pages": pages_all[fr2:to2],
                            "total": total,
                            "sizePreset": data_.get("sizePreset", "a4_portrait")})
            resp.headers["Cache-Control"] = "no-store"
            return resp
        return jsonify({"ok": False, "error": "없는 문서"}), 404
    # 디스크에 이미 gzip → 지원하면 압축 상태 그대로 스트리밍(대용량 절약)
    if "gzip" in (request.headers.get("Accept-Encoding") or "").lower():
        with open(path, "rb") as fp:
            body = fp.read()
        resp = Response(body, mimetype="application/json")
        resp.headers["Content-Encoding"] = "gzip"
    else:
        with _gz.open(path, "rt", encoding="utf-8") as fp:
            data_ = json.load(fp)
        resp = jsonify({"ok": True, "pages": data_.get("pages", []),
                        "sizePreset": data_.get("sizePreset", "a4_portrait")})
    resp.headers["Cache-Control"] = "no-store"
    return resp


# 수식 전용 글꼴. 이 목록은 의도적으로 보수적이다.
#
# Computer Modern Roman(cmr), STIXGeneral, XITS 같은 글꼴은 논문에서
# 본문에도 쓰인다. 예전에는 이름에 그 문자열이 있다는 이유만으로 span 전체를
# 수식 사진으로 잘라, 식 옆의 "where", "for" 같은 설명까지 이미지가 됐다.
# 전용 기호/수학 글꼴만 즉시 수식으로 인정하고, 겸용 글꼴은 아래의 기호·문맥
# 점수를 통과할 때만 수식으로 처리한다.
MATH_FONTS = (
    "cmmi", "cmsy", "cmex", "msam", "msbm", "mathjax",
    "cambriamath", "lmmath", "rsfs", "eufm", "wasy", "symbol",
    # 10.1 · pdflatex txfonts/newtx (rtxmi·rtxbmi·txsy·txex 계열)
    "txmi", "txbmi", "txsy", "txex",
)
MATH_AMBIG_FONTS = (
    "cmr", "stixgeneral", "stix", "xits", "euclid", "texgyre",
)


def _span_text(sp):
    """rawdict/dict 양쪽에서 span 글자를 안전하게 얻는다."""
    if sp.get("chars") is not None:
        return "".join(ch.get("c", "") for ch in sp.get("chars", []))
    return sp.get("text") or ""

# 물리·수학 논문은 일반 Times/Arial 글꼴로 식을 심는 경우가 많다.
# 글꼴 이름만 믿지 않고, 기호 밀도·첨자·관계식 패턴도 함께 본다.
_MATH_SIGNS = set("=≠≈≃≅≤≥±∓×÷·∂∇∫∮∑∏√∞∝∈∉⊂⊃∪∩→←↔⇒⇔∀∃∇∆∥⊥〈〉⟨⟩")
_GREEK_RE = re.compile(r"[αβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ]")


def _formula_score(text):
    t = (text or "").strip()
    if not t:
        return 0
    signs = sum(ch in _MATH_SIGNS for ch in t)
    greek = len(_GREEK_RE.findall(t))
    # e.g. E = mc², Ĥψ = Eψ, d²x/dt², k_B T
    score = signs * 2 + greek * 2
    if re.search(r"[A-Za-zα-ωΑ-Ω]\s*[=≈≃≠≤≥]\s*", t): score += 4
    if re.search(r"(?:d|∂)\s*[A-Za-zα-ω]|[A-Za-zα-ω]\s*/\s*(?:d|∂)", t): score += 3
    if re.search(r"[A-Za-zα-ω]\s*[_^]\s*[A-Za-z0-9α-ω]", t): score += 2
    if re.search(r"\b(?:sin|cos|tan|exp|log|lim|det|Tr|curl|grad|div)\b", t): score += 2
    return score


_MATH_WORDS = {
    "sin", "cos", "tan", "cot", "sec", "csc", "sinh", "cosh", "tanh", "coth",
    "arcsin", "arccos", "arctan", "arg", "exp", "log", "ln", "lg",
    "lim", "limsup", "liminf", "det", "tr", "trace", "curl", "grad", "div",
    "max", "min", "sup", "inf", "mod", "rank", "diag", "eff", "tot", "ext", "avg", "rms", "rmsnorm",
    "ric", "scal", "vol", "diam", "inj", "hess", "deg", "sgn", "sign", "span", "ker", "im", "re",
    "hom", "aut", "gl", "sl", "so", "su", "dim", "codim", "supp", "const", "var", "cov",
    "erf", "sinc", "pr", "opt", "attn", "top", "topk", "mlp", "ffn", "gelu", "relu",
    "softmax", "proj", "qkv", "mqa", "gqa", "kv", "hbm", "hbf", "gpu", "sram", "tsv", "ecc",
    "base", "out", "gate", "route", "gcd", "lcm"
}

_COMMON_PROSE = {
    "in", "if", "as", "to", "then", "where", "let", "we", "have", "and", "for", "with", "by", "is", "are",
    "of", "that", "this", "such", "each", "there", "from", "which", "also", "shows", "under", "above",
    "below", "given", "thus", "hence", "so", "holds", "true", "equal", "obtain", "proves", "note", "finally",
    "define", "using", "proof", "theorem", "lemma", "corollary", "proposition", "remark", "assume", "assuming",
    "suppose", "follows", "particular", "section", "recall", "fact", "indeed", "clearly", "similarly",
    "furthermore", "moreover", "therefore", "because", "since", "seen", "easily", "integrating", "integrate",
    "consider", "taking", "definition", "identity", "satisfies", "denote", "denotes", "denoting", "satisfy",
}


# ─────────────────────────────────────────────────────────────
#  9.0 · 인용 번호는 '언제나 텍스트'
# ─────────────────────────────────────────────────────────────
#  논문 본문의 [12] · [3–5] · [1, 4, 9] · (2024) · 위첨자 12,13 은
#  참고문헌 번호지 수식이 아니다. 예전에는 글꼴·크기·베이스라인에 따라
#  어떤 것은 LaTeX 로, 어떤 것은 텍스트로 갈렸다(= 사용자가 말한 '들쭉날쭉').
#  판정을 이 한 곳에 모아 두고, 수식 판정기 전부가 여기를 먼저 통과하게 해서
#  인용 번호는 예외 없이 일반 글자 상자로 남긴다.
_CITE_BRACKET = re.compile(
    r"^[\[\(]\s*\d{1,4}(?:\s*[-–—,;]\s*\d{1,4})*\s*[\]\)][.,;:]?$")
_CITE_BARE = re.compile(r"^\d{1,3}(?:\s*[-–—,]\s*\d{1,3})*[.,;:]?$")
_CITE_SUPCHARS = set("⁰¹²³⁴⁵⁶⁷⁸⁹⁻⁼⁽⁾,-–—")
_SUP_DIGITSET = set("⁰¹²³⁴⁵⁶⁷⁸⁹")


def _is_citation_token(text, superscript=False):
    """이 토막이 참고문헌/각주 번호인가?  (수식이면 안 되는 것들)"""
    t = (text or "").strip()
    if not t or len(t) > 24:
        return False
    if _CITE_BRACKET.match(t):
        return True
    # 유니코드 위첨자만으로 이뤄진 것 → ¹²·⁵⁻⁷
    if t and all(ch in _CITE_SUPCHARS for ch in t) and any(ch in _SUP_DIGITSET for ch in t):
        return True
    # 위첨자로 조판된 맨숫자 → 12, 3–5, 1,2
    if superscript and _CITE_BARE.match(t):
        return True
    return False


def _line_is_citation_only(text):
    """줄 전체가 인용 번호 뭉치인지 (예: '[12, 15–18]')"""
    t = (text or "").strip()
    if not t:
        return False
    parts = [p for p in re.split(r"\s+", t) if p]
    return bool(parts) and all(_is_citation_token(p) for p in parts)


def _is_formula_only(text):
    """수식 자체인지 판정한다.

    물리 논문의 설명문까지 그림으로 굳어 버리지 않도록, 기호 점수만 보지 않고
    일반 단어(prose)가 섞였는지 확인한다. 예: 'where E is energy'는 제외하고
    'E = mc²', 'iℏ∂ψ/∂t = Hψ'만 통과시킨다.
    """
    t=(text or "").strip()
    if not t or _formula_score(t) < 5 or len(t) > 180:
        return False
    words=re.findall(r"[A-Za-z]{2,}", t)
    prose=[w for w in words if w.lower() not in _MATH_WORDS]
    # 2글자 이상 일반 단어가 둘 이상이면 문장으로 간주한다.
    return len(prose) <= 1


def _is_math_span(sp):
    """span 하나가 정말 '수식만' 담고 있는지 보수적으로 판정한다.

    수식 전용 글꼴은 그대로 인정한다. 본문과 수식 양쪽에 쓰이는 CMR/STIX/XITS
    계열은 글꼴 이름만으로 인정하지 않고, span 자체가 수식 모양일 때만 인정한다.
    일반 글꼴 역시 강한 수식 패턴이 있을 때만 이미지화한다.
    """
    name = (sp.get("font") or "").lower().replace("-", "").replace(" ", "")
    text = _span_text(sp)
    # [12], [3–5], (2024), ¹² 같은 인용/각주 표시는 절대 수식이 되지 않는다. (9.0)
    if _is_citation_token(text, superscript=bool(sp.get("_sup"))):
        return False
    if any(k in name for k in MATH_FONTS):
        # 전용 글꼴이어도 긴 일반 단어(where, energy...)는 본문으로 남긴다.
        # 짧은 변수(E, mc, kB)와 함수(sin 등)는 수식 조각으로 인정한다.
        words = [w for w in re.findall(r"[A-Za-z]{2,}", text)
                 if w.lower() not in _MATH_WORDS]
        if len(words) > 1 or any(len(w) >= 4 for w in words):
            return False
        return bool(text.strip())
    if any(k in name for k in MATH_AMBIG_FONTS):
        return _is_formula_only(text) and _formula_score(text) >= 5
    if _is_formula_only(text):
        return True
    # '='·'≤' 같은 연산자가 별도 span으로 분리된 PDF. 그 조각도 수식 영역에
    # 포함해야 좌우의 변수 span과 병합되어 식 한 장이 된다.
    compact=(text or "").strip()
    return bool(compact and not re.search(r"[A-Za-z가-힣]", compact)
                and _formula_score(compact) >= 2)


def _spans_to_clean_latex(spans):
    """PDF spans → 고품질 LaTeX 수식 문자열 (13.4 · 위첨자/아래첨자/연산자/기호 정밀 보존)."""
    szs = [float(sp.get("size") or 10) for sp in spans if any((ch.get("c") or "").strip() for ch in sp.get("chars", []))]
    if not szs:
        szs = [float(sp.get("size") or 10) for sp in spans if sp.get("text", "").strip()]
    if not szs:
        return ""
    base_sz = sorted(szs)[len(szs) // 2]

    origins = []
    for sp in spans:
        for ch in sp.get("chars", []):
            if (ch.get("c") or "").strip():
                if ch.get("origin"):
                    origins.append(ch["origin"][1])
                elif ch.get("bbox"):
                    origins.append(ch["bbox"][3])
    base_y = sorted(origins)[len(origins) // 2] if origins else 0

    tokens = []
    for sp in spans:
        sz = float(sp.get("size") or 10)
        font = (sp.get("font") or "").lower()
        chars = sp.get("chars", [])
        if not chars:
            t = sp.get("text", "")
            if t:
                tokens.append({"t": t, "sz": sz, "y": base_y, "font": font, "mode": "normal"})
            continue
        for ch in chars:
            c = ch.get("c", "")
            if not c:
                continue
            oy = ch["origin"][1] if ch.get("origin") else ch["bbox"][3]
            mode = "normal"
            if sz < base_sz * 0.88 or abs(oy - base_y) >= base_sz * 0.14:
                if oy < base_y - base_sz * 0.10:
                    mode = "sup"
                elif oy > base_y + base_sz * 0.10:
                    mode = "sub"
            tokens.append({"t": c, "sz": sz, "y": oy, "font": font, "mode": mode})

    groups = []
    for tok in tokens:
        if groups and groups[-1]["mode"] == tok["mode"] and tok["mode"] in ("sub", "sup"):
            groups[-1]["t"] += tok["t"]
        elif groups and groups[-1]["mode"] == tok["mode"] == "normal" and tok["t"].isalnum() and groups[-1]["t"].isalnum():
            groups[-1]["t"] += tok["t"]
        else:
            groups.append(dict(tok))

    res = []
    for g in groups:
        t = g["t"]
        mode = g["mode"]
        for ch, sym in _LATEX_GREEK.items():
            t = t.replace(ch, f" \\{sym} ")
        for ch, sym in _LATEX_SYMBOLS.items():
            t = t.replace(ch, f" {sym} ")
        t = t.replace("−", "-").replace("¯g", r"\bar{g}").replace("¯", r"\bar")

        if mode == "sub":
            clean_sub = t.strip()
            if clean_sub.lower() in _MATH_WORDS:
                res.append(r"_{\text{" + clean_sub + r"}}")
            else:
                res.append("_{" + clean_sub + "}")
        elif mode == "sup":
            res.append("^{" + t.strip() + "}")
        else:
            if t.lower() in _MATH_WORDS and len(t) >= 2:
                if t.lower() in {"sin", "cos", "tan", "exp", "log", "ln", "lim", "det", "dim", "min", "max", "sup", "inf"}:
                    res.append(f"\\{t.lower()} ")
                else:
                    res.append(r"\text{" + t + r"} ")
            else:
                res.append(t)

    out = "".join(res)
    out = re.sub(r"\s+", " ", out).strip()
    return out


def _is_display_formula_line(ln):
    """한 줄 전체를 LaTeX로 쓸지, 전체를 텍스트로 쓸지 결정한다. (13.4)

    부분 변환은 하지 않는다. 설명 문장이 섞인 줄은 수식이 있어도 전부 텍스트,
    독립된 식으로 확실한 줄만 전부 LaTeX가 된다.
    """
    spans = ln.get("spans", [])
    text = "".join(_span_text(sp) for sp in spans).strip()
    if not text or len(text) > 240:
        return False
    # 9.0 · 인용 번호 줄은 언제나 텍스트
    if _line_is_citation_only(text):
        return False
    # 9.0 · 식 번호만 떼어낸 조각([12] 처럼 보이는 (3))도 텍스트로
    if _CITE_BRACKET.match(text):
        return False
    words = []
    for sp in spans:
        st = _span_text(sp)
        name = (sp.get("font") or "").lower().replace("-", "").replace(" ", "")
        if any(k in name for k in MATH_FONTS):
            continue
        words.extend(re.findall(r"[A-Za-z]{2,}", st))
    prose = [w for w in words if w.lower() in _COMMON_PROSE
             or (w.lower() not in _MATH_WORDS and len(w) >= 3)]
    # 설명 단어가 하나라도 있으면 수식 일부가 있어도 줄 전체를 텍스트로 유지
    # 14.3 · if/of/in 같은 2글자 산문도 수식 줄로 만들지 않는다.
    if prose:
        return False
    score = _formula_score(text)
    if score >= 4:
        return True
    total = math_n = 0
    for sp in spans:
        t = _span_text(sp).strip()
        if not t: continue
        total += len(t)
        name = (sp.get("font") or "").lower().replace("-", "").replace(" ", "")
        if any(k in name for k in MATH_FONTS): math_n += len(t)
    ratio = (math_n / total) if total else 0
    return ratio >= 0.35 or any(sym in text for sym in ["≥", "≤", "=", "→", "∈", "λ", "ϵ", "α", "δ", "η", "γ", "θ"])


def _display_region_of_line(ln, page=None, gtables=None, rules=None):
    """디스플레이 수식 줄에서 '식 본체'만의 영역을 돌려준다. (13.4)

    논문 수식 줄은 대개  ─  ψ(x) = A e^{ikx}          (3)  ─  처럼
    오른쪽 끝(가끔 왼쪽 끝)에 식 번호가 멀찍이 떨어져 붙어 있다.
    여기서 식 번호를 떼어내고, 식 본체는 고품질 KaTeX 수식으로 보존한다.

    반환: (x0, y0, x1, y1, text, size) 또는 None
    """
    spans = ln.get("spans", [])
    if not spans:
        return None
    toks = []
    for sp in spans:
        chars = sp.get("chars") or []
        size = float(sp.get("size") or 10)
        if not chars:
            t = _span_text(sp).strip()
            bb = sp.get("bbox")
            if t and bb:
                toks.append({"t": t, "b": list(bb), "sz": size})
            continue
        cur = []

        def _flush():
            if not cur:
                return
            txt = "".join(c.get("c") or "" for c in cur).strip()
            bxs = [c.get("bbox") for c in cur if c.get("bbox")]
            if txt and bxs:
                toks.append({"t": txt, "sz": size,
                             "b": [min(b[0] for b in bxs), min(b[1] for b in bxs),
                                   max(b[2] for b in bxs), max(b[3] for b in bxs)]})
            cur.clear()

        for ch in chars:
            if not (ch.get("c") or "").strip():
                _flush()
            else:
                cur.append(ch)
        _flush()
    if not toks:
        return None
    toks.sort(key=lambda u: u["b"][0])

    heights = sorted(max(1.0, u["b"][3] - u["b"][1]) for u in toks)
    hmed = heights[len(heights) // 2]

    def _eqnum(u):
        return bool(_CITE_BRACKET.match(u["t"]) or re.fullmatch(r"[\[\(]\s*[\dA-Za-z.\-]{1,8}\s*[\]\)]", u["t"]))

    # 오른쪽 끝: 큰 공백 뒤에 오는 (3) 꼴을 떼어낸다
    while len(toks) >= 2 and _eqnum(toks[-1]):
        gap = toks[-1]["b"][0] - toks[-2]["b"][2]
        if gap < max(6.0, hmed * 1.6):
            break
        toks.pop()
    # 왼쪽 끝도 같은 규칙
    while len(toks) >= 2 and _eqnum(toks[0]):
        gap = toks[1]["b"][0] - toks[0]["b"][2]
        if gap < max(6.0, hmed * 1.6):
            break
        toks.pop(0)
    if not toks:
        return None

    tex = _spans_to_clean_latex(spans)
    if not tex:
        tex = " ".join(u["t"] for u in toks).strip()
    if not tex:
        return None
    szs = sorted(u["sz"] for u in toks)
    x0 = min(u["b"][0] for u in toks); y0 = min(u["b"][1] for u in toks)
    x1 = max(u["b"][2] for u in toks); y1 = max(u["b"][3] for u in toks)
    # 14.0 · 줄 단위 1D 조립은 중첩 분수를 못 살린다. 2D 복원기가 되면 그걸 쓴다.
    if page is not None:
        try:
            tex2 = region_to_latex(page.parent, page, (x0, y0, x1, y1), gtables, rules)
            if tex2 and _latex_is_sane(tex2) and not _tex_is_figure_junk(tex2):
                tex = tex2
        except Exception:
            pass
    return (x0, y0, x1, y1, tex, szs[len(szs) // 2])


def _line_math_regions(ln):
    """한 줄을 토큰 단위로 나눠 수식 덩어리를 찾는다.

    PDF는 한 식 안에서도 변수는 CMMI, 등호는 CMR, 숫자는 Times처럼 여러
    span으로 쪼갠다. 예전에는 일부 span만 수식으로 빠지고 나머지가 글상자로
    남아 두 겹이 됐다. 이제 '='·그리스문자·수식 글꼴을 앵커로 삼고, 좌우의
    짧은 변수/숫자/첨자를 함께 한 LaTeX 영역으로 묶는다.
    """
    spans=ln.get("spans", [])
    all_sizes=[float(sp.get("size") or 10) for sp in spans if _span_text(sp).strip()]
    if not all_sizes: return [], set()
    ss=sorted(all_sizes); base_size=ss[len(ss)//2]
    origins=[]
    for sp in spans:
        for ch in sp.get("chars", []):
            if (ch.get("c") or "").strip():
                origins.append((ch.get("origin") or (0,ch.get("bbox",[0,0,0,0])[3]))[1])
    origins.sort(); base_y=origins[len(origins)//2] if origins else 0

    units=[]
    for sp in spans:
        chars=sp.get("chars") or []
        if not chars:
            txt=_span_text(sp).strip(); bb=sp.get("bbox")
            if txt and bb: units.append({"text":txt,"bbox":bb,"sp":sp,"chars":[]})
            continue
        cur=[]
        def flush():
            if not cur: return
            txt="".join(c.get("c") or "" for c in cur).strip()
            boxes=[c.get("bbox") for c in cur if c.get("bbox")]
            if txt and boxes:
                units.append({"text":txt,
                    "bbox":[min(b[0] for b in boxes),min(b[1] for b in boxes),
                            max(b[2] for b in boxes),max(b[3] for b in boxes)],
                    "sp":sp,"chars":list(cur)})
            cur.clear()
        for ch in chars:
            if not (ch.get("c") or "").strip(): flush()
            else: cur.append(ch)
        flush()
    if not units: return [],set()
    units.sort(key=lambda u:(u["bbox"][0],u["bbox"][1]))

    def _is_sup(u):
        """이 토막이 위첨자로 조판됐는가 (인용 번호 판정용)"""
        sz=float(u["sp"].get("size") or base_size)
        if sz>=base_size*0.9: return False
        oy=[(c.get("origin") or (0,c.get("bbox",[0,0,0,0])[3]))[1] for c in u["chars"]]
        if not oy: return False
        return (sum(oy)/len(oy)) < base_y-base_size*0.12

    def citation(t,u=None):
        return _is_citation_token(t, superscript=bool(u and _is_sup(u)))
    prose_short={"is","in","of","to","as","or","and","the","for","by","at","on","if","we","a","an"}
    def attrs(u):
        t=u["text"].strip(); sp=u["sp"]
        pseudo=dict(sp); pseudo["text"]=t
        pseudo["_sup"]=_is_sup(u)
        if u["chars"]: pseudo["chars"]=u["chars"]
        if t.lower() in prose_short: return False,False
        # 9.0 · 인용 번호는 앵커도 이웃도 될 수 없다 → 수식에 절대 안 딸려간다
        if citation(t,u): return False,False
        anchor=_is_math_span(pseudo) or _formula_score(t)>=2
        words=re.findall(r"[A-Za-z]{2,}",t)
        short_word=all((len(w)<=3 or w.lower() in _MATH_WORDS) and w.lower() not in prose_short for w in words)
        simple=bool(re.fullmatch(r"[A-Za-z0-9α-ωΑ-Ω.,+\-*/=<>:;()_^{}\\|]+",t))
        neighbor=(anchor or (simple and short_word and len(t)<=20)
                   or t.lower() in _MATH_WORDS)
        return anchor,neighbor
    flags=[attrs(u) for u in units]
    ranges=[]
    for i,(anchor,_) in enumerate(flags):
        if not anchor: continue
        a=b=i
        while a>0 and flags[a-1][1]:
            gap=units[a]["bbox"][0]-units[a-1]["bbox"][2]
            h=max(units[a]["bbox"][3]-units[a]["bbox"][1],units[a-1]["bbox"][3]-units[a-1]["bbox"][1],1)
            if gap>max(4.0,h*1.15): break
            a-=1
        while b+1<len(units) and flags[b+1][1]:
            gap=units[b+1]["bbox"][0]-units[b]["bbox"][2]
            h=max(units[b]["bbox"][3]-units[b]["bbox"][1],units[b+1]["bbox"][3]-units[b+1]["bbox"][1],1)
            if gap>max(4.0,h*1.15): break
            b+=1
        if ranges and a<=ranges[-1][1]+1: ranges[-1]=(ranges[-1][0],max(ranges[-1][1],b))
        else: ranges.append((a,b))

    regs=[]; skip=set()
    for a,b in ranges:
        part=units[a:b+1]
        raw=" ".join(u["text"] for u in part).strip()
        # 기호 하나뿐인 오탐은 버리되 변수 하나(CMMI)는 허용
        if not raw: continue
        pieces=[]
        for u in part:
            txt=u["text"]
            sp=u["sp"]; sz=float(sp.get("size") or base_size)
            oy=[]
            for ch in u["chars"]:
                oy.append((ch.get("origin") or (0,ch.get("bbox",[0,0,0,0])[3]))[1])
                skip.add(id(ch))
            y=sum(oy)/len(oy) if oy else base_y
            if sz<base_size*.84 or (sp.get("flags",0)&1):
                if y<base_y-base_size*.12: txt="^{"+txt+"}"
                elif y>base_y+base_size*.12: txt="_{"+txt+"}"
            pieces.append(txt)
        boxes=[u["bbox"] for u in part]
        regs.append({"x0":min(x[0] for x in boxes),"y0":min(x[1] for x in boxes),
                     "x1":max(x[2] for x in boxes),"y1":max(x[3] for x in boxes),
                     "display":False,"text":" ".join(pieces),"size":base_size})
    return regs,skip


def _math_ratio(blk):
    """블록에서 수식 글꼴이 차지하는 글자 비율 (dict/rawdict 공용)"""
    tot = mat = 0
    for ln in blk.get("lines", []):
        for sp in ln.get("spans", []):
            if sp.get("chars") is not None:                 # rawdict
                t = "".join(ch.get("c", "") for ch in sp["chars"]).strip()
            else:                                           # dict
                t = (sp.get("text") or "").strip()
            if not t:
                continue
            tot += len(t)
            if _is_math_span(sp):
                mat += len(t)
    return (mat / tot) if tot else 0.0


def _imp_uid(prefix):
    return f"{prefix}_{uuid.uuid4().hex[:9]}"


def _px(v):
    return int(round(v))


def _img_data_url(raw, ext="png", keep_big=False):
    """이미지 바이트 → data URL.

    keep_big=True 면 그림(도표·사진)을 또렷하게 보이도록
    큰 해상도를 유지하고 무손실(PNG)로 담는다.
    """
    try:
        im = Image.open(io.BytesIO(raw))
        if im.mode in ("P", "LA", "CMYK"):
            im = im.convert("RGBA" if "A" in im.mode else "RGB")
        limit = 3200 if keep_big else 1600
        if max(im.size) > limit:
            r = limit / max(im.size)
            im = im.resize((max(1, int(im.width * r)), max(1, int(im.height * r))),
                           Image.LANCZOS)
        out = io.BytesIO()
        if keep_big and im.mode != "RGBA":
            # 색이 적은 도표는 PNG 가 작고 선명하다.
            # 사진처럼 색이 많으면 PNG 가 지나치게 커지므로 고품질 JPEG.
            try:
                colors = im.convert("RGB").getcolors(4096)
            except Exception:
                colors = None
            if colors is not None:            # 색이 4096 가지 이하 = 도표
                rgb = im.convert("RGB")
                n = len(colors)
                if n > 256:
                    # 색을 256 가지로 정리 → 크기가 크게 줄고 선은 그대로 또렷
                    rgb = rgb.quantize(colors=256, method=Image.MEDIANCUT,
                                       dither=Image.Dither.NONE)
                rgb.save(out, "PNG", optimize=True)
                mime = "image/png"
            else:
                im.convert("RGB").save(out, "JPEG", quality=88, optimize=True,
                                       subsampling=0)
                mime = "image/jpeg"
        elif keep_big:
            im.save(out, "PNG", optimize=True)
            mime = "image/png"
        elif im.mode == "RGBA":
            im.save(out, "PNG", optimize=True)
            mime = "image/png"
        else:
            im.convert("RGB").save(out, "JPEG", quality=82, optimize=True)
            mime = "image/jpeg"
        b = base64.b64encode(out.getvalue()).decode()
        return f"data:{mime};base64,{b}"
    except Exception:
        return None


# ─────────────────────────────────────────────────────────────
#  PDF 단어 단위 추출
#
#  핵심 원칙:
#   · 배경(그래프·도형·표 테두리·색·선 굵기) = SVG/래스터로 원본 그대로 유지
#   · 글자 = 띄어쓰기 단위 '단어별' 상자를 만들어 정확한 절대 위치에 배치
#     (문단으로 묶을 때 생기던 "Core Core Core Core" 간격 오류 해결)
#   · 단어끼리 절대 겹치지 않게 패딩을 인접 단어/줄 간격 안으로 제한
#   · 원본 정렬(왼쪽/가운데/오른쪽)을 판별해 그대로 적용
#   · 표: 테두리는 배경에 남기고, 안쪽 글자만 단어 단위로 추출 → 중복 없음
# ─────────────────────────────────────────────────────────────

WORD_GAP_RATIO = 0.18   # 글자 사이 간격이 글자 크기의 이 비율을 넘으면 띄어쓰기
SUP_SUB_DY = 0.6        # 베이스라인이 이 비율(글자 크기 대비) 이상 다르면 다른 단어


# ═══════════════════════════════════════════════════════════
#  9.3 · 큰 수식 복원기 (oversized math reconstruction)
# ═══════════════════════════════════════════════════════════
#  물리 논문에서 '큰' 구조 — 키 큰 적분/시그마, 분자·분모가 큰 분수,
#  그리고 대입 기호(evaluation bar, \right|) — 는 지금까지 사진으로
#  굳어 버렸다. 이유는 두 가지였다.
#
#   ① 큰 기호는 CMEX/txex 같은 '확장 글꼴'로 조판된다. 이 글꼴의
#      글자를 그대로 읽으면 ∫ 는 'Z', ∑ 는 'P', 큰 괄호는 \x00 처럼
#      엉뚱한 값이 나온다. 그래서 수식 점수가 0점이 되어 텍스트로도,
#      LaTeX 로도 못 가고 배경 그림에 남았다.
#      → PDF 안에 들어 있는 글꼴 /Differences 인코딩을 읽어
#        'integraldisplay', 'radicalbig', 'braceleftbigg' 같은
#        진짜 글리프 이름을 얻고, 그것을 LaTeX 로 되돌린다.
#
#   ② 분수선은 벡터 선이 아니다. TeX(dvips)는 1x1 인라인 이미지를
#      납작하게 늘려 그린다. page.get_drawings() 로는 하나도 안 잡혀서
#      분자/분모를 나눌 수가 없었다.
#      → 콘텐츠 스트림에서 'cm ... BI' 패턴을 직접 읽어 분수선을 찾는다.
#
#  이 둘을 얻으면 글자들의 (x, y, 크기) 배치만으로 2차원 수식을
#  1차원 LaTeX 로 되돌릴 수 있다. 아래가 그 복원기다.


# ── TeX 확장 글꼴(CMEX/txex) 글리프 이름 → 의미 ──────────────
_OPEN  = {"parenleft":"(", "bracketleft":"[", "braceleft":r"\{",
          "angbracketleft":r"\langle", "floorleft":r"\lfloor", "ceilingleft":r"\lceil"}
_CLOSE = {"parenright":")", "bracketright":"]", "braceright":r"\}",
          "angbracketright":r"\rangle", "floorright":r"\rfloor", "ceilingright":r"\rceil"}
_BIGOP = {"integral":r"\int", "summation":r"\sum", "product":r"\prod",
          "union":r"\bigcup", "intersection":r"\bigcap", "coproduct":r"\coprod",
          "contintegral":r"\oint", "circlemultiply":r"\bigotimes", "circleplus":r"\bigoplus",
          "acute":r"\int", "circumflex":r"\int", "ffl":r"\oint",
          "int":r"\int", "iint":r"\iint", "iiint":r"\iiint", "oint":r"\oint",
          "sum":r"\sum", "prod":r"\prod", "coprod":r"\coprod"}
# 크기 접미사 (big / Big / bigg / Bigg / text / display / tp / bt / ex ...)
_SIZE_SUFFIX = re.compile(r"(big{1,2}|Big{1,2}|text|display|tp|bt|ex|mid)$", re.I)

# 14.0 · TeX CMEX 기본 인코딩. /Differences 와 글꼴 프로그램 표가 비어 있어도
# 제어문자(\x12=parenleftbigg 등)를 큰 괄호·적분으로 되돌린다.
_CMEX_STD = {
    0: "parenleft", 1: "parenright", 2: "bracketleft", 3: "bracketright",
    4: "floorleft", 5: "floorright", 6: "ceilingleft", 7: "ceilingright",
    8: "braceleft", 9: "braceright", 10: "angbracketleft", 11: "angbracketright",
    12: "vextendsingle", 13: "vextenddouble", 14: "slash", 15: "backslash",
    16: "parenleftbig", 17: "parenrightbig",
    18: "parenleftbigg", 19: "parenrightbigg",
    20: "parenleftBigg", 21: "parenrightBigg",
    22: "bracketleftbigg", 23: "bracketrightbigg",
    24: "braceleftbigg", 25: "bracerightbigg",
    26: "angbracketleftbigg", 27: "angbracketrightbigg",
    32: "integral", 33: "integraldisplay", 34: "contintegral",
    80: "summation", 81: "product", 86: "coproduct",
    112: "radical", 113: "radicalbig",
}


def classify_glyph(name):
    """CMEX 글리프 이름 → (역할, LaTeX 토막)

    역할: 'open' | 'close' | 'op' | 'radical' | 'vbar' | 'ext' | None
    """
    if not name:
        return (None, None)
    n = name.strip()
    if n in ("acute", "circumflex") or n.startswith("integral") or n in ("int", "iint", "iiint", "smallint"):
        return ("op", r"\int")
    if n in ("ffl", "oint", "oiint") or n.startswith("contintegral"):
        return ("op", r"\oint")
    if n.startswith("summation") or n in ("sum",):
        return ("op", r"\sum")
    if n.startswith("product") or n in ("prod",):
        return ("op", r"\prod")
    if n.startswith("coproduct") or n in ("coprod",):
        return ("op", r"\coprod")
    # 세로로 늘어나는 대입 기호 |  (\right| 로 쓰는 그것)
    if n.startswith("vextendsingle") or n in ("bar", "verticalbar"):
        return ("vbar", "|")
    if n.startswith("vextenddouble") or n == "arrowvert":
        return ("vbar", r"\|")
    if n.startswith("radical"):
        # radicalbt/vertex/tp 는 큰 루트의 조각들
        return ("radical", r"\sqrt")
    base = _SIZE_SUFFIX.sub("", n)
    if base in _OPEN:
        return ("open", _OPEN[base])
    if base in _CLOSE:
        return ("close", _CLOSE[base])
    if base in _BIGOP:
        return ("op", _BIGOP[base])
    # 큰 괄호의 위/중간/아래 조각 (parenlefttp, parenleftex ...)
    for k, v in _OPEN.items():
        if n.startswith(k):
            return ("open", v)
    for k, v in _CLOSE.items():
        if n.startswith(k):
            return ("close", v)
    for k, v in _BIGOP.items():
        if n.startswith(k):
            return ("op", v)
    return (None, None)


def font_glyph_tables(doc, page):
    """이 페이지 글꼴들의 /Differences 인코딩 표: {basefont: {code: glyphname}}"""
    out = {}
    try:
        fonts = page.get_fonts(full=True)
    except Exception:
        return out
    for f in fonts:
        xref = f[0]
        base = f[3] or ""
        try:
            obj = doc.xref_object(xref)
        except Exception:
            continue
        m = re.search(r"/Encoding\s+(\d+)\s+0\s+R", obj)
        if not m:
            continue
        try:
            enc = doc.xref_object(int(m.group(1)))
        except Exception:
            continue
        dm = re.search(r"/Differences\s*\[(.*?)\]", enc, re.S)
        if not dm:
            continue
        table, cur = {}, 0
        for tok in dm.group(1).split():
            if tok.startswith("/"):
                table[cur] = tok[1:]
                cur += 1
            else:
                try:
                    cur = int(tok)
                except ValueError:
                    pass
        if table:
            out[base] = table
            out[base.split("+")[-1]] = table
    return out


def _font_program_encoding(doc, xref):
    """10.1 · 임베디드 Type1(PFA/PFB) 글꼴 '파일 안'의 /Encoding 배열을 읽는다.

    pdflatex(txfonts 등) 글꼴은 PDF쪽 /Differences 가 비어 있어도 글꼴
    프로그램 헤더에 'dup 88 /summationdisplay put' 같은 진짜 글리프
    이름표를 갖고 있다. 이게 없으면 txex 의 Σ·∏·큰괄호가 'X','Y','!'
    같은 쓰레기 글자로 읽혀 수식이 박살난다.
    """
    try:
        buf = doc.extract_font(xref)[3]
        if not buf or len(buf) > 2_000_000:
            return {}
        txt = buf.decode("latin-1", "replace")
        m = re.search(r"/Encoding\s+256\s+array", txt)
        if not m:
            return {}
        seg = txt[m.end():m.end() + 16000]
        table = {}
        for a, b in re.findall(r"dup\s+(\d+)\s+/([A-Za-z0-9_.\-]+)\s+put", seg):
            table[int(a)] = b
        # 14.0 · 서브셋 CMEX(괄호 2개만 임베드)도 살려야 큰 괄호가 안 사라진다.
        if len(table) < 2:
            for a, b in re.findall(r"dup\s+(\d+)\s+/([A-Za-z0-9_.\-]+)\s+put", txt):
                table[int(a)] = b
        return table if table else {}
    except Exception:
        return {}


def glyph_tables_full(doc, page):
    """/Differences 와 글꼴 프로그램 인코딩을 합친 전체 글리프 이름표."""
    out = font_glyph_tables(doc, page)
    try:
        fonts = page.get_fonts(full=True)
    except Exception:
        return out
    for f in fonts:
        xref, base = f[0], (f[3] or "")
        short = base.split("+")[-1]
        t = _font_program_encoding(doc, xref)
        if t:
            if short in out:
                out[short].update(t)
            else:
                out[short] = t
            if base in out:
                out[base].update(t)
            else:
                out[base] = t
        if "ESINT" in base.upper() or "ESINT" in short.upper():
            es_map = {1: "integral", 2: "integral", 31: "contintegral", 94: "integral", 710: "integral"}
            if short not in out: out[short] = {}
            if base not in out: out[base] = {}
            out[short].update(es_map)
            out[base].update(es_map)
    return out


_NUM = r"[-+]?[\d.]+"
_CM_BI = re.compile(
    r"q\s+(" + _NUM + r")\s+(" + _NUM + r")\s+(" + _NUM + r")\s+(" + _NUM + r")\s+("
    + _NUM + r")\s+(" + _NUM + r")\s+cm\s+BI\b", re.S)


def page_rules(page):
    """가로 규칙선 = 분수선·루트 윗줄. (분수선을 못 찾으면 \\frac 복원이 불가능하다)

    TeX(dvips) 는 분수선을 벡터가 아니라 1x1 인라인 이미지를 납작하게
    늘려서 그린다. 그래서 get_drawings() 로는 하나도 안 잡힌다.
    """
    out = []
    ph = page.rect.height
    try:
        for d in page.get_drawings():
            r = d.get("rect")
            if r and r.width >= 3 and r.height <= 2.6:
                out.append([float(r.x0), float(r.y0), float(r.x1), float(r.y1)])
    except Exception:
        pass
    try:
        cs = page.read_contents().decode("latin-1", "replace")
    except Exception:
        cs = ""
    if cs:
        gs = re.search(r"q\s+(" + _NUM + r")\s+0\s+0\s+(" + _NUM + r")\s+0\s+0\s+cm", cs)
        sx = float(gs.group(1)) if gs else 1.0
        sy = float(gs.group(2)) if gs else 1.0
        for m in _CM_BI.finditer(cs):
            a, _b, _c, dd, e, f = [float(x) for x in m.groups()]
            w = abs(a) * sx
            h = abs(dd) * sy
            if w < 3 or h > 2.6:
                continue
            x0 = e * sx
            y0 = ph - (f * sy)
            out.append([x0, y0, x0 + w, y0 + h])
    ded = []
    for r in out:
        if not any(abs(r[0] - q[0]) < 0.6 and abs(r[1] - q[1]) < 0.6
                   and abs(r[2] - q[2]) < 0.6 for q in ded):
            ded.append(r)
    ded.sort(key=lambda r: (r[1], r[0]))
    return ded


# ═══════════════════════════════════════════════════════════
#  2차원 배치 → 1차원 LaTeX
# ═══════════════════════════════════════════════════════════
#  PDF 는 글자마다 (x, y, 크기) 만 준다. 분수는 '선 위/아래', 지수는
#  '작고 위', 큰 적분은 '아주 키 큰 글리프'로만 구분된다. 그래서
#  ① 분수선을 먼저 찾아 위/아래를 재귀로 나누고
#  ② 남은 글자들을 x 순서로 읽으며 첨자/큰 연산자/큰 괄호를 붙인다.

_GREEK = {
    "α": r"\alpha", "β": r"\beta", "γ": r"\gamma", "δ": r"\delta",
    "ε": r"\epsilon", "ζ": r"\zeta", "η": r"\eta", "θ": r"\theta",
    "ι": r"\iota", "κ": r"\kappa", "λ": r"\lambda", "μ": r"\mu", "µ": r"\mu",
    "ν": r"\nu", "ξ": r"\xi", "π": r"\pi", "ρ": r"\rho", "σ": r"\sigma",
    "τ": r"\tau", "υ": r"\upsilon", "φ": r"\phi", "ϕ": r"\phi", "χ": r"\chi",
    "ψ": r"\psi", "ω": r"\omega", "Γ": r"\Gamma", "Δ": r"\Delta",
    "Θ": r"\Theta", "Λ": r"\Lambda", "Ξ": r"\Xi", "Π": r"\Pi",
    "Σ": r"\Sigma", "Φ": r"\Phi", "Ψ": r"\Psi", "Ω": r"\Omega",
}
_SYM = {
    "∞": r"\infty", "∂": r"\partial", "∇": r"\nabla", "∫": r"\int",
    "∮": r"\oint", "∑": r"\sum", "∏": r"\prod", "√": r"\sqrt",
    "≈": r"\approx", "≃": r"\simeq", "≅": r"\cong", "≠": r"\ne",
    "≤": r"\le", "≥": r"\ge", "±": r"\pm", "∓": r"\mp", "×": r"\times",
    "÷": r"\div", "·": r"\cdot", "∝": r"\propto", "∈": r"\in",
    "∉": r"\notin", "⊂": r"\subset", "⊃": r"\supset", "∪": r"\cup",
    "∩": r"\cap", "→": r"\to", "←": r"\leftarrow", "↔": r"\leftrightarrow",
    "⇒": r"\Rightarrow", "⇔": r"\Leftrightarrow", "∀": r"\forall",
    "∃": r"\exists", "∥": r"\parallel", "⊥": r"\perp", "ℏ": r"\hbar",
    "ℓ": r"\ell", "−": "-", "≡": r"\equiv", "≫": r"\gg", "≪": r"\ll",
    "⊗": r"\otimes", "⊕": r"\oplus", "′": "'", "″": "''", "…": r"\dots",
    "⟨": r"\langle", "⟩": r"\rangle", "∧": r"\wedge", "∨": r"\vee",
    "ˆ": r"\int", "◦": r"^{\circ}", "□": r"\square",
}
_MATHOP = re.compile(r"^(sin|cos|tan|cot|sec|csc|arcsin|arccos|arctan|sinh|cosh|tanh|"
                     r"exp|log|ln|lim|det|dim|ker|deg|gcd|max|min|sup|inf|arg|Tr|tr)$")


# 같은 모양 다른 코드포인트 정규화 (Ω 옴 기호 U+2126, µ 마이크로 U+00B5 등)
_LOOKALIKE = {
    "\u2126": "\u03a9",   # OHM SIGN      → GREEK CAPITAL OMEGA
    "\u00b5": "\u03bc",   # MICRO SIGN    → GREEK SMALL MU
    "\u2206": "\u0394",   # INCREMENT     → GREEK CAPITAL DELTA
    "\u220f": "\u03a0", "\u2211": "\u03a3",
    "\u2212": "-", "\u2010": "-", "\u2011": "-", "\u2013": "-", "\u2014": "-",
    "\u02b9": "'", "\u2032": "'",
}


_LIGATURES = {"\ufb00": "ff", "\ufb01": "fi", "\ufb02": "fl",
              "\ufb03": "ffi", "\ufb04": "ffl", "\ufb05": "st", "\ufb06": "st"}


def _tok_tex(t):
    """한 글자를 LaTeX 로."""
    t = _LOOKALIKE.get(t, t)
    if t in _LIGATURES:
        return r"\mathrm{" + _LIGATURES[t] + "}"
    if t and ord(t[0]) < 0x20:
        return ""          # 인코딩을 못 얻은 제어문자는 버린다
    if t in _GREEK: return _GREEK[t]
    if t in _SYM:   return _SYM[t]
    if t in "%#&_{}$": return "\\" + t
    return t


class Box:
    """한 글자 또는 이미 조립된 덩어리."""
    __slots__ = ("x0", "y0", "x1", "y1", "tex", "size", "role", "atomic")

    def __init__(self, x0, y0, x1, y1, tex, size, role=None, atomic=False):
        self.x0, self.y0, self.x1, self.y1 = x0, y0, x1, y1
        self.tex = tex
        self.size = size
        self.role = role          # 'open'|'close'|'op'|'radical'|'vbar'|None
        self.atomic = atomic      # 이미 완성된 덩어리(재귀 결과)

    @property
    def cy(self): return (self.y0 + self.y1) / 2.0
    @property
    def h(self):  return self.y1 - self.y0
    @property
    def w(self):  return self.x1 - self.x0

    def __repr__(self):
        return f"Box({self.tex!r} x={self.x0:.0f}..{self.x1:.0f} y={self.y0:.0f}..{self.y1:.0f})"


def _median(v, d=10.0):
    v = sorted(v)
    return v[len(v) // 2] if v else d


def _covered(b, rx0, rx1):
    """상자가 분수선 폭 안에 '대부분' 들어가는가.

    느슨하게 보면 분수선 왼쪽의 '=' 나 오른쪽의 다음 항까지 분자로
    빨려 들어간다(U_0^4 = 2^7/3 ... 가 \\frac{=2^7 \\pi}{3} 로 깨지던 원인).
    """
    w = max(1e-6, b.x1 - b.x0)
    ov = min(b.x1, rx1 + 1.0) - max(b.x0, rx0 - 1.0)
    return ov > 0 and (ov / w) >= 0.6


def _classify_rules(boxes, rules):
    """규칙선을 '분수선'과 '근호 윗줄'로 나눈다.

    근호 윗줄은 왼쪽 끝에 루트 글리프가 딱 붙어 있다. 이걸 먼저 갈라놓지
    않으면 분수 분자까지 근호 안으로 빨려 들어간다(√(U²/4πgN) 오인).
    """
    rads = [b for b in boxes if b.role == "radical"]
    frac, radbar = [], {}
    for r in rules:
        owner = None
        for b in rads:
            if abs(r[0] - b.x1) <= 4.5 and (b.y0 - 4.5) <= (r[1] + r[3]) / 2.0 <= (b.y1 + 4.5):
                owner = b
                break
        if owner is not None:
            radbar[id(owner)] = r
        else:
            frac.append(r)
    return frac, radbar


def _row_split(boxes, rules):
    """여러 줄로 조판된 수식을 '줄' 단위로 가른다.

    기준은 '본문 크기 글자의 베이스라인'이다. 분수의 분자·분모나 첨자는
    베이스라인이 없고 큰 글자 주위에 딸려 있을 뿐이라, 본문 글자의
    아랫선만 모으면 진짜 줄이 몇 개인지 깔끔하게 드러난다.
    """
    real = [b for b in boxes if not b.atomic and b.role is None]
    if len(real) < 6:
        return None
    top = max(b.size for b in real)
    mains = [b for b in real if b.size >= top * 0.9]
    if len(mains) < 6:
        return None
    hh = _median([b.h for b in mains]) or 10.0

    bl = sorted(b.y1 for b in mains)
    groups = [[bl[0]]]
    for v in bl[1:]:
        if v - groups[-1][-1] > hh * 0.85:
            groups.append([v])
        else:
            groups[-1].append(v)
    # 글자가 몇 개 안 되는 무리는 첨자 잔재다 → 줄로 치지 않는다
    groups = [g for g in groups if len(g) >= 3]
    if len(groups) < 2:
        return None
    centers = [_median(g) for g in groups]

    # 넓은 분수선이 줄 경계를 넘으면 그건 줄바꿈이 아니라 분수다
    x0 = min(b.x0 for b in boxes); x1 = max(b.x1 for b in boxes)
    width = max(1e-6, x1 - x0)
    for rl in rules:
        rcy = (rl[1] + rl[3]) / 2.0
        # 10.0 · '분수선 폭 ≥ 전체 60%' 조건만으로는 부족했다.
        #  식 앞에 'F =' 나 식 번호 '(3)' 이 붙으면 분수선은 전체 폭의 60%가
        #  못 되고, 큰 분수가 두 '줄'로 쪼개져 \\begin{aligned} 로 뒤집혔다.
        #  폭 대신 '그 선 위·아래에 실제로 글자가 있는가'로 판정한다.
        span = [b for b in boxes if _covered(b, rl[0], rl[2])]
        up = [b for b in span if b.cy < rcy - 0.5]
        dn = [b for b in span if b.cy > rcy + 0.5]
        if not up or not dn:
            continue
        if (rl[2] - rl[0]) < width * 0.6 and len(up) < 3 and len(dn) < 3:
            continue
        for a, b2 in zip(centers, centers[1:]):
            if a < rcy < b2:
                return None

    # 분수는 분자가 윗줄 베이스라인에 가깝게 놓이는 일이 흔하다.
    # 분수선이 속한 줄을 먼저 정하고, 그 분수선이 덮는 글자는 전부
    # 같은 줄로 못박아 분자/분모가 두 줄로 찢어지지 않게 한다.
    rows = [[] for _ in centers]
    pinned = {}
    gaps = [abs(a - b2) for a, b2 in zip(centers, centers[1:])] or [hh * 2.0]
    reach = max(hh * 0.9, min(gaps) * 0.45)   # 이웃 줄까지 넘어가지 않는 거리
    for rl in rules:
        rcy = (rl[1] + rl[3]) / 2.0
        k = min(range(len(centers)), key=lambda j: abs(rcy - centers[j]))
        for b in boxes:
            if _covered(b, rl[0], rl[2]) and abs(b.cy - rcy) <= reach:
                pinned[id(b)] = k
    for b in boxes:
        k = pinned.get(id(b))
        if k is None:
            k = min(range(len(centers)), key=lambda j: abs(b.y1 - centers[j]))
        rows[k].append(b)
    rows = [r for r in rows if r]
    if len(rows) < 2:
        return None
    for r in rows:
        if (max(z.x1 for z in r) - min(z.x0 for z in r)) < width * 0.25:
            return None
    return rows


def assemble(boxes, rules, depth=0):
    """상자들 + 규칙선들 → LaTeX 문자열 (재귀)."""
    boxes = [b for b in boxes if (b.tex or "").strip() != "" or b.atomic]
    if not boxes:
        return ""
    if depth > 14:
        return " ".join(b.tex for b in sorted(boxes, key=lambda b: b.x0))

    if depth == 0:
        cased = _try_cases(boxes, rules, depth)
        if cased:
            return cased
        rws = _row_split(boxes, rules)
        if rws:
            parts = [assemble(r, rules, depth + 1) for r in rws]
            parts = [p for p in parts if p.strip()]
            if len(parts) > 1:
                return (r"\begin{aligned} " + r" \\ ".join(parts)
                        + r" \end{aligned}")

    fracs, radbar = _classify_rules(boxes, rules)

    # ── ① 가장 바깥(가장 넓은) 분수선으로 위/아래를 가른다 ──
    best = None
    for r in fracs:
        rx0, ry0, rx1, ry1 = r
        rcy = (ry0 + ry1) / 2.0
        span = [b for b in boxes if _covered(b, rx0, rx1)]
        above = [b for b in span if b.cy < rcy - 0.5]
        below = [b for b in span if b.cy > rcy + 0.5]
        if not above or not below:
            continue
        # 분수선은 분자·분모를 '거의' 덮어야 한다.
        # 10.0 · 큰 분수는 큰 괄호·적분 기호가 분수선 끝보다 살짝 더 나오는
        #  일이 흔하다. 분수선이 클수록 여유를 함께 늘린다(작은 분수는
        #  여전히 2.5pt 엄격 — 옆 글자가 분자로 빨려 들어가는 일은 그대로 방지).
        slack = max(2.5, 0.045 * (rx1 - rx0))
        if any(b.x0 < rx0 - slack or b.x1 > rx1 + slack for b in above + below):
            continue
        # 분자·분모가 둘 다 분수선 폭의 일부라도 실제로 차지해야 한다
        if not above or not below:
            continue
        w = rx1 - rx0
        if best is None or w > best[0]:
            best = (w, rcy, rx0, rx1, r)

    if best is not None:
        _w, rcy, rx0, rx1, used = best
        num, den, rest = [], [], []
        for b in boxes:
            if _covered(b, rx0, rx1):
                bar = radbar.get(id(b))
                eff_cy = (bar[1] + bar[3]) / 2.0 if (b.role == "radical" and bar) else b.cy
                (num if eff_cy < rcy else den).append(b)
            else:
                rest.append(b)
        sub = [r for r in rules if r is not used]
        frac = (r"\frac{" + assemble(num, sub, depth + 1) + "}{"
                + assemble(den, sub, depth + 1) + "}")
        if rest:
            # 이 덩어리의 '글자 크기' 는 분자·분모 내용의 크기다. 영역 전체
            # 중앙값을 쓰면 지수로 올라간 작은 분수를 본문 크기로 착각한다.
            fb = Box(rx0, min(b.y0 for b in num), rx1, max(b.y1 for b in den),
                     frac, _median([b.size for b in (num + den)]), atomic=True)
            return assemble(rest + [fb], sub, depth + 1)
        return frac

    # ── ② 분수선이 없다 → x 순서로 읽으며 첨자·큰 연산자를 붙인다 ──
    return _linear(boxes, rules, depth)


def _linear(boxes, rules, depth):
    r"""분수선이 없는 영역: 왼쪽→오른쪽으로 읽는다.

    · 큰 적분/시그마  → \int_{아래}^{위}
    · 큰 여는 괄호    → \left( ... \right)  (짝을 찾아 재귀)
    · 큰 세로 막대 |  → \right| _{아래}^{위}   ← 대입 기호(evaluation bar)
    · 큰 루트         → \sqrt{ 윗줄이 덮는 범위 }
    · 작고 위/아래    → ^{ } / _{ }
    """
    bs = sorted(boxes, key=lambda b: (b.x0, b.y0))
    _fr, radbar = _classify_rules(bs, rules)
    base_sz = _median([b.size for b in bs if not b.atomic] or [10.0])
    # 본문 글자 높이(첨자 판정 기준): 큰 글리프를 뺀 중앙값
    norm_h = _median([b.h for b in bs if not b.atomic and b.role is None] or [base_sz])
    out = []
    taken = []          # 방금 평범하게 찍은 글자들 (큰 연산자 한계로 회수될 수 있다)
    i = 0
    n = len(bs)
    while i < n:
        b = bs[i]

        # ── 큰 괄호: 짝을 찾아 그 안을 통째로 재귀 ──
        if b.role == "open":
            depthc = 0
            j = -1
            for k in range(i, n):
                if bs[k].role == "open":  depthc += 1
                elif bs[k].role == "close":
                    depthc -= 1
                    if depthc == 0:
                        j = k
                        break
            if j > i:
                inner = assemble(bs[i + 1:j], rules, depth + 1)
                last = bs[j]
                ref_sz = bs[j].size
                ref_h = max(ref_sz * 0.72, bs[j].h, norm_h)
                base_cy = bs[j].cy
                cluster, pend = [], []
                m_idx = j + 1
                while m_idx < n:
                    c = bs[m_idx]
                    if c.atomic:
                        if (c.size < ref_sz * 0.95 and (last.y1 - c.y1) >= ref_h * 0.25
                                and c.x0 - max([z.x1 for z in (cluster + pend)] or [last.x1]) <= ref_h * 0.5):
                            cluster.extend(pend); pend = []
                            cluster.append(c)
                            m_idx += 1
                            continue
                        break
                    if c.role is not None:
                        break
                    prev_x1 = max([z.x1 for z in (cluster + pend)] or [last.x1])
                    if c.x0 - prev_x1 > ref_h * 0.5:
                        break
                    small = c.size < ref_sz * 0.95
                    same_run = bool(cluster) and abs(c.cy - cluster[-1].cy) < ref_h * 0.25
                    if not (small or same_run):
                        break
                    off = abs(c.cy - base_cy)
                    if off >= ref_h * 0.12:
                        cluster.extend(pend); pend = []
                        cluster.append(c)
                        m_idx += 1
                        continue
                    if small and c.size < ref_sz * 0.9:
                        pend.append(c)
                        m_idx += 1
                        continue
                    break
                if pend:
                    m_idx -= len(pend)
                subs = [c for c in cluster if c.cy > base_cy]
                sups = [c for c in cluster if c.cy < base_cy]
                grp = r"\left" + b.tex + " " + inner + r" \right" + bs[j].tex
                if sups: grp += "^{" + assemble(sups, rules, depth + 1) + "}"
                if subs: grp += "_{" + assemble(subs, rules, depth + 1) + "}"
                out.append(grp)
                i = m_idx
                continue
            # 짝이 없으면 그냥 큰 괄호 하나
            out.append(r"\left" + b.tex + r" \right.")
            i += 1
            continue

        if b.role == "close":
            out.append(r"\left. \right" + b.tex)
            i += 1
            continue

        # ── 대입 기호(큰 세로 막대) + 위/아래 한계 ──
        if b.role == "vbar" and b.h <= norm_h * 1.35:
            # 키가 크지 않은 세로 막대는 일반 기호로 처리 (첨자 수집 허용)
            b.role = None

        if b.role == "vbar":
            # 대입 기호: 오른쪽에 붙은 작은 글자들이 아래/위 한계다 (|_{cut})
            # 기준선: 막대 자체의 한가운데가 아니라 '옆 본문 글자'의 중심.
            # 막대는 위아래로 길어서 그 중심을 쓰면 아래첨자가 위첨자로 뒤집힌다.
            neigh = [c for c in bs if c.role is None and not c.atomic
                     and c.size >= b.size * 0.95]
            ref_cy = _median([c.cy for c in neigh] or [b.cy])
            lo, hi = [], []
            j = i + 1
            while j < n:
                c = bs[j]
                if c.atomic or c.role is not None:
                    break
                if c.size >= b.size * 0.95:
                    break
                prev = max([x.x1 for x in (lo + hi)] or [b.x1])
                if c.x0 - prev > b.size * 0.6:
                    break
                (lo if c.cy > ref_cy else hi).append(c)
                j += 1
            # 대입 기호의 '키'를 원본만큼 살린다. \left.\right| 만 쓰면
            # 내용이 없어 막대가 한 줄 높이로 쪼그라든다. 원본 막대가
            # 본문보다 몇 배 큰지 재서 \middle| 대신 크기 명령을 붙인다.
            ratio = b.h / max(1e-6, norm_h)
            if   ratio >= 3.2: bar = r"\Biggr|"
            elif ratio >= 2.5: bar = r"\biggr|"
            elif ratio >= 1.9: bar = r"\Bigr|"
            else:              bar = r"\bigr|"
            piece = bar
            if lo: piece += "_{" + assemble(lo, rules, depth + 1) + "}"
            if hi: piece += "^{" + assemble(hi, rules, depth + 1) + "}"
            out.append(piece)
            i = j
            continue

        # ── 큰 루트: 윗줄(rule)이 덮는 x 범위가 근호 안이다 ──
        if b.role == "radical":
            bar = radbar.get(id(b))
            if bar:
                bar_cy = (bar[1] + bar[3]) / 2.0
                inside = [c for c in bs[i + 1:]
                          if c.x0 >= bar[0] - 2.5 and c.x1 <= bar[2] + 2.5 and abs(c.cy - bar_cy) <= 14.0]
                sub = [r for r in rules if r is not bar]
                out.append(r"\sqrt{" + assemble(inside, sub, depth + 1) + "}")
                keep = [c for c in bs[i + 1:] if c not in inside]
                bs = bs[:i] + keep
                n = len(bs)
                continue
            # 윗줄이 없으면 바로 뒤 한 덩어리를 근호 안으로
            if i + 1 < n:
                nxt = bs[i + 1]
                out.append(r"\sqrt{" + nxt.tex + "}")
                bs = bs[:i] + bs[i + 2:]
                n = len(bs)
                continue
            out.append(r"\sqrt{}")
            i += 1
            continue
            # 윗줄이 없으면 바로 뒤 한 덩어리를 근호 안으로
            if i + 1 < n:
                nxt = bs[i + 1]
                out.append(r"\sqrt{" + nxt.tex + "}")
                bs = bs[:i] + bs[i + 2:]
                n = len(bs)
                continue
            out.append(r"\sqrt{}")
            i += 1
            continue

        # ── 큰 연산자(∫ ∑ ∏): 위/아래 한계를 모은다 ──
        if b.role == "op" or (b.tex in (r"\int", r"\sum", r"\prod", r"\oint")
                              and b.h > norm_h * 1.5):
            # 큰 연산자의 한계는 연산자보다 살짝 왼쪽에서 시작하기도 한다
            # (\sum_{states} 의 s 가 ∑ 왼쪽에 걸침) → 이미 처리한 앞 글자도 회수한다
            lo, hi = [], []
            reach = b.x1 + max(3.0, b.w * 1.15)
            span0 = b.x0 - b.w * 0.55
            for c in list(taken):
                if c.x0 >= span0 and c.size < b.size * 0.95:
                    (lo if c.cy > b.cy else hi).append(c)
                    taken.remove(c)
                    if out and out[-1] == c.tex:
                        out.pop()
            # 한계 글자는 연산자 폭을 조금 넘어가기도 한다(\sum_{states}).
            # 작은 글자가 끊기지 않고 이어지는 동안 계속 받아들인다.
            k = i + 1
            while k < n:
                c = bs[k]
                if c.atomic or c.role is not None:
                    break
                small = (c.size < b.size * 0.95) or (c.h < norm_h * 0.95)
                if not small:
                    break
                prev = max([z.x1 for z in (lo + hi)] or [b.x0])
                if c.x0 > reach and c.x0 - prev > b.size * 0.28:
                    break
                (lo if c.cy > b.cy else hi).append(c)
                k += 1
            lo.sort(key=lambda z: z.x0); hi.sort(key=lambda z: z.x0)
            piece = b.tex
            if lo: piece += "_{" + assemble(lo, rules, depth + 1) + "}"
            if hi: piece += "^{" + assemble(hi, rules, depth + 1) + "}"
            out.append(piece)
            i = k
            continue

        # ── 보통 글자 (+ 뒤따르는 첨자) ──
        word = b.tex
        last = b               # 낱말로 합친 마지막 글자 (첨자 위치 기준)
        taken.append(b)
        i += 1
        # 여러 글자로 된 함수 이름 묶기 (sin, exp, lim …)
        while (i < n and not bs[i].atomic and bs[i].role is None
               and len(bs[i].tex) == 1 and bs[i].tex.isascii() and bs[i].tex.isalpha()
               and word.isascii() and word.isalpha()
               and abs(bs[i].cy - b.cy) < norm_h * 0.3
               and bs[i].x0 - bs[i - 1].x1 < norm_h * 0.18
               and abs(bs[i].size - b.size) < 0.4):
            word += bs[i].tex
            last = bs[i]
            i += 1
        if _MATHOP.match(word):
            out.append("\\" + word if word not in ("Tr", "tr") else r"\mathrm{Tr}")
        else:
            out.append(word if len(word) == 1 else word)

        # 첨자 수집: 바로 오른쪽에 붙은 '작은' 글자들.
        #   기준은 전체 중앙값이 아니라 '지금 이 글자(b)' 다. 중앙값을 쓰면
        #   M^{tree}_{(1)} 처럼 작은 글자가 더 많은 식에서 기준이 작아져
        #   첨자를 하나도 못 잡는다.
        #   위/아래 첨자는 x 로 섞여 나오므로(t ( r 1 e ) e) 한 덩어리로
        #   모은 뒤 높이로 가른다.
        # 관계·연산 기호(= + - < >)에는 첨자가 붙지 않는다. 이걸 막지 않으면
        # '= \sum_{states}' 의 s 가 '=' 의 아래첨자로 붙어버린다.
        if word in ("=", "+", "-", r"\pm", r"\mp", r"\to", r"\approx",
                    r"\equiv", r"\le", r"\ge", "<", ">", r"\ne", r"\simeq"):
            continue
        ref_sz = b.size
        # 기준 높이는 글자마다 들쭉날쭉하다('-' 는 낮고 'l' 은 높다).
        # 글꼴 크기를 1차 기준으로 쓰고, 높이는 보조로만 본다.
        ref_h = max(ref_sz * 0.72, b.h)
        base_cy = b.cy
        cluster = []
        pend = []          # 크기는 첨자인데 위/아래가 애매한 글자 (판단 보류)
        while i < n:
            c = bs[i]
            # 이미 조립된 덩어리(분수 등)도 작고 위로 떠 있으면 지수다.
            #   (V/(4/3)π)^{1/3} 의 1/3 이 이 경우.
            if c.atomic:
                if (c.size < ref_sz * 0.9
                        and (last.y1 - c.y1) >= ref_h * 0.30
                        and c.x0 - max([z.x1 for z in (cluster + pend)] or [last.x1]) <= ref_h * 0.5):
                    cluster.extend(pend); pend = []
                    cluster.append(c)
                    i += 1
                    continue
                break
            if c.role is not None:
                break
            prev_x1 = max([z.x1 for z in (cluster + pend)] or [last.x1])
            if c.x0 - prev_x1 > ref_h * 0.5:
                break
            small = c.size < ref_sz * 0.95
            same_run = bool(cluster) and abs(c.cy - cluster[-1].cy) < ref_h * 0.25
            if not (small or same_run):
                break
            bl_off = abs(c.y1 - last.y1)
            off = abs(c.cy - base_cy)
            if off >= ref_h * 0.12 or (small and bl_off >= ref_h * 0.05) or (small and c.size < ref_sz * 0.82):
                # 위/아래가 분명하다 → 보류분까지 함께 확정
                cluster.extend(pend); pend = []
                cluster.append(c)
                i += 1
                continue
            # 크기는 첨자인데 중심이 애매한 글자(σ^{⊗n} 의 ⊗ 처럼 글리프 상자가
            # 위아래로 긴 기호). 뒤에 확실한 첨자가 이어지면 그때 함께 넣는다.
            if small and c.size < ref_sz * 0.9:
                pend.append(c)
                i += 1
                continue
            break
        # 끝까지 확정되지 않은 보류분은 첨자가 아니다 → 되돌린다
        if pend:
            i -= len(pend)
        # 14.0 · 첨자 판정은 글자 상자 중심(cy)보다 베이스라인(y1)이 정확하다.
        #   칠판체 S^n 처럼 본문 글리프가 커서 cy 가 거의 같아도 n 은 위첨자다.
        ref_bl = last.y1
        subs, sups, rest = [], [], []
        for c in cluster:
            if c.y1 > ref_bl + max(0.25, ref_h * 0.06):
                subs.append(c)
            elif c.y1 < ref_bl - max(0.25, ref_h * 0.06):
                sups.append(c)
            else:
                rest.append(c)
        for c in rest:
            (subs if c.cy > last.cy else sups).append(c)
        if sups and out and out[-1].rstrip().endswith("'"):
            out[-1] = "{" + out[-1].strip() + "}"
        if sups: out.append("^{" + assemble(sups, rules, depth + 1) + "}")
        if subs: out.append("_{" + assemble(subs, rules, depth + 1) + "}")

    return " ".join(x for x in out if x).strip()


def region_boxes(doc, page, rect, gtables=None):
    """페이지의 한 영역에서 Box 목록을 만든다 (CMEX 글리프 해석 포함)."""
    if gtables is None:
        gtables = font_glyph_tables(doc, page)
    x0, y0, x1, y1 = rect
    out = []
    try:
        rd = page.get_text("rawdict")
    except Exception:
        return out
    for blk in rd.get("blocks", []):
        if blk.get("type") != 0:
            continue
        for ln in blk.get("lines", []):
            for sp in ln.get("spans", []):
                fname = (sp.get("font") or "")
                size = float(sp.get("size") or 10)
                short = fname.split("+")[-1]
                table = gtables.get(fname) or gtables.get(short) or {}
                ext = ("CMEX" in short.upper() or "TXEX" in short.upper()
                       or "EXTRA" in short.upper() or "LMEX" in short.upper()
                       or "ESINT" in short.upper())
                for ch in (sp.get("chars") or []):
                    c = ch.get("c") or ""
                    bb = ch.get("bbox")
                    if not bb:
                        continue
                    cx = (bb[0] + bb[2]) / 2.0
                    cy = (bb[1] + bb[3]) / 2.0
                    if not (x0 - 1 <= cx <= x1 + 1 and y0 - 1 <= cy <= y1 + 1):
                        continue
                    role = None
                    tex = None
                    if ext:
                        gname = table.get(ord(c)) if c else None
                        if not gname and c:
                            gname = _CMEX_STD.get(ord(c))
                        role, tex = classify_glyph(gname)
                        if tex is None:
                            # 이름을 못 얻었으면 이 글리프는 버린다.
                            # (그대로 두면 'Z','p' 같은 쓰레기 글자가 식에 박힌다)
                            # 단 PUA(F8xx) 조각(큰 괄호를 세로로 쌓은 조각)은
                            # 나중에 한 덩어리로 조립한다(cases 큰중괄호 등).
                            if c and 0xE000 <= ord(c) <= 0xF8FF:
                                out.append(Box(bb[0], bb[1], bb[2], bb[3],
                                               "\ue000", size, "pua"))
                            continue
                    else:
                        if not c.strip():
                            continue
                        tex = _tok_tex(c)
                        if c in ("|",):
                            role = "vbar"
                        elif c == "√":
                            role = "radical"
                        elif c in ("∫", "∑", "∏", "∮"):
                            role = "op"
                    out.append(Box(bb[0], bb[1], bb[2], bb[3], tex, size, role))
    out.sort(key=lambda b: (b.x0, b.y0))
    out = _merge_vbars(out)
    out = _merge_stacked_delims(out)
    out = _assemble_pua_pieces(out)
    return out


def _assemble_pua_pieces(boxes):
    """세로로 쌓인 큰-괄호 조각(PUA)들을 한 덩어리(기둥)로 조립한다.

    PDF에 따라 큰 괄호 조각이 ToUnicode 없이 읽혀 U+F8F1 같은 PUA 문자로
    나온다. 한 x-기둥에 잇달아 쌓인 조각들 = 큰 괄호 하나. 기둥이 내용의
    왼쪽 끝이면 여는 괄호, 오른쪽 끝이면 닫는 괄호로 쓴다(cases 중괄호 등).
    """
    try:
        pieces = [b for b in boxes if b.role == "pua"]
        if len(pieces) < 2:
            boxes2 = [b for b in boxes if b.role != "pua"]
            boxes2.sort(key=lambda b: (b.x0, b.y0))
            return boxes2
        keep = [b for b in boxes if b.role != "pua"]
        pieces.sort(key=lambda b: (b.x0, b.y0))
        cols = []
        for p in pieces:
            for c in cols:
                if (abs(c[0][1].x0 - p.x0) <= 1.5
                        and p.y0 <= c[-1][1].y1 + 3.0):
                    c.append((p.y0, p))
                    c.sort(key=lambda u: u[0])
                    break
            else:
                cols.append([(p.y0, p)])
        for c in cols:
            ps = [u[1] for u in c]
            if len(ps) < 3:
                continue                      # 3조각 미만은 확실한 괄호가 아니다
            span = max(b.y1 for b in ps) - min(b.y0 for b in ps)
            tallest = max(b.y1 - b.y0 for b in ps)
            if span < tallest * 1.7:
                continue   # 세로로 '쌓인' 모양이 아니다(underbrace 좌우 절반 등)
            x0 = min(b.x0 for b in ps); x1 = max(b.x1 for b in ps)
            y0 = min(b.y0 for b in ps); y1 = max(b.y1 for b in ps)
            lo = min([b.x0 for b in keep] + [x0])
            hi = max([b.x1 for b in keep] + [x1])
            fracL = (x0 - lo) / max(1e-6, hi - lo)
            if fracL >= 0.65:
                keep.append(Box(x0, y0, x1, y1, r"\}", ps[0].size, "close"))
            else:
                keep.append(Box(x0, y0, x1, y1, r"\{", ps[0].size, "open"))
        keep.sort(key=lambda b: (b.x0, b.y0))
        return keep
    except Exception:
        return boxes


def _merge_vbars(boxes):
    """세로 막대(대입 기호)는 조각을 여러 개 쌓아 그린다 → 하나로 합친다."""
    bars = [b for b in boxes if b.role == "vbar"]
    rest = [b for b in boxes if b.role != "vbar"]
    used, merged = set(), []
    for i, b in enumerate(bars):
        if i in used:
            continue
        y0, y1 = b.y0, b.y1
        for j in range(i + 1, len(bars)):
            if j in used:
                continue
            c = bars[j]
            if abs(c.x0 - b.x0) <= 1.2 and c.y0 <= y1 + 2.0 and c.y1 >= y0 - 2.0:
                y0 = min(y0, c.y0); y1 = max(y1, c.y1)
                used.add(j)
        merged.append(Box(b.x0, y0, b.x1, y1, b.tex, b.size, "vbar"))
    out = rest + merged
    out.sort(key=lambda b: (b.x0, b.y0))
    return out


def _merge_stacked_delims(boxes):
    """14.3 · 세로로 쌓인 큰 괄호({ } ( ))를 하나로 합친다.

    cases 중괄호는 위·아래 조각 두 개로 조판되는 일이 많다. 합치지 않으면
    \\left\\{ \\right. 가 두 번 나오고, 그 안의 분수가 한 덩어리로 뭉개진다.
    """
    kinds = ("open", "close")
    special = [b for b in boxes if b.role in kinds]
    rest = [b for b in boxes if b.role not in kinds]
    if len(special) < 2:
        return boxes
    used, merged = set(), []
    special.sort(key=lambda b: (round(b.x0, 1), b.y0))
    for i, b in enumerate(special):
        if i in used:
            continue
        y0, y1 = b.y0, b.y1
        x0, x1 = b.x0, b.x1
        for j in range(i + 1, len(special)):
            if j in used:
                continue
            c = special[j]
            if c.role != b.role or (c.tex or "") != (b.tex or ""):
                continue
            if abs(c.x0 - x0) <= 2.8 and c.y0 <= y1 + 6.0:
                y0 = min(y0, c.y0); y1 = max(y1, c.y1)
                x0 = min(x0, c.x0); x1 = max(x1, c.x1)
                used.add(j)
        merged.append(Box(x0, y0, x1, y1, b.tex, b.size, b.role))
    out = rest + merged
    out.sort(key=lambda b: (b.x0, b.y0))
    return out


def _cases_rows(boxes):
    """cases 오른쪽 글자를 조건식(k>0 등) 앵커 기준으로 줄 가른다."""
    if not boxes:
        return []
    rel = [b for b in boxes if (b.tex or "") in (
        ">", "<", "=", r"\ge", r"\le", r"\neq", r"\ne", r"\gt", r"\lt")]
    if len(rel) >= 2:
        rel = sorted(rel, key=lambda b: b.cy)
        # 너무 가까운 관계기호는 같은 줄
        hh = _median([b.h for b in boxes if b.role is None] or [10.0]) or 10.0
        anchors = [rel[0].cy]
        for b in rel[1:]:
            if abs(b.cy - anchors[-1]) > hh * 0.7:
                anchors.append(b.cy)
        if len(anchors) >= 2:
            rows = [[] for _ in anchors]
            for b in boxes:
                k = min(range(len(anchors)), key=lambda j: abs(b.cy - anchors[j]))
                rows[k].append(b)
            return [r for r in rows if r]
    real = [b for b in boxes if not b.atomic]
    if not real:
        return [boxes]
    hh = _median([b.h for b in real if b.role is None] or [10.0]) or 10.0
    items = sorted(real, key=lambda b: b.y1)
    groups = [[items[0]]]
    for b in items[1:]:
        prev = groups[-1]
        cy = _median([z.y1 for z in prev])
        if abs(b.y1 - cy) > hh * 0.72:
            groups.append([b])
        else:
            prev.append(b)
    rows = []
    for g in groups:
        ids = {id(z) for z in g}
        extra = [b for b in boxes if id(b) not in ids and any(
            abs(b.cy - z.cy) < hh * 0.6 for z in g)]
        rows.append(g + extra)
    return [r for r in rows if r]


def _try_cases(boxes, rules, depth):
    r"""큰 왼쪽 중괄호 + 여러 줄 → \begin{cases} … \end{cases}."""
    try:
        opens = [b for b in boxes if b.role == "open" and (
            (b.tex or "").endswith("{") or (b.tex or "") in ("{", r"\{"))]
        if not opens:
            return None
        body_h = _median([b.h for b in boxes if b.role is None and not b.atomic] or [10.0]) or 10.0
        tall = [b for b in opens if b.h >= max(22.0, body_h * 2.15)]
        if not tall:
            return None
        brace = min(tall, key=lambda b: (b.x0, -b.h))
        prefix = [b for b in boxes if id(b) != id(brace) and b.x1 <= brace.x0 + 1.2]
        right = [b for b in boxes if id(b) != id(brace) and b.x0 >= brace.x1 - 2.5]
        if len(right) < 3:
            return None
        right.sort(key=lambda b: b.x0)
        cut = None
        for i in range(len(right) - 1):
            gap = right[i + 1].x0 - right[i].x1
            if gap >= 18.0:
                left_n = i + 1
                body_part = right[:left_n]
                if len(_cases_rows(body_part)) >= 2:
                    cut = left_n
                    break
        suffix = []
        if cut is not None:
            suffix = right[cut:]
            right = right[:cut]
        rows = _cases_rows(right)
        if len(rows) < 2:
            return None
        parts = []
        for row in rows:
            row = sorted(row, key=lambda b: b.x0)
            piece = None
            if len(row) >= 3:
                gaps = [(row[i + 1].x0 - row[i].x1, i) for i in range(len(row) - 1)]
                gap, gi = max(gaps)
                med = _median([b.w for b in row]) or 8.0
                if gap > max(8.0, med * 1.4):
                    expr = assemble(row[:gi + 1], rules, depth + 1)
                    cond = assemble(row[gi + 1:], rules, depth + 1)
                    piece = (expr or "") + " & " + (cond or "")
            if piece is None:
                piece = assemble(row, rules, depth + 1)
            if (piece or "").strip():
                parts.append(piece.strip())
        if len(parts) < 2:
            return None
        body = r"\begin{cases} " + r" \\ ".join(parts) + r" \end{cases}"
        out = []
        if prefix:
            px = assemble(prefix, rules, depth + 1)
            if px:
                out.append(px)
        out.append(body)
        if suffix:
            sx = assemble(suffix, rules, depth + 1)
            if sx:
                out.append(sx)
        return " ".join(out).strip()
    except Exception:
        return None


_CTRL_RE = re.compile(r"[\x00-\x1f\x7f]")


def _tidy_latex(t):
    """복원 결과를 KaTeX 가 반드시 읽을 수 있는 모양으로 다듬는다.

    KaTeX 는 x^{a}^{b} (이중 위첨자) 를 오류로 낸다. 조판상 첨자가 여러
    번 붙는 경우가 실제로 있으므로(∫Ldt = 4.8 fb^-1) 하나로 합친다.
    """
    if not t:
        return ""
    t = _CTRL_RE.sub("", t)
    # 14.3 · 유니코드 프라임/바는 KaTeX 가 빨간 오류 또는 빈 글리프로 낸다.
    t = (t.replace("\u2032", "'").replace("\u2033", "''").replace("\u2034", "'''")
           .replace("\u00b4", "'").replace("\u2019", "'").replace("\u2018", "'")
           .replace("\u2212", "-"))
    t = re.sub(r"([A-Za-z])\s*[\u00af\u02c9]", r"\\bar{\1}", t)
    t = t.replace("\u00af", r"\bar{}").replace("\u02c9", r"\bar{}")
    # 이계도함수: j ^{' '} / z ' '  →  j'' / z''   (이중 위첨자 오류 원천 차단)
    t = re.sub(r"\^\{\s*'\s*'\s*\}", "''", t)
    t = re.sub(r"\^\{\s*'\s*\}", "'", t)
    t = re.sub(r"(\\[A-Za-z]+|[A-Za-z])\s+'\s+'(?![A-Za-z])", r"\1''", t)
    # ^{a}^{b} 와 _{a}_{b} 를 각각 하나로. 두 종류가 번갈아 나오는 경우
    # (_{=}^{fb}_{4.8}) 도 있으므로 더 이상 줄지 않을 때까지 반복한다.
    for _ in range(8):
        prev = t
        t = re.sub(r"\^\{([^{}]*)\}((?:\s*_\{[^{}]*\})*)\s*\^\{([^{}]*)\}",
                   r"^{\1 \3}\2", t)
        t = re.sub(r"_\{([^{}]*)\}((?:\s*\^\{[^{}]*\})*)\s*_\{([^{}]*)\}",
                   r"_{\1 \3}\2", t)
        if t == prev:
            break
    # 깨진 분수 지수 결합: e.g. ) ^{-} \frac{1}{2} -> )^{- \frac{1}{2}}
    t = re.sub(r"(\)|\]|\}|[a-zA-Z0-9\|])\s*\^\{([^{}]*)\}\s*\\frac\{([^{}]*)\}\{([^{}]*)\}",
               r"\1^{\2 \\frac{\3}{\4}}", t)
    t = re.sub(r"(\)|\]|\}|[a-zA-Z0-9\|])\s*_\{([^{}]*)\}\s*\\frac\{([^{}]*)\}\{([^{}]*)\}",
               r"\1_{\2 \\frac{\3}{\4}}", t)
    t = re.sub(r"\\bigr\|\s*([A-Za-z0-9]+)\s*\\bigr\|", r"| \1 |", t)
    t = re.sub(r"\\biggr\|\s*([A-Za-z0-9]+)\s*\\biggr\|", r"| \1 |", t)
    t = re.sub(r"\\Bigr\|\s*([A-Za-z0-9]+)\s*\\Bigr\|", r"| \1 |", t)
    t = re.sub(r"(?<![|\\])\|\s*\|(?!\|)", r"\\parallel ", t)
    # 인수 없는 \sqrt 는 파스 오류다
    t = re.sub(r"\\sqrt(?!\s*[{\[])", r"\\sqrt{}", t)
    t = re.sub(r"\\sqrt\{\}", "", t)
    t = re.sub(r"\s+", " ", t).strip()
    while t.count(r"\left") > t.count(r"\right"):
        t += r" \right."
    while t.count(r"\right") > t.count(r"\left"):
        t = r"\left. " + t
    return t


def _latex_is_sane(t):
    """분명히 깨진 결과는 버린다 (그 자리는 기존 경로가 처리한다)."""
    if not t:
        return False
    if _CTRL_RE.search(t):
        return False
    # 10.1 · 이스케이프된 \{ \} 는 braces 짝에서 제외한다
    #   (cases 큰중괄호 = \left\{ … \right. 처럼 짝이 없는 정상적인 경우가 있다)
    if (len(re.findall(r"(?<!\\)\{", t)) != len(re.findall(r"(?<!\\)\}", t))):
        return False
    if t.count(r"\left") != t.count(r"\right"):
        return False
    # 알맹이가 거의 없는 것 (\sqrt 하나 등)
    if len(re.sub(r"[\s\\{}^_]|left|right|begin|end|aligned", "", t)) < 2:
        return False
    return True


def region_to_latex(doc, page, rect, gtables=None, rules=None):
    """PDF 영역 → LaTeX 한 줄."""
    boxes = region_boxes(doc, page, rect, gtables)
    if not boxes:
        return ""
    if rules is None:
        rules = page_rules(page)
    x0, y0, x1, y1 = rect
    # 10.0 · 분수선은 분자·분모 중 '넓은 쪽' 폭이라 밴드보다 더 넓은 경우가
    #  흔하다. 통째로 들어오는 선만 쓰면 그런 분수선이 전부 버려져 \frac 이
    #  만들어지지 않았다 → 밴드와 가로로 겹치는 선은 모두 포함한다.
    rs = [r for r in rules
          if r[0] < x1 + 3 and r[2] > x0 - 3 and y0 - 3 <= (r[1] + r[3]) / 2 <= y1 + 3]
    return _tidy_latex(assemble(boxes, rs))


def _expand_math_bands(rd, bands):
    """cases처럼 키 큰 밴드만, 짧은 이웃 수식 토큰을 좌우로 끌어들인다.

    아무 밴드나 키우면 본문 문장이 수식에 빨려 들어가므로
    높이가 큰(중괄호/분수 여러 줄) 밴드에만 적용한다.
    """
    if not bands:
        return bands
    glyphs = []
    for blk in rd.get("blocks", []):
        if blk.get("type") != 0:
            continue
        for ln in blk.get("lines", []):
            bb = ln.get("bbox")
            if not bb:
                continue
            if (bb[2] - bb[0]) > 120:
                continue
            txt = "".join(_span_text(sp) for sp in ln.get("spans", []))
            words = re.findall(r"[A-Za-z]+", txt)
            prose_hits = [w for w in words if w.lower() in _COMMON_PROSE]
            non_math = [w for w in words if w.lower() not in _MATH_WORDS and len(w) >= 3]
            if prose_hits or len(non_math) >= 2:
                continue
            glyphs.append(bb)
    out = [list(m) for m in bands]
    for m in out:
        if (m[3] - m[1]) < 40:
            continue
        for _ in range(3):
            grew = False
            for bb in glyphs:
                vov = min(bb[3], m[3]) - max(bb[1], m[1])
                if vov <= 2.0:
                    continue
                hgap = max(0.0, max(bb[0] - m[2], m[0] - bb[2]))
                if hgap > 40.0:
                    continue
                nx0, ny0 = min(m[0], bb[0]), min(m[1], bb[1])
                nx1, ny1 = max(m[2], bb[2]), max(m[3], bb[3])
                if ny1 - ny0 > (m[3] - m[1]) + 12:
                    continue
                if nx1 - nx0 > (m[2] - m[0]) + 90:
                    continue
                if (nx0, ny0, nx1, ny1) != (m[0], m[1], m[2], m[3]):
                    m[0], m[1], m[2], m[3] = nx0, ny0, nx1, ny1
                    grew = True
            if not grew:
                break
    return [m[:4] for m in out]


def _drop_contained_math(regs):
    """다른 수식 상자 안에 들어 있는 조각(첨자·프라임)을 버린다. 겹침 방지."""
    if not regs:
        return []
    ordered = sorted(regs, key=lambda r: -max(1.0, (r["x1"] - r["x0"]) * (r["y1"] - r["y0"])))
    keep = []
    for r in ordered:
        ra = max(1e-6, (r["x1"] - r["x0"]) * (r["y1"] - r["y0"]))
        contained = False
        for o in keep:
            ox = min(o["x1"], r["x1"]) - max(o["x0"], r["x0"])
            oy = min(o["y1"], r["y1"]) - max(o["y0"], r["y0"])
            if ox > 0 and oy > 0 and (ox * oy) / ra >= 0.72:
                contained = True
                break
            if (o["x0"] - 1.2 <= r["x0"] and o["y0"] - 1.2 <= r["y0"]
                    and o["x1"] + 1.2 >= r["x1"] and o["y1"] + 1.2 >= r["y1"]):
                contained = True
                break
        if not contained:
            keep.append(r)
    keep.sort(key=lambda r: (r["y0"], r["x0"]))
    return keep


def _big_math_bands(page, avoid=None):
    """이 쪽에서 '큰 수식' 이 놓인 줄 영역(밴드)들을 찾는다. (9.3)

    큰 수식의 표시는 두 가지다.
      · 확장 글꼴(CMEX/txex) 글리프 — 키 큰 적분·시그마·괄호·근호·대입 막대
      · 분수선 — 위/아래에 글자를 거느린 가로 규칙선
    한 식이 여러 '줄'로 쪼개져 나오므로(분자/분모/첨자가 각각 줄이 된다)
    세로로 이어지는 것들을 하나의 밴드로 합친다.
    """
    try:
        rd = page.get_text("rawdict")
    except Exception:
        return [], [], {}
    try:
        doc = page.parent
        gtables = glyph_tables_full(doc, page)   # 10.1 · 글꼴 프로그램 이름표 포함
    except Exception:
        doc, gtables = None, {}
    try:
        rules = page_rules(page)
    except Exception:
        rules = []

    # 10.4 · 표 안쪽의 가로선은 분수선이 아니다 (표→rac 오인 방지)
    if avoid:
        def _rule_in_avoid(r):
            cx, cy = (r[0] + r[2]) / 2.0, (r[1] + r[3]) / 2.0
            return any(a[0] - 1 <= cx <= a[2] + 1 and a[1] - 1 <= cy <= a[3] + 1
                       for a in avoid)
        rules = [r for r in rules if not _rule_in_avoid(r)]

    # '본문 문장'은 절대 밴드 씨앗이 되지 못한다.
    # (본문 안에 복잡한 인라인 수식이 있어 확장글꼴이 섞여 있어도 마찬가지 —
    #  그런 줄이 씨앗이 되면 옆의 디스플레이 수식과 합쳐지고, 산문과 겹쳐
    #  밴드 전체가 버려지면서 진짜 수식까지 사라졌다)
    def _is_prose(ln):
        txt = "".join(_span_text(sp) for sp in ln.get("spans", []))
        words = re.findall(r"[A-Za-z]+", txt)
        if not words:
            return False
        prose_hits = [w for w in words if w.lower() in _COMMON_PROSE]
        non_math = [w for w in words if w.lower() not in _MATH_WORDS and len(w) >= 3]
        return bool(prose_hits or len(non_math) >= 2)

    seeds = []
    for blk in rd.get("blocks", []):
        if blk.get("type") != 0:
            continue
        for ln in blk.get("lines", []):
            bb = ln.get("bbox")
            if not bb or _is_prose(ln):
                continue
            has_ext = False
            for sp in ln.get("spans", []):
                fn = (sp.get("font") or "").split("+")[-1].upper()
                if any(x in fn for x in ("CMEX", "TXEX", "EXTRA", "LMEX", "MSAM", "MSBM", "ESINT")):
                    has_ext = True
                    break
            # 10.0 · '큰 분수' 시드 조건 완화.
            #  예전 조건(분수선이 줄 bbox 안에 '통째로' 들어갈 것)은 분자·분모
            #  폭이 서로 다른 큰 분수에서는 한쪽이 항상 실패했다 — 분수선은 둘 중
            #  넓은 쪽 폭이므로, 좁은 쪽 줄보다 분수선이 길다. 그러면 그 줄은
            #  시드가 못 되고, 결국 분수가 분자/분모 둘로 쪼개져 버렸다.
            #  이제 ① 세로 여유를 줄 높이에 비례하고 ② 가로는 '분수선과 40% 이상
            #  겹치면' 통과시킨다.
            vtol = max(7.0, (bb[3] - bb[1]) * 0.35)
            near_rule = False
            for r in rules:
                rcy = (r[1] + r[3]) / 2.0
                if not (bb[1] - vtol <= rcy <= bb[3] + vtol):
                    continue
                hov = min(bb[2] + 4, r[2]) - max(bb[0] - 4, r[0])
                if (bb[0] - 4 <= r[0] and r[2] <= bb[2] + 4) \
                        or (bb[0] >= r[0] - 4 and bb[2] <= r[2] + 4) \
                        or hov >= 0.4 * max(1e-6, r[2] - r[0]):
                    near_rule = True
                    break
            if has_ext or near_rule:
                seeds.append([bb[0], bb[1], bb[2], bb[3]])
    if not seeds:
        return [], rules, gtables

    # 세로로 겹치거나 맞닿은 조각들을 한 식으로 합친다
    # 세로로 '실제로 겹치는' 조각만 한 식으로 본다. 단순히 맞닿았다고
    # 합치면 위아래로 나란한 별개의 식 두 개가 한 덩어리가 된다.
    seeds.sort(key=lambda r: (r[1], r[0]))
    bands = []      # [x0,y0,x1,y1,capH] — capH: 이 밴드가 자랈 수 있는 최대 높이
    for r in seeds:
        placed = False
        for m in bands:
            vov = min(r[3], m[3]) - max(r[1], m[1])
            # 10.1 · 세로로 겹치더라도 가로로 아득히 떨어진 조각끼리 붙이면
            #   한 '줄'의 서로 다른 식이 통째로 합쳐져 밴드가 페이지 전체로
            #   자라나고, 결국 산문과 겹쳐 통째로 버려졌다. 가로 근접 필수.
            hgap = max(0.0, max(r[0] - m[2], m[0] - r[2]))
            if vov > min(r[3] - r[1], m[3] - m[1]) * 0.28 and hgap <= 36.0:
                cap = max(m[4], max(46.0, (r[3] - r[1]) * 1.35))
                m[0] = min(m[0], r[0]); m[1] = min(m[1], r[1])
                m[2] = max(m[2], r[2]); m[3] = max(m[3], r[3])
                m[4] = max(cap, 46.0)
                placed = True
                break
        if not placed:
            bands.append([r[0], r[1], r[2], r[3],
                          max(46.0, (r[3] - r[1]) * 1.35)])

    # 10.0 · 분수선을 사이에 둔 두 밴드는 한 식이다.
    #   분자 밴드와 분모 밴드는 세로로 전혀 겹치지 않아, 위의 '세로 겹침'
    #   병합으로는 절대 합쳐지지 않는다. 그래서 (큰 수식)/(큰 수식) 분수가
    #   분자·분모 두 개의 서로 다른 수식으로 갈라져 버렸다.
    #   같은 분수선을 사이에 두고 마주 보는 밴드들을 하나로 묶는다.
    for r in rules:
        rcy = (r[1] + r[3]) / 2.0
        rw = max(1e-6, r[2] - r[0])
        grp = []
        for m in bands:
            vgap = max(0.0, max(m[1] - rcy, rcy - m[3]))   # 밴드가 분수선과 떨어진 거리
            hov = min(m[2], r[2]) - max(m[0], r[0])
            inside = (m[0] >= r[0] - 4 and m[2] <= r[2] + 4)   # 밴드가 분수선 폭 안에
            if vgap <= max(14.0, (m[3] - m[1]) * 0.8) and (inside or hov >= 0.4 * rw):
                grp.append(m)
        if len(grp) >= 2:
            grp.sort(key=lambda m: m[1])
            # 분수선을 기준으로 위·아래가 마주 보는 쌍일 때만 합친다
            if grp[0][3] <= rcy + 3.0 and grp[-1][1] >= rcy - 3.0:
                m0 = grp[0]
                for m in grp[1:]:
                    m0[0] = min(m0[0], m[0]); m0[1] = min(m0[1], m[1])
                    m0[2] = max(m0[2], m[2]); m0[3] = max(m0[3], m[3])
                    m0[4] = max(m0[4], m[4], 46.0)
                    bands.remove(m)

    for _ in range(4):
        grew = False
        # (a) 줄 조각 흡수 — 세로로 겹치거나 한 줄 높이 안쪽으로 붙고,
        #     가로로 겹치거나 바로 옆에 붙었으면 같은 식이다.
        for blk in rd.get("blocks", []):
            if blk.get("type") != 0:
                continue
            for ln in blk.get("lines", []):
                bb = ln.get("bbox")
                if not bb or _is_prose(ln):
                    continue
                lh = max(2.0, bb[3] - bb[1])
                for m in bands:
                    vov = min(bb[3], m[3]) - max(bb[1], m[1])
                    vgap = max(0.0, max(bb[1] - m[3], m[1] - bb[3]))
                    hov = min(bb[2], m[2]) - max(bb[0], m[0])
                    if not (vov > 0.5 or vgap <= 0.55 * lh):
                        continue
                    if lh > max(14.0, (m[3] - m[1]) * 2.6):
                        continue                      # 훨씬 큰 덩어리는 못 붙인다
                    hgap = max(0.0, max(bb[0] - m[2], m[0] - bb[2]))
                    ok = (hov >= (bb[2] - bb[0]) * 0.55
                          or hgap <= max(8.0, lh * 2.2)
                          or hov >= 0.4 * min(bb[2] - bb[0], m[2] - m[0]))
                    if not ok:
                        continue
                    nx0 = min(m[0], bb[0]); nx1 = max(m[2], bb[2])
                    ny0 = min(m[1], bb[1]); ny1 = max(m[3], bb[3])
                    if ny1 - ny0 > m[4]:
                        continue          # 키 제한 초과 — 여기서 자라면 폭주다
                    if (nx0, ny0, nx1, ny1) != (m[0], m[1], m[2], m[3]):
                        m[0], m[1], m[2], m[3] = nx0, ny0, nx1, ny1
                        grew = True
                    break
        # (b) 작은 밴드(확장글꼴 글리프 하나짜리: ⟨ ⟩ ∏ …)끼리, 또는
        #     이웃 밴드에 붙여 합친다 — 분리돼 있으면 식이 조각난다.
        bands.sort(key=lambda m: (m[1], m[0]))
        i = 0
        while i < len(bands):
            a = bands[i]
            for j in range(len(bands)):
                if j == i:
                    continue
                b = bands[j]
                if a is b:
                    continue
                vgap = max(0.0, max(b[1] - a[3], a[1] - b[3]))
                hgap = max(0.0, max(b[0] - a[2], a[0] - b[2]))
                small = (a[2] - a[0] < 12.0 or b[2] - b[0] < 12.0)
                if (small and vgap <= 26.0
                        and hgap <= max(14.0, (a[2] - a[0] + b[2] - b[0]))
                        and max(a[3], b[3]) - min(a[1], b[1]) <= max(a[4], b[4])):
                    a[0] = min(a[0], b[0]); a[1] = min(a[1], b[1])
                    a[2] = max(a[2], b[2]); a[3] = max(a[3], b[3])
                    a[4] = max(a[4], b[4], 46.0)
                    bands.pop(j)
                    grew = True
                    if j < i:
                        i -= 1
                    break
            else:
                i += 1
        if not grew:
            break

        # 10.4 · 한 등식이 여러 줄로 쪼개진 경우: 세로로 바로 붙은 두 밴드가
        #   가로로 많이 겹치면 한 식으로 합친다 (줄바꿈된 긴 등식).
        for _pass in range(2):
            hit = False
            for i in range(len(bands)):
                for j in range(i + 1, len(bands)):
                    a, b = bands[i], bands[j]
                    vgap = max(0.0, max(b[1] - a[3], a[1] - b[3]))
                    hov = min(a[2], b[2]) - max(a[0], b[0])
                    if (vgap <= 4.0 and hov >= 0.6 * min(a[2] - a[0], b[2] - b[0])
                            and max(a[3], b[3]) - min(a[1], b[1]) <= 55.0):
                        a[0] = min(a[0], b[0]); a[1] = min(a[1], b[1])
                        a[2] = max(a[2], b[2]); a[3] = max(a[3], b[3])
                        a[4] = max(a[4], b[4], 55.0)
                        bands.pop(j)
                        hit = True
                        break
                if hit:
                    break
            if not hit:
                break

    # 중복/포함된 내부 밴드 정리
    final_bands = []
    for b in bands:
        if any(b is not o and o[0] <= b[0] + 1 and o[1] <= b[1] + 1 and o[2] >= b[2] - 1 and o[3] >= b[3] - 1 for o in bands):
            continue
        final_bands.append(b)

    # 산문 줄과 세로로 겹쳐 버린 밴드는 신뢰할 수 없다 → 큰 수식 처리 포기
    bands = [m[:4] for m in final_bands]
    good = []
    for m in bands:
        clash = False
        for blk in rd.get("blocks", []):
            if blk.get("type") != 0:
                continue
            for ln in blk.get("lines", []):
                bb = ln.get("bbox")
                if not bb or not _is_prose(ln):
                    continue
                vov = min(bb[3], m[3]) - max(bb[1], m[1])
                hov = min(bb[2], m[2]) - max(bb[0], m[0])
                # 10.1 · 1~6pt 스치는 겹침(디센더 등)으로 식 전체를 버리던
                #   일이 있었다 → 산문 한 줄을 통째로 삼킬 때만 포기한다
                if vov > 7.0 and hov > 40.0:
                    clash = True
                    break
            if clash:
                break
        if not clash:
            good.append(m)
    bands = good

    # 14.3 · cases/큰 괄호 밴드는 좌우 이웃 수식 토큰(s_k(t)=, k>0, …)까지 포함한다.
    bands = _expand_math_bands(rd, bands)

    return bands, rules, gtables


def _pdf_page_lines(page, avoid=None):
    """한 페이지에서 단어를 뽑는다 (rawdict, 문자 단위 좌표 기반).

    반환: (lines, math_regions)
      lines = [{x0,y0,x1,y1,base,words:[...]}]
      math_regions = [{x0,y0,x1,y1,display:bool,text:str}] — LaTeX로 바꿀 수식 영역
    수식은 일반 글자 상자에서 제외하고, 하나의 LaTeX 요소로 만들도록 영역과 원문을 모은다.
    """
    try:
        rd = page.get_text("rawdict")
    except Exception:
        return [], []

    lines = []
    math_regions = []

# ── 9.3 · 큰 수식 먼저 ────────────────────────────────────
    # 키 큰 적분/시그마, 큰 분수, 대입 기호가 있는 영역은 일반 줄 판정에
    # 넘기지 않는다. 확장 글꼴 글자가 'Z','P' 같은 쓰레기로 읽혀서
    # 점수가 0점이 되고, 결국 배경 사진으로 굳어 버리기 때문이다.
    big_bands = []
    _gt = {}
    _rules = []
    try:
        bands, _rules, _gt = _big_math_bands(page, avoid=avoid)
        doc = page.parent
        for bb in bands:
            if (bb[2] - bb[0]) < 6 or (bb[3] - bb[1]) < 6:
                continue
            try:
                tex = region_to_latex(doc, page, tuple(bb), _gt, _rules)
            except Exception:
                tex = ""
            if not tex or len(tex) > 1200 or not _latex_is_sane(tex):
                continue
            # 10.4 · 그림 안 글자가 수식으로 오인된 쓰레기 버림
            if _tex_is_figure_junk(tex):
                continue
            # 글자가 거의 없는(그림에 가까운) 밴드는 건드리지 않는다
            if len(re.sub(r"[\s\\{}^_]", "", tex)) < 2:
                continue
            szs = []
            for blk in rd.get("blocks", []):
                if blk.get("type") != 0:
                    continue
                for ln in blk.get("lines", []):
                    lb = ln.get("bbox")
                    if not lb:
                        continue
                    if lb[1] >= bb[1] - 1 and lb[3] <= bb[3] + 1 \
                       and lb[0] >= bb[0] - 2 and lb[2] <= bb[2] + 2:
                        for sp in ln.get("spans", []):
                            if _span_text(sp).strip():
                                szs.append(float(sp.get("size") or 10))
            msz = sorted(szs)[len(szs) // 2] if szs else 10.0
            math_regions.append({"x0": bb[0], "y0": bb[1], "x1": bb[2], "y1": bb[3],
                                 "display": True, "text": tex, "size": msz,
                                 "big": True})
            big_bands.append(bb)
    except Exception as e:
        print(f"[import] 큰 수식 복원 건너뜀: {e}")

    def _in_big(bb):
        """이 줄이 이미 '큰 수식' 밴드에 들어갔는가."""
        if not bb:
            return False
        cx = (bb[0] + bb[2]) / 2.0
        cy = (bb[1] + bb[3]) / 2.0
        for m in big_bands:
            if m[0] - 2 <= cx <= m[2] + 2 and m[1] - 2 <= cy <= m[3] + 2:
                return True
        return False

    def _ln_in_avoid(bb):
        if not avoid:
            return False
        la = max(1e-6, (bb[2] - bb[0]) * (bb[3] - bb[1]))
        for a in avoid:
            ox = min(bb[2], a[2]) - max(bb[0], a[0])
            oy = min(bb[3], a[3]) - max(bb[1], a[1])
            if ox > 0 and oy > 0 and (ox * oy) / la >= 0.5:
                return True
        return False

    for bi, blk in enumerate(rd.get("blocks", [])):
        if blk.get("type") != 0:
            continue
        blines = blk.get("lines", [])
        # 회전된 글자는 배경에 원본 모습 그대로 둔다
        if any(abs((ln.get("dir") or (1, 0))[0]) < 0.98 for ln in blines):
            continue
        # 8.20: 블록 일부를 수식으로 잘라내지 않는다. 각 줄을 아래에서
        # "줄 전체 LaTeX" 또는 "줄 전체 텍스트" 중 하나로만 결정한다.
        for ln in blines:
            # 10.4 · 표 안 글자는 칸 텍스트로 옮겨지므로 줄 경로에서는 뺀다
            if _ln_in_avoid(ln.get("bbox") or [0, 0, 0, 0]):
                continue
            # 9.3 · 큰 수식 밴드에 이미 들어간 줄은 건너뛴다.
            # (같은 글자가 LaTeX 와 텍스트 두 겹으로 나오는 것을 막는다)
            if _in_big(ln.get("bbox")):
                continue
            _sanitize_line_glyphs(ln, _gt)   # 10.1 · 확장글꼴 글자 정리
            # 독립 수식으로 확실한 줄만 줄 전체를 LaTeX로 보존한다.
            if _is_display_formula_line(ln):
                try:
                    # 9.0 · 줄 bbox 를 통째로 쓰지 않는다. 식 번호 '(3)' 을 떼고
                    #       식 본체 영역만 LaTeX 로 잡아야 폭이 부풀지 않아
                    #       옆·아래 글자와 겹치지 않는다.
                    reg = _display_region_of_line(ln, page, _gt, _rules)
                    if reg:
                        x0, y0, x1, y1, mtext, msz = reg
                    else:
                        x0, y0, x1, y1 = ln.get("bbox", [0, 0, 0, 0])
                        ls=[float(s.get("size") or 10) for s in ln.get("spans",[]) if _span_text(s).strip()]
                        mtext="".join(_span_text(s) for s in ln.get("spans", [])).strip()
                        msz=sorted(ls)[len(ls)//2] if ls else 10
                    if x1 > x0 and y1 > y0 and mtext and not _tex_is_figure_junk(mtext):
                        math_regions.append({"x0": x0, "y0": y0, "x1": x1, "y1": y1,
                                             "display": True, "text": mtext, "size": msz})
                        # 식 번호는 아래 일반 텍스트 경로가 알아서 상자로 만든다.
                        # (수식 영역 밖이므로 제거되지 않는다)
                        ln["_math_cut"] = (x0, y0, x1, y1)
                except Exception:
                    pass
            # LaTeX로 확정되지 않은 줄은 수식 글꼴이 섞여 있어도 전부 텍스트로 남긴다.
            chars = []
            fcount = {}
            # 9.0 · 이 줄에서 수식 본체를 이미 떼어냈으면 남은 글자(식 번호 등)는
            #       단어 단위로 정밀 redact 해야 배경에 잔상이 안 남는다.
            ln_has_math = [bool(ln.get("_math_cut"))]
            for sp in ln.get("spans", []):
                size = sp.get("size") or 10
                fname = (sp.get("font") or "").lower()
                flags = sp.get("flags", 0)
                bold = bool(flags & 16) or "bold" in fname
                ital = bool(flags & 2) or "italic" in fname or "oblique" in fname
                col = sp.get("color", 0)
                for ch in sp.get("chars", []):
                    c = ch.get("c") or ""
                    bb = ch.get("bbox")
                    if not c.strip() or not bb:
                        continue          # 공백은 아래 간격 판정으로 처리
                    chars.append({
                        "c": c, "b": bb,
                        "o": ch.get("origin") or (bb[0], bb[3]),
                        "sz": size, "bold": bold, "ital": ital, "col": col,
                    })
                    fcount[fname] = fcount.get(fname, 0) + 1
            if not chars:
                continue
            oys = sorted(ch["o"][1] for ch in chars)
            base_y = oys[len(oys) // 2]          # 이 줄의 베이스라인

            # 띄어쓰기(간격) 기준으로 자른다
            words, cur, prev = [], [], None
            for ch in chars:
                if prev is not None:
                    gap = ch["b"][0] - prev["b"][2]
                    thr = max(prev["sz"], ch["sz"]) * WORD_GAP_RATIO
                    dy = abs(ch["o"][1] - prev["o"][1])
                    if gap > thr or dy > ch["sz"] * SUP_SUB_DY:
                        if cur:
                            words.append(cur)
                        cur = []
                cur.append(ch)
                prev = ch
            if cur:
                words.append(cur)

            lw = []
            for wd in words:
                x0 = min(ch["b"][0] for ch in wd); y0 = min(ch["b"][1] for ch in wd)
                x1 = max(ch["b"][2] for ch in wd); y1 = max(ch["b"][3] for ch in wd)
                # 위/아래 첨자를 뺀 기준 글자 크기
                bases = [ch["sz"] for ch in wd
                         if abs(ch["o"][1] - base_y) <= max(ch["sz"] * 0.25, 1)]
                size = sum(bases) / len(bases) if bases else wd[0]["sz"]
                lw.append({"x0": x0, "y0": y0, "x1": x1, "y1": y1,
                           "size": size, "chars": wd})
            has_math = ln_has_math[0]
            if lw:
                lines.append({
                    "x0": min(w["x0"] for w in lw), "y0": min(w["y0"] for w in lw),
                    "x1": max(w["x1"] for w in lw), "y1": max(w["y1"] for w in lw),
                    "base": base_y, "words": lw, "blk": bi, "has_math": has_math,
                    "font": max(fcount, key=fcount.get) if fcount else "",
                })
    # 인접한 인라인 수식 영역을 병합 (한 수식이 여러 span 으로 쪼개진 경우)
    math_regions = _merge_math_regions(math_regions)
    math_regions = _drop_contained_math(math_regions)
    # 9.0 · 식에서 떨어져 나온 위/아래첨자 조각을 되찾아 온다.
    _absorb_orphan_scripts(lines, math_regions)
    lines = [ln for ln in lines if ln["words"]]
    for ln in lines:
        ln["x0"] = min(w["x0"] for w in ln["words"]); ln["y0"] = min(w["y0"] for w in ln["words"])
        ln["x1"] = max(w["x1"] for w in ln["words"]); ln["y1"] = max(w["y1"] for w in ln["words"])
    return lines, math_regions


def _tex_is_figure_junk(tex):
    """그림 라벨·underbrace 조각·엉켜든 글자가 '수식 밴드'로 오인된 결과인지."""
    try:
        raw = tex or ""
        # \left. / \right. 의 베어 점은 도트 지도자(목차 점선)가 아니다
        raw = re.sub(r"\\(?:left|right)\s*\.", " ", raw)
        t = re.sub(r"\\[a-zA-Z]+", " ", raw)
        toks = [x for x in t.split() if x]
        if not toks:
            return True
        # 도트 지도자: 마침표/가운뎃점만 센다 (쉼표·마이너스는 수식에 흔하다)
        dots = sum(1 for x in toks if set(x) <= {".", "\u00b7"})
        if dots >= 3:
            return True
        # 관계식/연산자/첨자가 있어야 수식이다 ('|','⟨' 같은 기호만으론 부족)
        has_rel = bool(re.search(r"[=<>\u2264\u2265\u2248\u00b1\u00d7\u00f7\u2211\u222b\u220f\u221a^_\\]", raw))
        alnum = sum(len(x) for x in toks if any(ch.isalnum() for ch in x))
        singles = sum(1 for x in toks if len(x) == 1)
        digits = any(ch.isdigit() for ch in raw)
        if not has_rel:
            # 단어 1~2개짜리 그림 라벨('Interposer','MC')도 버린다
            if alnum < 4 or singles / len(toks) > 0.55 or len(toks) <= 2:
                return True
        else:
            # \frac 이 있어도: 관계식 없고 숫자 없고 낱글자 비율 높으면
            # 두 줄이 엉켜 들어간 쓰레기('Cond. o P n re ly…')다
            has_eq = bool(re.search(r"[=<>\u2264\u2265\u2248]", raw)) or digits
            if not has_eq and toks and singles / len(toks) > 0.6:
                return True
        return False
    except Exception:
        return False


_EXT_UNICODE = {
"(": "(", ")": ")", "[": "[", "]": "]", r"\{": "{", r"\}": "}",
r"\langle": "⟨", r"\rangle": "⟩", r"\lfloor": "⌊", r"\rfloor": "⌋",
r"\lceil": "⌈", r"\rceil": "⌉",
r"\int": "∫", r"\sum": "∑", r"\prod": "∏", r"\oint": "∮",
r"\bigcup": "⋃", r"\bigcap": "⋂", r"\coprod": "∐",
r"\bigotimes": "⊗", r"\bigoplus": "⊕",
r"\sqrt": "√", "|": "|", r"\|": "‖",
}


def _sanitize_line_glyphs(ln, gtables):
    """10.1 · 확장글꼴(txex 등) 글자를 읽을 수 있는 기호로 바꾼다.

    줄(텍스트) 경로로 나가는 수식에서 'R','X','\x10', U+F8F1 같은
    글자몰漁垃圾 대신 ∫·∑·√·( 같은 진짜 기호가 보이게 한다.
    매핑에 실패한 조각(PUA/제어문자)은 지운다 — 깨진 글자로 남는 것보다 낫다.
    """
    try:
        for sp in ln.get("spans", []):
            short = (sp.get("font") or "").split("+")[-1].upper()
            if not ("CMEX" in short or "TXEX" in short or "EXTRA" in short
                    or "LMEX" in short or "MSAM" in short or "MSBM" in short):
                continue
            table = gtables.get(sp.get("font")) or gtables.get(
                (sp.get("font") or "").split("+")[-1]) or {}
            for ch in (sp.get("chars") or []):
                c = ch.get("c") or ""
                if not c:
                    continue
                gname = table.get(ord(c)) if c else None
                if not gname:
                    gname = _CMEX_STD.get(ord(c))
                if 0xE000 <= ord(c) <= 0xF8FF or ord(c) < 0x20:
                    role, tex = classify_glyph(gname)
                    ch["c"] = _EXT_UNICODE.get(tex, "") if tex else ""
                    continue
                role, tex = classify_glyph(gname)
                if tex is None:
                    continue
                ch["c"] = _EXT_UNICODE.get(tex, "")
    except Exception:
        pass





def _absorb_orphan_scripts(lines, math_regions):
    """수식 바로 옆에 홀로 남은 위/아래첨자 조각을 수식 안으로 흡수한다. (9.0)

    PDF 는 지수 'E = mc²' 의 ² 를 베이스라인이 다르다는 이유로 별개의 줄로
    떼어 놓는 경우가 많다. 그러면 식은 'E = mc' 로 잘리고 ² 만 덩그러니
    글상자로 남아, 수식 옆에 숫자가 떠다니는 것처럼 보였다.
    식 오른쪽에 딱 붙은 작은 조각만(인용 번호·식 번호는 제외) 되찾아 온다.
    """
    if not math_regions or not lines:
        return
    for r in math_regions:
        # 9.3 · 큰 수식은 첨자까지 이미 제자리에 복원돼 있다. 더 흡수하면
        #       식 뒤의 글자를 지수로 잘못 빨아들인다.
        if r.get("big"):
            continue
        rh = max(1.0, r["y1"] - r["y0"])
        rsz = float(r.get("size") or 10)
        rcy = (r["y0"] + r["y1"]) / 2
        again = True
        while again:
            again = False
            for ln in lines:
                for wd in list(ln["words"]):
                    txt = "".join(ch["c"] for ch in wd["chars"]).strip()
                    if not txt or len(txt) > 4:
                        continue
                    # [12]·(1) 같은 괄호 번호는 어떤 경우에도 흡수하지 않는다
                    if _CITE_BRACKET.match(txt):
                        continue
                    # 맨숫자 위첨자는 본문에서는 인용 번호다. 다만 '독립 수식(display)'
                    # 줄의 식 끝에 붙은 것은 인용이 아니라 지수(mc²)다.
                    if not r.get("display") and _is_citation_token(txt, superscript=True):
                        continue
                    if not re.fullmatch(r"[0-9A-Za-zα-ωΑ-Ω+\-*/=,.]{1,4}", txt):
                        continue
                    # 첨자 크기여야 한다 (본문 글자는 건드리지 않음)
                    if wd["size"] >= rsz * 0.92:
                        continue
                    gap = wd["x0"] - r["x1"]
                    if not (-1.0 <= gap <= max(2.5, rh * 0.62)):
                        continue
                    cy = (wd["y0"] + wd["y1"]) / 2
                    if not (r["y0"] - rh * 0.6 <= cy <= r["y1"] + rh * 0.6):
                        continue
                    r["x1"] = max(r["x1"], wd["x1"])
                    r["y0"] = min(r["y0"], wd["y0"]); r["y1"] = max(r["y1"], wd["y1"])
                    r["text"] = (r.get("text") or "") + \
                        (("^{" + txt + "}") if cy < rcy else ("_{" + txt + "}"))
                    ln["words"].remove(wd)
                    again = True
                    break
                if again:
                    break


def _merge_math_regions(regions):
    """같은 줄에 붙어 있는 수식 조각들을 하나의 영역으로 합친다.

    인테그랄·분수처럼 여러 span(글꼴)이 이어붙어 하나의 수식을 이루는 경우
    조각조각 잘리지 않도록, 세로가 겹치고 가로가 인접한 것들을 병합한다.
    """
    if not regions:
        return []
    regs = sorted(regions, key=lambda r: (r["y0"], r["x0"]))
    merged = []
    for r in regs:
        if not merged:
            merged.append(dict(r))
            continue
        m = merged[-1]
        # 위첨자·아래첨자는 세로로 어긋난 span이므로 단순 overlap 대신
        # 중심 거리까지 본다. 인접 조각을 통째 수식 하나로 보존한다.
        v_overlap = min(m["y1"], r["y1"]) - max(m["y0"], r["y0"])
        h_gap = r["x0"] - m["x1"]
        mh = max(1.0, m["y1"] - m["y0"])
        rh = max(1.0, r["y1"] - r["y0"])
        v_close = abs((m["y0"] + m["y1"]) / 2 - (r["y0"] + r["y1"]) / 2) <= max(mh, rh) * 1.15
        # 수식 조각 사이의 실제 조판 간격만 합친다. 허용 폭이 글자 높이보다
        # 크던 예전 값은 식 뒤의 짧은 단어까지 한 이미지로 합칠 수 있었다.
        join_gap = max(2.2, max(mh, rh) * 0.68)
        if (v_overlap > -min(mh, rh) * 0.35 or v_close) and h_gap <= join_gap \
           and not m["display"] and not r["display"] \
           and not m.get("big") and not r.get("big"):
            m["x1"] = max(m["x1"], r["x1"])
            m["y0"] = min(m["y0"], r["y0"])
            m["y1"] = max(m["y1"], r["y1"])
            m["text"] = ((m.get("text") or "") + " " + (r.get("text") or "")).strip()
            m["size"] = max(float(m.get("size") or 0),float(r.get("size") or 0)) or 10
        else:
            merged.append(dict(r))

    # ── 9.0 · 수식끼리 겹치지 않게 마무리 ──────────────────────
    # 인라인 수식 두 개가 살짝 포개지면 KaTeX 상자 두 장이 겹쳐 글자가
    # 두 겹으로 보였다. 서로 크게 겹치면 하나로 합치고, 살짝 스치면
    # 경계를 가운데에서 잘라 절대 포개지지 않게 만든다.
    merged.sort(key=lambda r: (round(r["y0"], 1), r["x0"]))
    out = []
    for r in merged:
        if out:
            p = out[-1]
            ox = min(p["x1"], r["x1"]) - max(p["x0"], r["x0"])
            oy = min(p["y1"], r["y1"]) - max(p["y0"], r["y0"])
            if ox > 0 and oy > 0:
                ra = max(1e-6, (r["x1"] - r["x0"]) * (r["y1"] - r["y0"]))
                pa = max(1e-6, (p["x1"] - p["x0"]) * (p["y1"] - p["y0"]))
                if p.get("big") or r.get("big"):
                    # 9.3 · 큰 수식은 이미 완성된 하나의 식이다. 다른 조각과
                    #       합치거나 경계를 자르면 LaTeX 가 깨진다.
                    if (ox * oy) / min(ra, pa) > 0.5:
                        keep = p if p.get("big") else r
                        drop = r if keep is p else p
                        if keep is r:
                            out[-1] = r
                        continue
                    out.append(r)
                    continue
                if (ox * oy) / min(ra, pa) > 0.5:       # 사실상 같은 식 → 합친다
                    p["x0"] = min(p["x0"], r["x0"]); p["y0"] = min(p["y0"], r["y0"])
                    p["x1"] = max(p["x1"], r["x1"]); p["y1"] = max(p["y1"], r["y1"])
                    p["text"] = ((p.get("text") or "") + " " + (r.get("text") or "")).strip()
                    p["display"] = bool(p.get("display") or r.get("display"))
                    continue
                if p["x1"] > r["x0"]:                    # 가로로 살짝 겹침 → 반씩 양보
                    cut = (p["x1"] + r["x0"]) / 2.0
                    p["x1"] = min(p["x1"], cut); r = dict(r); r["x0"] = max(r["x0"], cut)
                    if r["x1"] - r["x0"] < 0.5:
                        continue
        out.append(r)
    return out


_SUP_MAP = str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾", "0123456789+-=()")
_SUB_MAP = str.maketrans("₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎", "0123456789+-=()")
_LATEX_SYMBOLS = {
    "∞": r"\infty", "∂": r"\partial", "∇": r"\nabla", "∆": r"\Delta",
    "∫": r"\int", "∮": r"\oint", "∑": r"\sum", "∏": r"\prod",
    "√": r"\sqrt", "≈": r"\approx", "≃": r"\simeq", "≅": r"\cong",
    "≠": r"\ne", "≤": r"\le", "≥": r"\ge", "±": r"\pm", "∓": r"\mp",
    "×": r"\times", "÷": r"\div", "·": r"\cdot", "∝": r"\propto",
    "∈": r"\in", "∉": r"\notin", "⊂": r"\subset", "⊃": r"\supset",
    "∪": r"\cup", "∩": r"\cap", "→": r"\to", "←": r"\leftarrow",
    "↔": r"\leftrightarrow", "⇒": r"\Rightarrow", "⇔": r"\Leftrightarrow",
    "∀": r"\forall", "∃": r"\exists", "∥": r"\parallel", "⊥": r"\perp",
    "ℏ": r"\hbar", "ℓ": r"\ell", "ℜ": r"\Re", "ℑ": r"\Im",
    "ˆ": r"\int", "◦": r"^{\circ}", "□": r"\square",
}
_LATEX_GREEK = {
    # 같은 모양 다른 코드포인트도 함께 (µ MICRO SIGN, Ω OHM SIGN, ∆ INCREMENT)
    "\u00b5":"mu", "\u2126":"Omega", "\u2206":"Delta", "\u03d5":"phi", "\u03f5":"epsilon",
    "α":"alpha","β":"beta","γ":"gamma","δ":"delta","ε":"epsilon","ζ":"zeta",
    "η":"eta","θ":"theta","ι":"iota","κ":"kappa","λ":"lambda","μ":"mu","ν":"nu",
    "ξ":"xi","ο":"omicron","π":"pi","ρ":"rho","σ":"sigma","τ":"tau","υ":"upsilon",
    "φ":"phi","χ":"chi","ψ":"psi","ω":"omega","Γ":"Gamma","Δ":"Delta","Θ":"Theta",
    "Λ":"Lambda","Ξ":"Xi","Π":"Pi","Σ":"Sigma","Φ":"Phi","Ψ":"Psi","Ω":"Omega",
}


def _pdf_text_to_latex(text):
    """PDF에서 얻은 유니코드 수식을 KaTeX가 읽을 수 있는 LaTeX로 정리."""
    t = (text or "").replace("\u00a0", " ").replace("−", "-")
    t = re.sub(r"[\u200b-\u200f\u2060\ufeff]", "", t).strip()
    already = bool(re.search(r"\\(frac|left|right|begin|int|sum|sqrt|cases)", t))
    # 유니코드 위/아래첨자를 연속 묶음으로 변환
    t = re.sub(r"[⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾]+", lambda m: "^{" + m.group().translate(_SUP_MAP) + "}", t)
    t = re.sub(r"[₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎]+", lambda m: "_{" + m.group().translate(_SUB_MAP) + "}", t)
    for ch, name in _LATEX_GREEK.items():
        t = t.replace(ch, "\\" + name + " ")
    for ch, val in _LATEX_SYMBOLS.items():
        t = t.replace(ch, val + " ")
    t = t.replace("⟨", r"\langle ").replace("⟩", r"\rangle ")
    # 14.3 · 이미 조립된 cases 의 & 는 열 구분 문자이므로 이스케이프하지 않는다.
    if already:
        t = re.sub(r"(?<!\\)([%#])", r"\\\1", t)
    else:
        t = re.sub(r"(?<!\\)([%#&])", r"\\\1", t)
    if len(t)>1 and t.startswith("$") and t.endswith("$"): t=t[1:-1]
    t = re.sub(r"\s+", " ", t).strip()
    # 9.3 · 유니코드 첨자가 연달아 나오면 x^{a}^{b} 가 되어 KaTeX 가
    #       '이중 위첨자' 오류를 낸다(∫Ldt = 4.8 fb^-1 처럼). 하나로 합친다.
    return _tidy_latex(t)

def _mark_justify(lines, page_w):
    """양쪽 정렬(justify) 줄 판별 — 페이지·칼럼·문단 단위.

    PyMuPDF 가 문단을 줄별로 잘게 쪼개도 동작하도록 블록이 아닌
    페이지 전체의 칼럼(왼쪽/오른쪽/꽉찬) 기준으로 오른쪽 도달을 본다.
    문단 마지막 줄(다음 줄 들여쓰기/큰 세로 간격/없음)은 제외.
    """
    if not lines:
        return
    mid = page_w / 2.0
    cols = {"L": [], "R": [], "F": []}
    for ln in lines:
        cx = (ln["x0"] + ln["x1"]) / 2.0
        if ln["x1"] - ln["x0"] > page_w * 0.6:
            cols["F"].append(ln)
        elif cx < mid:
            cols["L"].append(ln)
        else:
            cols["R"].append(ln)
    all_sz = [w["size"] for ln in lines for w in ln["words"]]
    avg_sz = (sum(all_sz) / len(all_sz)) if all_sz else 10.0
    for col in cols.values():
        if len(col) < 2:
            continue
        col.sort(key=lambda l: l["y0"])
        c0 = min(l["x0"] for l in col)
        c1 = max(l["x1"] for l in col)
        cw = max(1e-6, c1 - c0)
        for i, ln in enumerate(col):
            reach = (ln["x1"] - c0) >= cw * 0.97
            nxt = col[i + 1] if i + 1 < len(col) else None
            if nxt is None:
                last = True
            else:
                gap_v = nxt["y0"] - ln["y1"]
                indent = (nxt["x0"] - c0) > max(2.0, avg_sz * 0.8)
                last = indent or gap_v > avg_sz * 0.9
            ln["just"] = bool(reach and not last)
            ln["j_right"] = c1


def _font_map(name):
    """PDF 원본 글꼴 이름 → 모양/자간이 가장 비슷한 앱 글꼴."""
    n = (name or "").lower()
    if any(k in n for k in ("times", "roman", "serif", "georgia", "garamond",
                            "batang", "myeongjo", "myungjo", "바탕", "명조")):
        return "times"
    if any(k in n for k in ("courier", "mono", "consolas")):
        return "mono"
    if any(k in n for k in ("helv", "arial", "roboto", "sans")):
        return "inter"   # 산세리프는 Inter 로 → 원본과 자폭 유사
    return "noto"


def _line_aligns(lines, page_w):
    """줄 정렬 판별 — 논문 2단 레이아웃 대응.

    왼쪽/오른쪽 기둥(칼럼)을 나눠 각 칼럼 안에서 여백을 계산해야
    2단 논문의 가운데/오른쪽 정렬이 엉뚱하게 잡히지 않는다.
    """
    if not lines:
        return
    mid = page_w / 2.0
    groups = {}
    for ln in lines:
        cx = (ln["x0"] + ln["x1"]) / 2.0
        lw = ln["x1"] - ln["x0"]
        if lw > page_w * 0.6:
            g = "full"
        elif cx < mid:
            g = "L"
        else:
            g = "R"
        groups.setdefault(g, []).append(ln)
    for gl in groups.values():
        cx0 = min(l["x0"] for l in gl)
        cx1 = max(l["x1"] for l in gl)
        cw = max(1e-6, cx1 - cx0)
        tol = max(6.0, cw * 0.035)
        for ln in gl:
            lw = ln["x1"] - ln["x0"]
            lg = ln["x0"] - cx0
            rg = cx1 - ln["x1"]
            if lw >= cw * 0.92:
                ln["align"] = "left"          # 꽉 찬 줄
                continue
            mid_c = abs(lg - rg) <= max(4.0, cw * 0.05)
            mid_p = abs(ln["x0"] - (page_w - ln["x1"])) <= max(6.0, page_w * 0.035)
            if (mid_c or mid_p) and lw < page_w * 0.85:
                ln["align"] = "center"
            else:
                # 10.4 · '오른쪽 정렬' 표시는 내보내지 않는다.
                #   가져온 글자는 절대좌표로 배치되는데, 정렬 표시가 남으면
                #   편집할 때 글이 오른쪽으로 붙어 어색해진다.
                ln["align"] = "left"


def _word_html(wd, base_sz, base_y):
    """단어 문자열 → HTML. 같은 스타일의 연속 문자는 태그 하나로 묶는다. (13.4 · 첨자 역치 정밀화)"""
    groups = []
    for ch in wd["chars"]:
        mode = ""
        c = ch.get("c", "")
        if not c:
            continue
        oy = ch["o"][1] if "o" in ch else base_y
        sz = ch.get("sz", base_sz)
        if base_sz > 0:
            if sz < base_sz * 0.88 or abs(oy - base_y) >= base_sz * 0.14:
                if oy < base_y - base_sz * 0.10:
                    mode = "sup"
                elif oy > base_y + base_sz * 0.10:
                    mode = "sub"
        key = (ch["bold"], ch["ital"], mode, ch["col"] or 0)
        if groups and groups[-1][0] == key:
            groups[-1][1] += c
        else:
            groups.append([key, c])
    out = []
    for (bold, ital, mode, col), t in groups:
        esc = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if col:
            esc = f'<span style="color:#{col & 0xFFFFFF:06x}">{esc}</span>'
        if mode:
            esc = f"<{mode}>{esc}</{mode}>"
        if bold:
            esc = f"<b>{esc}</b>"
        if ital:
            esc = f"<i>{esc}</i>"
        out.append(esc)
    return "".join(out)



def detect_page_figures_and_tables(page, pw, ph):
    """13.4 · 피규어(도표/그림) 및 표(Table) 영역 정밀 감지.
    2단 논문 칼럼 경계 자동 인식 및 피규어/표/수식 중복 추출 방지.
    """
    blocks = page.get_text("blocks")
    drawings = page.get_drawings()
    images = page.get_image_info()

    tab_caps = []
    fig_caps = []
    for b in blocks:
        text = b[4].strip()
        if b[3] > ph * 0.97 or b[1] < ph * 0.03:
            continue
        is_fig_cap = bool(re.match(r"^(?:Fig(?:ure|\.)?|FIG(?:URE|\.)?)\s*\d+\s*[:.]", text, re.IGNORECASE))
        is_tab_cap = bool(re.match(r"^(?:Tab(?:le|\.)?|TABLE)\s*[\d\w.\-]+\s*[:.]", text, re.IGNORECASE)) or bool(re.match(r"^TABLE\s+[IVXLCDM\d]+", text))

        if is_tab_cap:
            tab_caps.append({"rect": pymupdf.Rect(b[:4]), "text": text})
        elif is_fig_cap:
            fig_caps.append({"rect": pymupdf.Rect(b[:4]), "text": text})

    figure_boxes = []
    table_boxes = []
    avoid = []

    # 1. Figures: 캡션 상단에 위치한 그래픽 요소(드로잉/이미지) 클러스터링 (칼럼 경계 보호)
    for fc in fig_caps:
        cr = fc["rect"]
        is_full_width = (cr.width > pw * 0.6) or (cr.x0 < pw * 0.3 and cr.x1 > pw * 0.7)
        col_left = 0 if is_full_width else (0 if cr.x1 <= pw * 0.55 else pw * 0.45)
        col_right = pw if is_full_width else (pw * 0.55 if cr.x1 <= pw * 0.55 else pw)

        fig_cluster = None
        for d in drawings:
            r = d.get("rect")
            if not r: continue
            if abs(r.y1 - r.y0) <= 2.5 and (r.x1 - r.x0) < 60: continue
            if r.y1 <= cr.y0 + 4 and r.y0 >= cr.y0 - 450:
                if r.x0 >= col_left - 10 and r.x1 <= col_right + 10:
                    dr_box = pymupdf.Rect(r.x0, r.y0, r.x1, r.y1 if r.y1 > r.y0 else r.y0 + 1)
                    fig_cluster = dr_box if fig_cluster is None else (fig_cluster | dr_box)
        for im in images:
            ir = pymupdf.Rect(im["bbox"])
            if ir.y1 <= cr.y0 + 4 and ir.y0 >= cr.y0 - 450:
                if ir.x0 >= col_left - 10 and ir.x1 <= col_right + 10:
                    fig_cluster = ir if fig_cluster is None else (fig_cluster | ir)

        if fig_cluster and not fig_cluster.is_empty:
            for b in blocks:
                br = pymupdf.Rect(b[:4])
                if br.y1 <= cr.y0 + 2 and br.y0 >= fig_cluster.y0 - 8:
                    if br.x0 >= col_left - 10 and br.x1 <= col_right + 10:
                        if br.intersects(pymupdf.Rect(fig_cluster.x0 - 20, fig_cluster.y0 - 15, fig_cluster.x1 + 25, fig_cluster.y1 + 15)):
                            if len(b[4].strip().split('\n')) <= 3 and len(b[4].strip()) < 180:
                                fig_cluster = fig_cluster | br

            r = pymupdf.Rect(max(0, fig_cluster.x0 - 2), max(0, fig_cluster.y0 - 2),
                             min(pw, fig_cluster.x1 + 2), min(ph, fig_cluster.y1 + 2))
            if r.width >= 25 and r.height >= 25:
                figure_boxes.append({"rect": r, "caption": fc["text"]})
                avoid.append([r.x0, r.y0, r.x1, r.y1])

    # 2. Booktabs Tables: Table 캡션 하단에 위치한 가로 규칙선(Rules) 기반 감지
    for tc in tab_caps:
        tcr = tc["rect"]
        is_full_width = (tcr.width > pw * 0.6) or (tcr.x0 < pw * 0.3 and tcr.x1 > pw * 0.7)
        col_left = 0 if is_full_width else (0 if tcr.x1 <= pw * 0.55 else pw * 0.45)
        col_right = pw if is_full_width else (pw * 0.55 if tcr.x1 <= pw * 0.55 else pw)

        tab_lines = []
        for d in drawings:
            r = d.get("rect")
            if not r: continue
            if (r.x1 - r.x0) >= 40 and abs(r.y1 - r.y0) <= 3.0:
                if r.y0 >= tcr.y1 - 5 and r.y0 <= tcr.y1 + 400:
                    if r.x0 >= col_left - 10 and r.x1 <= col_right + 10:
                        tab_lines.append(r)
        if len(tab_lines) >= 2:
            tbox = pymupdf.Rect(min(r.x0 for r in tab_lines) - 2,
                                tcr.y0 - 2,
                                max(r.x1 for r in tab_lines) + 2,
                                max(r.y1 for r in tab_lines) + 2)
            if not any(tbox.intersects(pymupdf.Rect(a)) for a in avoid):
                table_boxes.append({"rect": tbox, "type": "booktabs", "caption": tc["text"]})
                avoid.append([tbox.x0, tbox.y0, tbox.x1, tbox.y1])

    # 3. Grid Tables: 명시적 표 캡션이 있거나 진짜 격자 구조를 가진 경우만 추출
    try:
        tbl_finder = page.find_tables()
        for t in (tbl_finder.tables if tbl_finder else []):
            tb = pymupdf.Rect(t.bbox)
            if tb.width < 30 or tb.height < 12: continue
            if any(tb.intersects(pymupdf.Rect(a)) for a in avoid): continue
            grid = t.extract() or []
            filled = [c for r in grid for c in (r or []) if c and str(c).strip()]
            has_tab_cap = any(abs(tc["rect"].y1 - tb.y0) < 60 or abs(tc["rect"].y0 - tb.y1) < 40 for tc in tab_caps)
            if len(grid) >= 2 and len(grid[0] or []) >= 2 and (has_tab_cap or (len(filled) >= 6 and len(filled) / max(1, len(grid)*len(grid[0])) >= 0.4)):
                table_boxes.append({"rect": tb, "type": "grid", "table": t})
                avoid.append([tb.x0, tb.y0, tb.x1, tb.y1])
    except Exception:
        pass

    # 4. Standalone raster images (캡션이 없는 독립 사진/다이어그램)
    for im in images:
        ir = pymupdf.Rect(im["bbox"])
        if ir.width >= 120 and ir.height >= 90:
            if not any(ir.intersects(pymupdf.Rect(a)) for a in avoid):
                figure_boxes.append({"rect": ir, "caption": "Image"})
                avoid.append([ir.x0, ir.y0, ir.x1, ir.y1])

    return figure_boxes, table_boxes, avoid


def _pdf_one_page(doc, pno, on_page_done, cache=None,
                  target_w=PAGE_W, target_h=PAGE_H):
    """1쪽 변환: 추출 → 글자 제거 → 배경 저장 → 단어 상자. (13.4 고정밀 피규어/표/수식 분리)"""
    try:
        page = doc[pno]
        pw, ph = page.rect.width, page.rect.height
        if pw <= 0 or ph <= 0:
            return {"id": _imp_uid("p"), "els": [], "tables": []}
        sc = min(target_w / pw, target_h / ph)
        offx = (target_w - pw * sc) / 2

        def X(v):
            return _px(v * sc + offx)

        def Y(v):
            return _px(v * sc)

        fig_table_els = []
        tbl_els, tbl_meta = [], []

        # ── 13.4 · 피규어 및 표 정밀 감지 ──────────────────────────
        figure_boxes, table_boxes, avoid = detect_page_figures_and_tables(page, pw, ph)

        # 1) 피규어 영역을 고화질(200 DPI) 독립 이미지 요소로 분리
        for fb in figure_boxes:
            try:
                r = fb["rect"]
                pm = page.get_pixmap(clip=r, dpi=200, alpha=False)
                furl = _save_pixmap_bg(pm)
                pm = None
                if furl:
                    fig_table_els.append({
                        "type": "image", "id": _imp_uid("i"), "url": furl,
                        "x": X(r.x0), "y": Y(r.y0),
                        "w": max(10, _px(r.width * sc)), "h": max(10, _px(r.height * sc)),
                        "imported": 1, "locked": True,
                    })
            except Exception as e:
                print(f"[import] 피규어 추출 실패: {e}")

        # 2) 표(Table) 처리: Booktabs는 고화질 이미지로, Grid 표는 앱 표 요소로 추출
        for tb_info in table_boxes:
            r = tb_info["rect"]
            if tb_info.get("type") == "booktabs":
                try:
                    pm = page.get_pixmap(clip=r, dpi=200, alpha=False)
                    turl = _save_pixmap_bg(pm)
                    pm = None
                    if turl:
                        fig_table_els.append({
                            "type": "image", "id": _imp_uid("i"), "url": turl,
                            "x": X(r.x0), "y": Y(r.y0),
                            "w": max(10, _px(r.width * sc)), "h": max(10, _px(r.height * sc)),
                            "imported": 1, "locked": True,
                        })
                except Exception as e:
                    print(f"[import] 북탭 표 추출 실패: {e}")
            elif tb_info.get("type") == "grid":
                tab = tb_info["table"]
                try:
                    tb = tab.bbox
                    grid = tab.extract() or []
                    ex = []
                    for row in tab.rows:
                        for c in (row.cells or []):
                            if c: ex.extend([c[0], c[2]])
                    if ex:
                        ex.sort()
                        edges = [ex[0]]
                        for v in ex[1:]:
                            if v - edges[-1] > 2.5: edges.append(v)
                        if len(edges) >= 3:
                            cw = [max(6.0, edges[k + 1] - edges[k]) for k in range(len(edges) - 1)]
                            ch = []
                            for row in tab.rows:
                                cs = [c for c in (row.cells or []) if c]
                                ch.append(max(6.0, (max(c[3] for c in cs) - min(c[1] for c in cs))) if cs else 10.0)
                            tid = _imp_uid("tb")
                            tbl_meta.append({"id": tid, "x": X(tb[0]), "y": Y(tb[1]),
                                             "cw": [_px(w * sc) for w in cw],
                                             "ch": [_px(h * sc) for h in ch],
                                             "color": "#d1d5db", "lw": 1, "bg": 1})
                            for ri, row in enumerate(tab.rows):
                                for ci, c in enumerate(row.cells or []):
                                    if not c: continue
                                    txt = ""
                                    if ri < len(grid) and ci < len(grid[ri] or []):
                                        txt = str(grid[ri][ci] or "").strip()
                                    if not txt: continue
                                    xc = (c[0] + c[2]) / 2.0
                                    col = min(range(len(edges) - 1),
                                              key=lambda k: abs((edges[k] + edges[k + 1]) / 2 - xc))
                                    html = (txt.replace("&", "&amp;").replace("<", "&lt;")
                                                .replace(">", "&gt;").replace("\n", "<br>"))
                                    fs = round(max(6.0, min(15.0, min(9.0, ch[ri] * 0.66) * sc)), 1)
                                    tbl_els.append({
                                        "type": "text", "id": _imp_uid("t"), "html": html,
                                        "x": X(edges[col]) + 3, "y": Y(c[1]) + 2,
                                        "w": max(14, _px((edges[col + 1] - edges[col]) * sc) - 6),
                                        "h": max(10, _px(ch[ri] * sc) - 3),
                                        "fontSize": fs, "align": "left",
                                        "tbl": {"tid": tid, "r": ri, "c": col},
                                    })
                except Exception:
                    pass

        lines, math_regions = _pdf_page_lines(page, avoid=avoid)
        # 표/피규어 영역에 밴드가 이미 만들어졌다면 한 번 더 정리
        if avoid and math_regions:
            def _mr_in_avoid(m):
                cx, cy = (m["x0"] + m["x1"]) / 2.0, (m["y0"] + m["y1"]) / 2.0
                return any(a[0] - 1 <= cx <= a[2] + 1 and a[1] - 1 <= cy <= a[3] + 1
                           for a in avoid)
            math_regions = [m for m in math_regions if not _mr_in_avoid(m)]

        # ── 반복 구조 예측 캐시 ──
        sig = None
        if cache is not None:
            try:
                head = "".join(
                    f"{round(ln['x0'])},{round(ln['y0'])}," +
                    "".join(ch["c"] for wd in ln["words"] for ch in wd["chars"])
                    for ln in lines[:80])
                math_sig = "|".join((r.get("text") or "") for r in math_regions)
                sig = hashlib.md5(
                    (head + "|M:" + math_sig + f"|{len(lines)}|{round(pw)}x{round(ph)}")
                    .encode("utf-8")).hexdigest()
                if tbl_meta or fig_table_els:
                    sig = None
                hit = cache.get(sig)
                if hit:
                    els = []
                    for el in hit:
                        ne = dict(el)
                        ne["id"] = _imp_uid("t" if ne["type"] == "text" else "m" if ne["type"] == "latex" else "i")
                        els.append(ne)
                    if on_page_done:
                        on_page_done()
                    return {"id": _imp_uid("p"), "els": els, "tables": []}
            except Exception:
                sig = None

        _line_aligns(lines, pw)
        _mark_justify(lines, pw)

        # 중복 제거: 같은 자리의 같은 글자·인용 번호가 두 겹인 PDF 대응
        seen = {}
        for ln in lines:
            kept = []
            for wd in ln["words"]:
                text = "".join(ch["c"] for ch in wd["chars"]).strip()
                box = (wd["x0"], wd["y0"], wd["x1"], wd["y1"])
                dup = any(max(abs(box[i] - old[i]) for i in range(4)) <= 1.25
                          for old in seen.get(text, []))
                if dup:
                    continue
                seen.setdefault(text, []).append(box)
                kept.append(wd)
            ln["words"] = kept
        lines = [ln for ln in lines if ln["words"]]

        # 수식 영역 안에 남은 일반 텍스트 조각을 한 번 더 제거
        if math_regions:
            for ln in lines:
                keep = []
                for wd in ln["words"]:
                    wa = max(1e-6, (wd["x1"] - wd["x0"]) * (wd["y1"] - wd["y0"]))
                    inside = False
                    for r in math_regions:
                        ox = min(wd["x1"], r["x1"] + 0.4) - max(wd["x0"], r["x0"] - 0.4)
                        oy = min(wd["y1"], r["y1"] + 0.6) - max(wd["y0"], r["y0"] - 0.6)
                        if ox > 0 and oy > 0 and (ox * oy) / wa >= 0.45:
                            inside = True
                            break
                    if not inside:
                        keep.append(wd)
                ln["words"] = keep
            lines = [ln for ln in lines if ln["words"]]
            for ln in lines:
                ln["x0"] = min(w["x0"] for w in ln["words"])
                ln["y0"] = min(w["y0"] for w in ln["words"])
                ln["x1"] = max(w["x1"] for w in ln["words"])
                ln["y1"] = max(w["y1"] for w in ln["words"])

        # ── ⓪ 수식은 LaTeX 요소 하나로 만든다 ──
        math_els = []
        for mr in math_regions:
            try:
                rx0, ry0, rx1, ry1 = mr["x0"], mr["y0"], mr["x1"], mr["y1"]
                if rx1 <= rx0 or ry1 <= ry0:
                    continue
                latex = _pdf_text_to_latex(mr.get("text") or "")
                if not latex or not _latex_is_sane(latex):
                    continue
                display = bool(mr.get("display"))
                pad_x, pad_y = ((0.6, 0.8) if display else (0.0, 0.4))
                x0 = max(0.0, rx0 - pad_x); y0 = max(0.0, ry0 - pad_y)
                x1 = min(pw, rx1 + pad_x); y1 = min(ph, ry1 + pad_y)
                fs = round(max(6.0, min(52.0, float(mr.get("size") or 10) * sc)), 1)
                ink_h = _px((y1 - y0) * sc)
                hh = max(8, ink_h if ink_h >= _px(fs * 0.75) else _px(fs * (1.15 if display else 1.05)))
                math_els.append({
                    "type": "latex", "id": _imp_uid("m"), "latex": latex,
                    "x": X(x0), "y": Y(y0),
                    "w": max(8, _px((x1 - x0) * sc)), "h": hh,
                    "fontSize": fs,
                    "displayMath": 1 if display else 0,
                    "inkW": max(8, _px((x1 - x0) * sc)), "inkH": max(8, ink_h),
                    "imported": 1, "locked": True,
                })
            except Exception as e:
                print(f"[import] {pno+1}쪽 LaTeX 변환 실패: {e}")

        # ── ① 배경: 지우고 래스터화 ──
        bg_els = []
        try:
            has_img = False
            try:
                has_img = len(page.get_images(full=False)) > 0
            except Exception:
                pass
            need_bg = True
            try:
                if not has_img and not page.get_drawings():
                    need_bg = False
            except Exception:
                pass
            if not need_bg:
                raise _NoBackground

            # 피규어 및 표 영역 지우기 (배경 잔상/중복 방지)
            for a in avoid:
                try:
                    page.add_redact_annot(pymupdf.Rect(
                        a[0] - 0.3, a[1] - 0.5, a[2] + 0.3, a[3] + 0.5))
                except Exception:
                    pass
            # 수식 영역 지우기
            for mr in math_regions:
                try:
                    page.add_redact_annot(pymupdf.Rect(
                        mr["x0"] - 0.3, mr["y0"] - 0.5,
                        mr["x1"] + 0.3, mr["y1"] + 0.5))
                except Exception:
                    pass
            # 텍스트 줄 지우기
            for ln in lines:
                if ln.get("has_math"):
                    for wd in ln["words"]:
                        page.add_redact_annot(pymupdf.Rect(
                            wd["x0"] - 0.3, wd["y0"] - 0.5,
                            wd["x1"] + 0.3, wd["y1"] + 0.5))
                else:
                    page.add_redact_annot(pymupdf.Rect(
                        ln["x0"] - 0.3, ln["y0"] - 0.6,
                        ln["x1"] + 0.3, ln["y1"] + 0.6))

            kw = {"images": getattr(pymupdf, "PDF_REDACT_IMAGE_PIXELS",
                                   pymupdf.PDF_REDACT_IMAGE_NONE)}
            touched = getattr(pymupdf, "PDF_REDACT_LINE_ART_REMOVE_IF_TOUCHED", None)
            if touched is not None:
                kw["graphics"] = touched
            try:
                page.apply_redactions(**kw)
            except TypeError:
                try:
                    page.apply_redactions(images=kw["images"])
                except Exception:
                    page.apply_redactions(images=pymupdf.PDF_REDACT_IMAGE_NONE)

            bg_url = None
            try:
                dpi = BG_DPI_CUR
                area_in2 = (pw / 72.0) * (ph / 72.0)
                if area_in2 * dpi * dpi > BG_PX_CUR:
                    dpi = max(96, int((BG_PX_CUR / max(1e-6, area_in2)) ** 0.5))
                pm = page.get_pixmap(dpi=dpi, alpha=False)
                try:
                    zoom = dpi / 72.0
                    wipes = []
                    for mr in math_regions:
                        wipes.append((mr["x0"] - 1.8, mr["y0"] - 1.8,
                                      mr["x1"] + 1.8, mr["y1"] + 1.8))
                    for a in (avoid or []):
                        wipes.append((a[0] - 1.0, a[1] - 1.0, a[2] + 1.0, a[3] + 1.0))
                    try:
                        for rr in page_rules(page):
                            wipes.append((rr[0] - 0.8, rr[1] - 1.4,
                                          rr[2] + 0.8, rr[3] + 1.4))
                    except Exception:
                        pass
                    for x0, y0, x1, y1 in wipes:
                        pr = pymupdf.Rect(x0 * zoom, y0 * zoom, x1 * zoom, y1 * zoom)
                        try:
                            pm.set_rect(pr, (255, 255, 255))
                        except Exception:
                            pass
                except Exception:
                    pass
                bg_url = _save_pixmap_bg(pm)
                pm = None
            except Exception as e:
                print(f"[import] {pno+1}쪽 래스터 실패: {e}")

            if not bg_url and not has_img:
                try:
                    svg = page.get_svg_image(text_as_path=True)
                    if svg and len(svg) < 6 * 1024 * 1024:
                        bg_url = _save_import_svg(svg)
                except Exception as e:
                    print(f"[import] {pno+1}쪽 SVG 실패: {e}")

            if bg_url:
                bg_els.append({
                    "type": "image", "id": _imp_uid("i"), "url": bg_url,
                    "x": _px(offx), "y": 0,
                    "w": _px(pw * sc), "h": _px(ph * sc),
                    "isBg": 1, "locked": True,
                })
        except _NoBackground:
            pass
        except Exception as e:
            print(f"[import] {pno+1}쪽 배경 실패: {e}")

        # ── 문단 단위 글상자 (단어는 내부 절대좌표 스팬) ──
        text_els = []
        groups, by_id = [], {}
        for ln in lines:
            k = ln.get("blk", id(ln))
            if k not in by_id:
                by_id[k] = []
                groups.append(by_id[k])
            by_id[k].append(ln)

        for glines in groups:
            x0 = min(l["x0"] for l in glines); y0 = min(l["y0"] for l in glines)
            x1 = max(l["x1"] for l in glines); y1 = max(l["y1"] for l in glines)
            aligns = [l.get("align", "left") for l in glines]
            align = max(set(aligns), key=aligns.count)
            pad = 1.0
            x = X(x0 - pad)
            y = Y(y0 - pad)
            w = max(10, _px((x1 - x0 + pad * 2) * sc))
            h = max(10, _px((y1 - y0 + pad * 2) * sc))
            el_fs = round(max(
                5.0, max(wd["size"] for l in glines for wd in l["words"]) * sc), 1)
            fcnt = {}
            for l in glines:
                fn = l.get("font") or ""
                if fn:
                    fcnt[fn] = fcnt.get(fn, 0) + 1
            el_font = _font_map(max(fcnt, key=fcnt.get) if fcnt else "")

            spans = []
            for ln in glines:
                fs_ln = max(wd["size"] for wd in ln["words"])
                vs = max(2.0, fs_ln * sc * 0.15)
                ltop = _px((ln["y0"] - y0 + pad) * sc - vs)
                lh = max(6, _px((ln["y1"] - ln["y0"]) * sc) + int(round(vs * 2)))
                just = ln.get("just", False)
                jattr = ' data-j="1"' if just else ""
                jpos = None
                if just and len(ln["words"]) > 1:
                    wpts = [w["x1"] - w["x0"] for w in ln["words"]]
                    first_x = ln["words"][0]["x0"]
                    gap_pts = max(0.5, (ln["j_right"] - first_x - sum(wpts))
                                  / (len(wpts) - 1))
                    jpos = []
                    cx_ = first_x
                    for k_ in range(len(wpts)):
                        jpos.append(cx_)
                        cx_ += wpts[k_] + gap_pts
                for wi, wd in enumerate(ln["words"]):
                    wx = jpos[wi] if jpos is not None else wd["x0"]
                    lx = _px((wx - x0 + pad) * sc)
                    fs = round(max(5.0, wd["size"] * sc), 1)
                    inner = _word_html(wd, wd["size"], ln["base"])
                    spans.append(
                        f'<span data-fs="{fs}"{jattr} style="position:absolute;'
                        f'left:{lx}px;top:{ltop}px;line-height:{lh}px;'
                        f'font-size:{fs}px;white-space:nowrap">'
                        f'{inner}<i class="zsp"> </i></span>')
            text_els.append({
                "type": "text", "id": _imp_uid("t"),
                "html": "".join(spans),
                "x": x, "y": y, "w": w, "h": h,
                "fontSize": el_fs, "font": el_font,
                "align": align, "tight": 1,
            })

        els = bg_els + fig_table_els + math_els + text_els + tbl_els
        if sig is not None:
            cache[sig] = [dict(e) for e in els]
        return {"id": _imp_uid("p"), "els": els,
                "tables": tbl_meta if tbl_meta else []}
    except Exception as e:
        print(f"[import] {pno+1}쪽 변환 실패, 건너: {e}")
        return {"id": _imp_uid("p"), "els": [], "tables": []}
    finally:
        if on_page_done:
            try:
                on_page_done()
            except Exception:
                pass


def _pdf_to_pages(data, on_page=None, target_w=PAGE_W, target_h=PAGE_H):
    """PDF → 페이지 목록 (청크 스트리밍 방식).

    · 문서를 4쪽 단위 청크로 나눠 열고→변환→닫는다.
      → 메모리가 '전체 문서'가 아니라 '쪽 1개' 수준에서만 논다.
      → 저사양 서버에서도 큰 PDF 가 죽지 않는다.
    · 문서 사본을 두 개 열던 옛 방식을 버려 속도도 2배 가깝게 개선.
    · 쪽이 끝날 때마다 배경은 즉시 디스크 저장 + 진행 콜백 호출.
    """
    import pymupdf

    d0 = pymupdf.open(stream=data, filetype="pdf")
    total = min(d0.page_count, IMPORT_MAX_PAGES)
    d0.close()

    pages = []
    state = {"done": 0}

    def _bump():
        state["done"] += 1
        if on_page:
            try:
                on_page(state["done"], total)
            except Exception:
                pass

    CHUNK = 4
    for start in range(0, total, CHUNK):
        doc = pymupdf.open(stream=data, filetype="pdf")
        try:
            for pno in range(start, min(start + CHUNK, total)):
                pages.append(_pdf_one_page(doc, pno, _bump,
                                           target_w=target_w, target_h=target_h))
        finally:
            doc.close()          # 청크 끝나면 즉시 메모리 반납
    return pages


def _pdf_to_pages_safe(data, on_page=None, target_w=PAGE_W, target_h=PAGE_H):
    """한 페이지가 깨져도 문서 전체가 실패하지 않도록 감싼다."""
    try:
        return _pdf_to_pages(data, on_page=on_page,
                             target_w=target_w, target_h=target_h)
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"[import] 전체 변환 실패 → 글자만 뽑아 재시도: {e}")
    # 최후 수단: 글자만 단순하게 뽑는다
    import pymupdf
    doc = pymupdf.open(stream=data, filetype="pdf")
    pages = []
    try:
        for pno in range(min(doc.page_count, IMPORT_MAX_PAGES)):
            page = doc[pno]
            pw, ph = page.rect.width, page.rect.height
            if pw <= 0 or ph <= 0:
                continue
            sc = min(target_w / pw, target_h / ph)
            offx = (target_w - pw * sc) / 2
            els = []
            try:
                blocks = page.get_text("blocks")
            except Exception:
                blocks = []
            for bl in blocks:
                try:
                    x0, y0, x1, y1, txt = bl[0], bl[1], bl[2], bl[3], bl[4]
                    if not (txt or "").strip():
                        continue
                    html = (txt.replace("&", "&amp;").replace("<", "&lt;")
                               .replace(">", "&gt;").replace("\n", " "))
                    els.append({
                        "type": "text", "id": _imp_uid("t"), "html": html,
                        "x": _px(x0 * sc + offx) - 2, "y": _px(y0 * sc) - 2,
                        "w": max(24, _px((x1 - x0) * sc) + 5),
                        "h": max(16, _px((y1 - y0) * sc) + 5),
                        "fontSize": 14, "fit": 1,
                    })
                except Exception:
                    continue
            pages.append({"id": _imp_uid("p"), "els": els, "tables": []})
    finally:
        doc.close()
    return pages


def _pdf_one_page_safe(doc, pno, target_w=PAGE_W, target_h=PAGE_H):
    """최소 변환: 배경 없이 단어 상자만. 자식 프로세스가 죽었을 때 쓰는 안전 모드."""
    try:
        page = doc[pno]
        pw, ph = page.rect.width, page.rect.height
        if pw <= 0 or ph <= 0:
            return {"id": _imp_uid("p"), "els": [], "tables": []}
        sc = min(target_w / pw, target_h / ph)
        offx = (target_w - pw * sc) / 2
        lines, math_regions = _pdf_page_lines(page)
        _line_aligns(lines, pw)
        _mark_justify(lines, pw)
        els = []
        for mr in math_regions:
            latex=_pdf_text_to_latex(mr.get("text") or "")
            if not latex: continue
            x=_px(mr["x0"]*sc+offx); y=_px(mr["y0"]*sc)
            fs=round(max(6.0,min(52.0,float(mr.get("size") or 10)*sc)),1)
            w=max(8,_px((mr["x1"]-mr["x0"])*sc))
            ink_h=_px((mr["y1"]-mr["y0"])*sc)
            h=max(8, ink_h if ink_h>=_px(fs*0.75) else _px(fs*1.1))
            els.append({"type":"latex","id":_imp_uid("m"),"latex":latex,
                        "x":x,"y":y,"w":w,"h":h,
                        "fontSize":fs,
                        "displayMath":1 if mr.get("display") else 0,
                        "inkW":w,"inkH":max(8,ink_h),
                        "imported":1,"locked":True})
        # 수식 영역과 겹친 텍스트는 안전 모드에서도 면적 기준으로 제거 (9.0)
        for ln in lines:
            keep=[]
            for wd in ln["words"]:
                wa=max(1e-6,(wd["x1"]-wd["x0"])*(wd["y1"]-wd["y0"]))
                hit=False
                for r in math_regions:
                    ox=min(wd["x1"],r["x1"])-max(wd["x0"],r["x0"])
                    oy=min(wd["y1"],r["y1"])-max(wd["y0"],r["y0"])
                    if ox>0 and oy>0 and (ox*oy)/wa>=0.45: hit=True; break
                if not hit: keep.append(wd)
            ln["words"]=keep
        lines=[ln for ln in lines if ln["words"]]
        for ln in lines:
            ln["x0"]=min(w_["x0"] for w_ in ln["words"]); ln["y0"]=min(w_["y0"] for w_ in ln["words"])
            ln["x1"]=max(w_["x1"] for w_ in ln["words"]); ln["y1"]=max(w_["y1"] for w_ in ln["words"])
        # 안전 모드에서도 겹친 인용/글자 레이어는 하나만 남긴다.
        seen_words={}
        for ln in lines:
            keep=[]
            for wd in ln["words"]:
                tx="".join(ch["c"] for ch in wd["chars"]).strip()
                bb=(wd["x0"],wd["y0"],wd["x1"],wd["y1"])
                if any(max(abs(bb[i]-o[i]) for i in range(4))<=1.25 for o in seen_words.get(tx,[])):
                    continue
                seen_words.setdefault(tx,[]).append(bb); keep.append(wd)
            ln["words"]=keep
        lines=[ln for ln in lines if ln["words"]]
        groups, by_id = [], {}
        for ln in lines:
            k = ln.get("blk", id(ln))
            if k not in by_id:
                by_id[k] = []
                groups.append(by_id[k])
            by_id[k].append(ln)
        for glines in groups:
            x0 = min(l["x0"] for l in glines); y0 = min(l["y0"] for l in glines)
            x1 = max(l["x1"] for l in glines); y1 = max(l["y1"] for l in glines)
            aligns = [l.get("align", "left") for l in glines]
            align = max(set(aligns), key=aligns.count)
            pad = 1.0
            x = _px((x0 - pad) * sc + offx)
            y = _px((y0 - pad) * sc)
            w = max(10, _px((x1 - x0 + pad * 2) * sc))
            h = max(10, _px((y1 - y0 + pad * 2) * sc))
            # 대표 크기·글꼴 (5.2 렌더 유지 + 글꼴 매핑만 부가)
            el_fs = round(max(
                5.0, max(wd["size"] for l in glines for wd in l["words"]) * sc), 1)
            fcnt = {}
            for l in glines:
                fn = l.get("font") or ""
                if fn:
                    fcnt[fn] = fcnt.get(fn, 0) + 1
            el_font = _font_map(max(fcnt, key=fcnt.get) if fcnt else "")

            # ── 5.2 그대로: 단어는 원본 좌표 절대 배치, 픽셀 크기,
            #    세로는 줄 전체 잉크 기준(높이 균일), 끝에 일반 공백 한 칸.
            # 양쪽 정렬 판별은 페이지 단위(_mark_justify)에서 완료됨.

            spans = []
            for ln in glines:
                # 줄 상자에 상하 숨통: 브라우저 글꼴의 디센더(g,y,p 꼬리)와
                # 어센더가 줄 잉크보다 깊어도 잘리지·겹치지 않게.
                fs_ln = max(wd["size"] for wd in ln["words"])
                vs = max(2.0, fs_ln * sc * 0.15)
                ltop = _px((ln["y0"] - y0 + pad) * sc - vs)
                lh = max(6, _px((ln["y1"] - ln["y0"]) * sc) + int(round(vs * 2)))
                just = ln.get("just", False)
                jattr = ' data-j="1"' if just else ""
                # 양쪽 정렬 줄: 서버가 PDF 좌표만으로 균등 간격을 구해 굽는다.
                # (클라이언트는 이 줄을 일절 수정 안 함 → 환경 차이·키 흔들림 제로)
                jpos = None
                if just and len(ln["words"]) > 1:
                    wpts = [w["x1"] - w["x0"] for w in ln["words"]]
                    first_x = ln["words"][0]["x0"]
                    gap_pts = max(0.5, (ln["j_right"] - first_x - sum(wpts))
                                  / (len(wpts) - 1))
                    jpos = []
                    cx_ = first_x
                    for k_ in range(len(wpts)):
                        jpos.append(cx_)
                        cx_ += wpts[k_] + gap_pts
                for wi, wd in enumerate(ln["words"]):
                    wx = jpos[wi] if jpos is not None else wd["x0"]
                    lx = _px((wx - x0 + pad) * sc)
                    fs = round(max(5.0, wd["size"] * sc), 1)
                    inner = _word_html(wd, wd["size"], ln["base"])
                    spans.append(
                        f'<span data-fs="{fs}"{jattr} style="position:absolute;'
                        f'left:{lx}px;top:{ltop}px;line-height:{lh}px;'
                        f'font-size:{fs}px;white-space:nowrap">'
                        f'{inner}<i class="zsp"> </i></span>')
            els.append({
                "type": "text", "id": _imp_uid("t"),
                "html": "".join(spans),
                "x": x, "y": y, "w": w, "h": h,
                "fontSize": el_fs, "font": el_font,
                "align": align, "tight": 1,
            })
        return {"id": _imp_uid("p"), "els": els, "tables": []}
    except Exception as e:
        print(f"[import] {pno+1}쪽 최소 변환 실패: {e}")
        return {"id": _imp_uid("p"), "els": [], "tables": []}


# ── 격리 변환: 변환은 '별도 프로세스'에서 ────────────────────
# 자식이 세그폴트/OOM 으로 죽어도 메인 서버는 절대 죽지 않는다.
# 죽으면 안전 모드(텍스트만)로 해당 청크를 재시도하고,
# 그것도 죽으면 빈 쪽으로 채워 '무조건 끝까지' 완료한다.
# 메모리/동시성은 배포 시 RAM 에 맞춰 apply.sh 가 주입한다
# (SDY_IMP_MAX_CONCURRENT, SDY_IMP_CHILD_MEM_MB).
# 기본값은 저사양 박스 보호용으로 보수적으로 유지.
def _imp_env_int(name, default):
    try:
        v = int(os.environ.get(name, ""))
        return v if v >= 1 else default
    except (TypeError, ValueError):
        return default


IMP_CHILD_MEM_MB = _imp_env_int("SDY_IMP_CHILD_MEM_MB", 2000)   # 자식 메모리 상한
IMP_MAX_CONCURRENT = _imp_env_int("SDY_IMP_MAX_CONCURRENT", 2)  # 동시 변환 잡 상한
IMP_SLICE = int(os.environ.get("SDY_IMP_SLICE", "8") or 8)      # 보관 배치 크기
_imp_sem = threading.Semaphore(IMP_MAX_CONCURRENT)
# 전역 '자식 프로세스' 상한. 잡별 동시성(IMP_MAX_CONCURRENT)만으로는 여러 잡의
# 청크가 겹쳐(예: 잡 3 × 청크 3) 자식 총합이 RAM 을 넘었다. apply.sh 가 12GB
# 박스 기준 SDY_IMP_MAX_CHUNKS(=워커 예산 ÷ 자식 메모리, 예: 5×1GB)를 주입하고,
# 모든 잡의 청크가 이 세마포를 공유해 실제 동시 자식 수를 한 번에 잠근다.
IMP_MAX_CHUNKS = _imp_env_int("SDY_IMP_MAX_CHUNKS", max(2, IMP_MAX_CONCURRENT * 2))
_imp_chunk_sem = threading.BoundedSemaphore(IMP_MAX_CHUNKS)
# 현재 가져오기 잡을 몇 개나 처리 중인지(세마포가 잡혀 있는 수). 여러 잡이
# 동시에 청크를 띄울 때 자식 총합이 RAM 을 넘지 않도록 문서당 병렬을 줄인다.
_imp_active = 0
_imp_active_lock = threading.Lock()
IMP_CHUNK = 4
IMP_CHUNK_TIMEOUT = 240       # 청크 하나 당 허용 시간(초)


def _chunk_pages(src, pnos, out_path, safe, total=1,
                 target_w=PAGE_W, target_h=PAGE_H):
    """자식 프로세스 본문. 결과(쪽 목록)를 파일로 남기고 끝난다."""
    try:
        try:
            import resource
            lim = IMP_CHILD_MEM_MB * 1024 * 1024
            resource.setrlimit(resource.RLIMIT_AS, (lim, lim))
        except Exception:
            pass
        global BG_DPI_CUR, BG_PX_CUR
        BG_DPI_CUR, BG_PX_CUR = _bg_tier(total)
        import pymupdf
        doc = pymupdf.open(src)          # 경로는 스트림 복사 없이 디스크에서 읽음
        pages = []
        try:
            cache = {}
            # 예전엔 큰 파일이면 '쪽마다' 문서를 새로 열었다.
            # 373쪽짜리 원서에선 이 재파싱이 변환 시간을 몇 배로 늘린다.
            # 자식 프로세스가 슬라이스(8쪽)마다 새로 뜨므로 메모리는 이미
            # 그 단위에서 반납된다 → 한 번만 열고 재사용해도 안전하다.
            for pno in pnos:
                pages.append(
                    _pdf_one_page_safe(doc, pno, target_w, target_h) if safe
                    else _pdf_one_page(doc, pno, None, cache, target_w, target_h)
                )
        finally:
            doc.close()
        with open(out_path, "w", encoding="utf-8") as fp:
            json.dump(pages, fp, ensure_ascii=False)
        os._exit(0)
    except BaseException as e:
        try:
            with open(out_path, "w", encoding="utf-8") as fp:
                json.dump({"error": f"{type(e).__name__}: {e}"}, fp)
        except Exception:
            pass
        os._exit(3)


def _run_chunk_start(src, pnos, safe, total=1,
                     target_w=PAGE_W, target_h=PAGE_H):
    """자식 프로세스 시작만 (병렬 실행용). 전역 청크 세마포로 자식 총합을 잠근다."""
    import multiprocessing
    out = os.path.join(IMG_DIR, f"chunk_{uuid.uuid4().hex}.json")
    ctx = multiprocessing.get_context("fork")
    _imp_chunk_sem.acquire()
    try:
        proc = ctx.Process(target=_chunk_pages,
                           args=(src, pnos, out, safe, total, target_w, target_h))
        proc.start()
    except BaseException:
        _imp_chunk_sem.release()
        raise
    return proc, out


def _run_chunk_collect(proc, out):
    """자식 수확. 죽으면 None (호출측이 재시도)."""
    try:
        proc.join(IMP_CHUNK_TIMEOUT)
        if proc.is_alive():
            proc.kill()
            proc.join()
        code = proc.exitcode
        try:
            if code == 0 and os.path.exists(out):
                with open(out, encoding="utf-8") as fp:
                    data_ = json.load(fp)
                if isinstance(data_, list):
                    return data_
        finally:
            try:
                os.remove(out)
            except Exception:
                pass
        return None
    finally:
        # 시작 시 잡은 전역 청크 세마포를 항상 여기서 반납한다(예외 포함).
        _imp_chunk_sem.release()


def _run_chunk(src, pnos, safe, total=1,
               target_w=PAGE_W, target_h=PAGE_H):
    return _run_chunk_collect(*_run_chunk_start(
        src, pnos, safe, total, target_w, target_h
    ))


def _imp_convert_pdf(src, jid, target_w=PAGE_W, target_h=PAGE_H):
    """PDF 전체를 청크 격리 변환으로. 서버는 절대 죽지 않는다."""
    import pymupdf
    d0 = pymupdf.open(src)
    total = min(d0.page_count, IMPORT_MAX_PAGES)
    d0.close()
    # 총 쪽수를 '시작하자마자' 알려 준다.
    # 예전엔 첫 묶음(8쪽)이 끝나야 total 이 잡혀서, 그전까지 진행률이
    # 엉뚱한 값(0/1 쪽)으로 계산돼 막대가 뒤로 튀어 보였다.
    _imp_job(jid, page=0, total=total)

    sz = os.path.getsize(src)
    # 단일 문서 안의 청크 병렬성. 동시에 여러 '문서' 잡이 돌 수 있으므로
    # (세마포 IMP_MAX_CONCURRENT), 지금 활성 잡 수에 비례해 문서당 상한을
    # 낮춰 자식 프로세스 총합이 RAM 을 넘지 않게 한다.
    with _imp_active_lock:
        active = _imp_active
    cpu = os.cpu_count() or 2
    # 활성 잡 1개 → 최대 3, 2개 → 2, 3개 → 1 (대/중형 파일은 더 낮춤)
    share = max(1, (IMP_MAX_CONCURRENT + 1) // max(1, active))
    cap = max(1, min(cpu, IMP_MAX_CONCURRENT, share, 3))
    if sz > 60 * 1024 * 1024:
        CONC = 1                      # 거대 파일: 직렬로 안정적으로
    elif sz > 25 * 1024 * 1024:
        CONC = min(2, cap)
    else:
        CONC = cap
    pending = list(range(0, total, IMP_SLICE))
    running = []
    by_start = {}
    done_n = 0

    def start_one(s0):
        pnos = list(range(s0, min(s0 + IMP_SLICE, total)))
        proc, out = _run_chunk_start(src, pnos, False, total,
                                     target_w, target_h)
        running.append([s0, pnos, proc, out])

    for k in range(min(CONC, len(pending))):
        start_one(pending[k])
    nxt = min(CONC, len(pending))
    import time as _t
    while running:
        while all(r[2].exitcode is None for r in running):
            _t.sleep(0.12)
        for r in list(running):
            if r[2].exitcode is None:
                continue
            running.remove(r)
            s0, pnos, proc, out = r
            got = _run_chunk_collect(proc, out)
            if got is None or len(got) != len(pnos):
                got = _run_chunk(src, pnos, True, total,
                                 target_w, target_h)
            if got is None or len(got) != len(pnos):
                got = [{"id": _imp_uid("p"), "els": [], "tables": []} for _ in pnos]
            by_start[s0] = got
            done_n += len(pnos)
            _imp_job(jid, page=done_n, total=total)
            if nxt < len(pending):
                start_one(pending[nxt])
                nxt += 1
    pages = []
    for s0 in sorted(by_start):
        pages.extend(by_start[s0])
    if not pages:
        pages = [{"id": _imp_uid("p"), "els": [], "tables": []}]
    return pages


def _emu_px(v):
    """EMU(914400 = 1inch) → 화면 픽셀 (96dpi 기준)"""
    try:
        return v / 914400.0 * 96.0
    except Exception:
        return 0


def _docx_to_pages(data, target_w=PAGE_W, target_h=PAGE_H):
    """Word(.docx) → 페이지 목록. 문단을 위에서 아래로 흘려 배치한다."""
    import docx
    from docx.shared import RGBColor

    f = docx.Document(io.BytesIO(data))
    margin_x, margin_y = 64, 60
    max_y = target_h - margin_y
    width = target_w - margin_x * 2

    pages, els = [], []
    y = margin_y

    def new_page():
        nonlocal els, y
        pages.append({"id": _imp_uid("p"), "els": els})
        els, y = [], margin_y

    # 문서에 들어있는 그림들 (문단 순서대로 꺼내 쓴다)
    images = []
    for rel in f.part.rels.values():
        if "image" in rel.reltype:
            try:
                images.append(rel.target_part.blob)
            except Exception:
                pass
    img_i = 0

    for para in f.paragraphs:
        text = (para.text or "").strip()
        style = (para.style.name or "").lower() if para.style is not None else ""

        # 그림이 들어있는 문단
        if "graphic" in para._p.xml or "<w:drawing" in para._p.xml:
            if img_i < len(images):
                url = _img_data_url(images[img_i]); img_i += 1
                if url:
                    try:
                        im = Image.open(io.BytesIO(images[img_i - 1]))
                        iw, ih = im.size
                    except Exception:
                        iw, ih = 400, 300
                    sc = min(width / iw, 1.0)
                    w, h = _px(iw * sc), _px(ih * sc)
                    if y + h > max_y:
                        new_page()
                    els.append({"type": "image", "id": _imp_uid("i"), "url": url,
                                "x": margin_x, "y": _px(y), "w": w, "h": h})
                    y += h + 12
            if not text:
                continue

        if not text:
            y += 14                      # 빈 줄
            continue

        # 제목 스타일이면 크게
        fs = 16
        if "heading 1" in style or "title" in style:
            fs = 28
        elif "heading 2" in style:
            fs = 23
        elif "heading 3" in style:
            fs = 20

        html = ""
        for run in para.runs:
            t = (run.text or "")
            if not t:
                continue
            t = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            col = None
            try:
                if run.font.color is not None and run.font.color.rgb is not None:
                    col = str(run.font.color.rgb)
            except Exception:
                col = None
            if col and col != "000000":
                t = f'<span style="color:#{col}">{t}</span>'
            if run.bold:
                t = f"<b>{t}</b>"
            if run.italic:
                t = f"<i>{t}</i>"
            if run.underline:
                t = f"<u>{t}</u>"
            html += t
        if not html.strip():
            html = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if "heading" in style or "title" in style:
            html = f"<b>{html}</b>" if "<b>" not in html else html

        # 글자 수에 따라 높이 추정 (한 줄에 들어가는 글자 수 기준)
        per_line = max(10, int(width / (fs * 0.62)))
        lines = max(1, (len(text) + per_line - 1) // per_line)
        h = _px(lines * fs * 1.55) + 14

        if y + h > max_y:
            new_page()

        els.append({"type": "text", "id": _imp_uid("t"), "html": html,
                    "x": margin_x, "y": _px(y), "w": width, "h": h,
                    "fontSize": fs})
        y += h + 4

    # 표는 줄글로 옮긴다
    for tb in f.tables:
        for row in tb.rows:
            cells = [c.text.strip().replace("&", "&amp;").replace("<", "&lt;") for c in row.cells]
            line = "  |  ".join([c for c in cells if c])
            if not line:
                continue
            h = 34
            if y + h > max_y:
                new_page()
            els.append({"type": "text", "id": _imp_uid("t"), "html": line,
                        "x": margin_x, "y": _px(y), "w": width, "h": h,
                        "fontSize": 15})
            y += h + 2

    if els or not pages:
        pages.append({"id": _imp_uid("p"), "els": els})
    return pages


# ── 가져오기 비동기 잡 (디스크 기반) ─────────────────────────
# 잡 상태를 메모리가 아닌 디스크에 저장한다.
# → 서버가 재시작되든, 프로세스가 여러 개이든, 어떤 경우에도
#   상태 조회가 일관되게 동작한다 ("없는 작업" 문제 원천 차단)
IMP_JOB_TTL = 900          # 잡 결과 보관 시간(초)


def _imp_job_path(jid):
    safe = re.sub(r"[^0-9a-zA-Z_\-]", "", jid or "")[:40]
    return os.path.join(JOBS_DIR, f"{safe}.json")


def _imp_job(jid, **kw):
    p = _imp_job_path(jid)
    try:
        cur = {}
        if os.path.exists(p):
            try:
                with open(p, encoding="utf-8") as fp:
                    cur = json.load(fp)
            except Exception:
                cur = {}
        cur.update(kw)
        cur["ts"] = time.time()
        tmp = "%s.tmp.%s" % (p, uuid.uuid4().hex[:8])
        with open(tmp, "w", encoding="utf-8") as fp:
            json.dump(cur, fp, ensure_ascii=False)
        os.replace(tmp, p)
    except Exception as e:
        print("[import] 잡 저장 실패:", e)


def _imp_job_load(jid):
    p = _imp_job_path(jid)
    try:
        if not os.path.exists(p):
            return None
        if time.time() - os.path.getmtime(p) > IMP_JOB_TTL:
            os.remove(p)
            return None
        with open(p, encoding="utf-8") as fp:
            return json.load(fp)
    except Exception:
        return None


def _imp_mark_stopped(p, why):
    try:
        with open(p, encoding="utf-8") as fp:
            j = json.load(fp)
        if j.get("status") == "working":
            j["status"] = "error"
            j["error"] = why
            tmp = "%s.tmp.%s" % (p, uuid.uuid4().hex[:8])
            with open(tmp, "w", encoding="utf-8") as fp:
                json.dump(j, fp, ensure_ascii=False)
            os.replace(tmp, p)
    except Exception:
        pass


def _imp_sweep_dead():
    """서버 시작 시점의 '변환 중' 잡은 전부 죽은 것(재시작으로 워커 소멸)."""
    try:
        for fn in os.listdir(JOBS_DIR):
            if fn.endswith(".json"):
                _imp_mark_stopped(os.path.join(JOBS_DIR, fn),
                                  "서버가 재시작되어 변환이 중단되었습니다. 다시 가져오기를 시도해 주세요")
    except Exception:
        pass


def _imp_detect_kind(raw, ext):
    """파일 내용(매직 바이트)으로 실제 형식을 판별해 정확한 안내를 준다."""
    if raw[:4] == b"%PDF":
        return "pdf", None
    if raw[:4] == b"PK\x03\x04":
        if ext in ("docx", "docm"):
            return "word", None
        return None, ("이 파일은 PDF 가 아닙니다(압축/Word 계열). "
                      "PDF 또는 .docx 로 저장해서 올려주세요")
    if raw[:4] == b"\xd0\xcf\x11\xe0":
        return None, ("HWP 또는 옛 .doc 형식은 지원하지 않습니다. "
                      "한글/워드에서 PDF 로 저장한 뒤 올려주세요")
    if raw[:8].startswith(b"\x89PNG") or raw[:3] == b"\xff\xd8\xff":
        return None, ("이미지 파일입니다. 사진 버튼으로 삽입하거나, "
                      "문서라면 PDF 로 저장해서 올려주세요")
    if ext in ("docx", "docm"):
        return "word", None
    if ext == "doc":
        return None, "옛 .doc 형식은 지원하지 않습니다. .docx 로 저장 후 올려주세요"
    if ext == "pdf":
        return "pdf", None
    return None, "PDF 또는 Word(.docx) 파일만 가져올 수 있습니다"


def _imp_worker(jid, src, name, kind):
    """백그라운드 변환. 본문은 디스크(docfile)에만 두고 잡엔 메타만.
    동시에 여러 잡이 메모리를 쌓지 않도록 세마포어로 제한한다."""
    global _imp_active
    import traceback
    _imp_sem.acquire()
    with _imp_active_lock:
        _imp_active += 1
    try:
        if kind == "pdf":
            size_preset = _pdf_size_preset(src)
            target_w, target_h = _preset_dims(size_preset)
            pages = _imp_convert_pdf(src, jid, target_w, target_h)
        else:
            _imp_job(jid, page=0, total=1)
            with open(src, "rb") as fp:
                word_data = fp.read()
            size_preset = _docx_size_preset(word_data)
            target_w, target_h = _preset_dims(size_preset)
            pages = _docx_to_pages(word_data, target_w, target_h)
            _imp_job(jid, page=1, total=1)

        if not pages:
            pages = [{"id": _imp_uid("p"), "els": [], "tables": []}]

        # 그림은 서버에 파일로 두고 문서에는 주소만 넣는다 (용량 문제 해결)
        for pg_ in pages:
            for el in pg_.get("els", []):
                u = el.get("url") or ""
                if el.get("type") == "image" and u.startswith("data:"):
                    el["url"] = _save_import_img(u)

        total_els = sum(len(p.get("els", [])) for p in pages)
        title = name.rsplit(".", 1)[0][:80] or "가져온 문서"

        # 대용량 문서: 브라우저 저장소 대신 서버에 본문 보관 (gzip)
        doc_ref = None
        try:
            import gzip as _gz
            total_n = len(pages)
            for s0 in range(0, total_n, IMP_SLICE):
                sp = os.path.join(DOCS_DIR, f"{jid}.s{s0}.gz.tmp")
                with _gz.open(sp, "wt", encoding="utf-8") as fp:
                    json.dump({"ok": True, "pages": pages[s0:s0 + IMP_SLICE],
                               "total": total_n}, fp, ensure_ascii=False)
                os.replace(sp, os.path.join(DOCS_DIR, f"{jid}.s{s0}.gz"))
            mp = os.path.join(DOCS_DIR, f"{jid}.meta.json.tmp")
            with open(mp, "w", encoding="utf-8") as fp:
                json.dump({"total": total_n, "sizePreset": size_preset,
                           "version": time.time()}, fp)
            os.replace(mp, os.path.join(DOCS_DIR, f"{jid}.meta.json"))
            doc_ref = jid
        except Exception as e:
            print("[import] docfile 저장 실패:", e)

        print(f"[import] {kind} '{name}' → {len(pages)}쪽 / 요소 {total_els}개")
        if doc_ref is None or len(pages) <= 60:
            # 작거나 디스크 보관 실패 시엔 잡에 본문도 태움(옛 프런트 호환)
            _imp_job(jid, status="done", kind=kind, title=title, pages=pages,
                     count=total_els, sizePreset=size_preset,
                     docRef=doc_ref, total=len(pages))
        else:
            _imp_job(jid, status="done", kind=kind, title=title,
                     count=total_els, sizePreset=size_preset,
                     docRef=doc_ref, total=len(pages))
        _notify_add("convert_done", "문서 변환이 끝났어요 ✓",
                    f"{title} · {len(pages)}쪽을 노트로 준비했습니다.",
                    dedupe=f"import-done:{jid}", meta={"job": jid, "pages": len(pages)})
    except MemoryError:
        traceback.print_exc()
        _imp_job(jid, status="error",
                 error="서버 메모리가 부족합니다. 쪽수가 적은 파일로 나눠서 올려주세요")
        _notify_add("error", "문서 변환을 마치지 못했어요",
                    f"{name} · 메모리가 부족합니다. 파일을 나눠 다시 시도해 주세요.",
                    dedupe=f"import-error:{jid}")
    except Exception as e:
        tb = traceback.format_exc()
        print(f"[import] 변환 실패: {e}\n{tb}")
        msg = str(e) or e.__class__.__name__
        low = msg.lower()
        if "password" in low or "encrypt" in low:
            msg = "비밀번호로 잠긴 PDF 입니다. 잠금을 푼 뒤 다시 올려주세요"
        elif "find_tables" in low or "no attribute" in low:
            msg = ("서버 라이브러리가 오래되었습니다. apply.sh 를 다시 실행해 "
                   "pymupdf 를 갱신해 주세요 (" + msg + ")")
        _imp_job(jid, status="error", error=f"변환 실패: {msg}",
                 detail=tb[-800:])
        _notify_add("error", "문서 변환을 마치지 못했어요",
                    f"{name} · {msg}", dedupe=f"import-error:{jid}")
    finally:
        with _imp_active_lock:
            _imp_active = max(0, _imp_active - 1)
        try:
            _imp_sem.release()
        except Exception:
            pass
        try:
            # 원본은 재변환(자가치유)을 위해 보관, 하루 지난 것은 정리
            if src and os.path.exists(src):
                import shutil
                shutil.move(src, os.path.join(DOCS_DIR, f"{jid}.src"))
            now = time.time()
            for fn in os.listdir(DOCS_DIR):
                if fn.endswith(".src"):
                    p = os.path.join(DOCS_DIR, fn)
                    if now - os.path.getmtime(p) > 86400:
                        os.remove(p)
        except Exception:
            pass


@app.route("/api/import/upload", methods=["POST"])
def import_upload():
    """큰 파일을 청크로 나눠 받는 조립소. 크기 제한 없이 업로드 가능."""
    uid = re.sub(r"[^0-9a-zA-Z_\-]", "", (request.form.get("uploadId") or ""))[:40]
    if not uid:
        return jsonify({"ok": False, "error": "uploadId 없음"}), 400
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "청크 없음"}), 400
    try:
        chunk = int(request.form.get("chunk", "0"))
        total = max(1, int(request.form.get("total", "1")))
    except ValueError:
        return jsonify({"ok": False, "error": "청크 번호 오류"}), 400
    if chunk >= total or total > 20000:
        return jsonify({"ok": False, "error": "청크 범위 오류"}), 400
    # 1시간 넘은 찌꺼기 업로드 정리
    try:
        now = time.time()
        for fn in os.listdir(UPLOAD_DIR):
            p = os.path.join(UPLOAD_DIR, fn)
            if now - os.path.getmtime(p) > 3600:
                os.remove(p)
    except Exception:
        pass
    part = os.path.join(UPLOAD_DIR, f"{uid}.part")
    try:
        with open(part, "wb" if chunk == 0 else "ab") as fp:
            fp.write(f.read())
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    done = chunk + 1 >= total
    if done:
        os.replace(part, os.path.join(UPLOAD_DIR, f"{uid}.bin"))
    return jsonify({"ok": True, "done": done, "chunk": chunk, "total": total})


@app.route("/api/import/doc", methods=["POST"])
def import_doc():
    """PDF / Word → 편집 가능한 노트. 잡 번호를 바로 돌려준다(비동기).

    큰 파일은 /api/import/upload 으로 청크 조립 후 uploadId 로 요청하고,
    작은 파일은 예전처럼 file 로 바로 보내도 된다. 크기 제한 없음.
    """
    uid = re.sub(r"[^0-9a-zA-Z_\-]", "", (request.form.get("uploadId") or ""))[:40]
    if uid:
        src = os.path.join(UPLOAD_DIR, f"{uid}.bin")
        if not os.path.exists(src):
            return jsonify({"ok": False, "error": "업로드가 완성되지 않았습니다"}), 400
        name = request.form.get("name") or "document"
    elif "file" in request.files:
        f = request.files["file"]
        name = f.filename or "document"
        uid = uuid.uuid4().hex[:12]
        src = os.path.join(UPLOAD_DIR, f"{uid}.bin")
        f.save(src)
    else:
        return jsonify({"ok": False, "error": "파일이 없습니다"}), 400

    try:
        size = os.path.getsize(src)
    except OSError:
        size = 0
    if not size:
        try:
            os.remove(src)
        except Exception:
            pass
        return jsonify({"ok": False, "error": "빈 파일입니다"}), 400

    with open(src, "rb") as fp:
        head = fp.read(16)

    ext = (name.rsplit(".", 1)[-1] if "." in name else "").lower()
    kind, err = _imp_detect_kind(head, ext)
    if err:
        try:
            os.remove(src)
        except Exception:
            pass
        return jsonify({"ok": False, "error": err}), 400

    # PDF 는 시작 전에 미리 열어본다 — 죽은 파일이면 여기서 정확히 안내
    if kind == "pdf":
        import pymupdf
        try:
            d = pymupdf.open(src)
            locked = d.is_encrypted and not d.authenticate("")
            n_pages = d.page_count
            d.close()
            if locked:
                os.remove(src)
                return jsonify({"ok": False,
                                "error": "비밀번호로 잠긴 PDF 입니다. 잠금을 푼 뒤 다시 올려주세요"}), 400
            if n_pages <= 0:
                os.remove(src)
                return jsonify({"ok": False, "error": "페이지가 없는 PDF 입니다"}), 400
        except Exception as e:
            os.remove(src)
            return jsonify({"ok": False,
                            "error": ("PDF 를 열 수 없습니다. 파일이 깨졌거나 "
                                      "다른 형식일 수 있어요. PDF 로 다시 저장해서 "
                                      f"올려주세요 ({str(e) or e.__class__.__name__})")}), 400

    jid = uuid.uuid4().hex[:12]
    _imp_job(jid, status="working", page=0, total=0)
    _notify_add("convert", "문서 변환을 시작했어요", f"{name} 파일을 읽고 있습니다.",
                dedupe=f"import-start:{jid}", meta={"job": jid, "kind": kind})
    threading.Thread(target=_imp_worker, args=(jid, src, name, kind),
                     daemon=True).start()
    return jsonify({"ok": True, "job": jid})


@app.route("/api/import/reconv", methods=["POST"])
def import_reconv():
    """비어 있게 변환된 슬라이스를 안전 모드로 재변환해 덮어쓴다 (자가치유)."""
    d = request.get_json(silent=True) or {}
    ref = re.sub(r"[^0-9a-zA-Z_\-]", "", d.get("ref") or "")
    try:
        s0 = int(d.get("from") or 0)
        total = int(d.get("total") or 0)
    except ValueError:
        return jsonify({"ok": False, "error": "인자 오류"}), 400
    srcp = os.path.join(DOCS_DIR, f"{ref}.src")
    if not ref or not os.path.exists(srcp):
        return jsonify({"ok": False, "error": "원본 없음"}), 404
    size_preset = "a4_portrait"
    try:
        with open(os.path.join(DOCS_DIR, f"{ref}.meta.json"), encoding="utf-8") as fp:
            size_preset = json.load(fp).get("sizePreset", size_preset)
    except Exception:
        pass
    target_w, target_h = _preset_dims(size_preset)
    pnos = list(range(s0, min(s0 + IMP_SLICE, total or s0 + IMP_SLICE)))
    proc, out = _run_chunk_start(srcp, pnos, True,
                                 total or s0 + IMP_SLICE, target_w, target_h)
    got = _run_chunk_collect(proc, out)
    if got is None:
        return jsonify({"ok": False, "error": "재변환 실패"}), 500
    import gzip as _gz
    sp = os.path.join(DOCS_DIR, f"{ref}.s{s0}.gz")
    tmp = "%s.tmp.%s" % (sp, uuid.uuid4().hex[:8])
    with _gz.open(tmp, "wt", encoding="utf-8") as fp:
        json.dump({"ok": True, "pages": got, "total": total or len(got)},
                  fp, ensure_ascii=False)
    os.replace(tmp, sp)
    return jsonify({"ok": True, "pages": got})


@app.route("/api/import/status", methods=["GET"])
def import_status():
    """변환 잡 진행/결과 조회 (디스크 기반 — 재시작/다중프로세스 무관).
    완료된 잡은 한 번 조회하면 폐기한다."""
    jid = request.args.get("id", "")
    out = _imp_job_load(jid)
    if not out:
        return jsonify({"ok": False, "error": "없는 작업입니다", "status": "gone"}), 404
    if out.get("status") == "working":
        # 5분 동안 진행 업데이트가 없으면 워커가 죽은 것 → 멈춤으로 확정
        try:
            if time.time() - os.path.getmtime(_imp_job_path(jid)) > 300:
                _imp_mark_stopped(_imp_job_path(jid),
                                  "변환이 중단되었습니다. 다시 가져오기를 시도해 주세요")
                out = _imp_job_load(jid) or out
        except Exception:
            pass
    if out.get("status") == "done":
        try:
            os.remove(_imp_job_path(jid))
        except Exception:
            pass
        return jsonify({"ok": True, "status": "done",
                        "kind": out.get("kind"), "title": out.get("title"),
                        "pages": out.get("pages"), "count": out.get("count"),
                        "total": out.get("total"),
                        "sizePreset": out.get("sizePreset"),
                        "docRef": out.get("docRef")})
    if out.get("status") == "error":
        try:
            os.remove(_imp_job_path(jid))
        except Exception:
            pass
        return jsonify({"ok": False, "status": "error",
                        "error": out.get("error", "변환에 실패했습니다"),
                        "detail": out.get("detail", "")})
    return jsonify({"ok": True, "status": "working",
                    "page": out.get("page", 0), "total": out.get("total", 0)})


# ============ 읽는 동안 배경 고화질 업그레이드 ============
# 큰 문서는 배경을 낮은 DPI 로 저장해 용량/시간을 아낀다.
# 사용자가 실제로 보는 쪽은 원본(.src)에서 300dpi 로 다시 렌더해
# 점점 또렷해지게 한다. (글자는 개별 텍스트 상자로 이미 분리돼 있으므로
# 배경에서는 redact 로 지운 뒤 순수 배경만 렌더한다.)
_HIBG_CACHE = {}          # (ref, pno) -> url
_HIBG_LOCK = threading.Lock()


def _render_hi_bg(src, pno):
    """원본 PDF 의 pno 쪽 배경을 고해상도(300dpi)로 렌더해 URL 을 돌려준다."""
    import pymupdf
    doc = pymupdf.open(src)
    try:
        page = doc[pno]
        pw, ph = page.rect.width, page.rect.height
        if pw <= 0 or ph <= 0:
            return None
        lines, math_regions = _pdf_page_lines(page)
        # 글자·수식 영역을 redact 해 순수 배경만 남긴다
        try:
            for mr in math_regions:
                page.add_redact_annot(pymupdf.Rect(
                    mr["x0"] - 0.3, mr["y0"] - 0.5,
                    mr["x1"] + 0.3, mr["y1"] + 0.5))
            for ln in lines:
                if ln.get("has_math"):
                    for wd in ln["words"]:
                        page.add_redact_annot(pymupdf.Rect(
                            wd["x0"] - 0.3, wd["y0"] - 0.5,
                            wd["x1"] + 0.3, wd["y1"] + 0.5))
                else:
                    page.add_redact_annot(pymupdf.Rect(
                        ln["x0"] - 0.3, ln["y0"] - 0.6,
                        ln["x1"] + 0.3, ln["y1"] + 0.6))
            kw = {"images": getattr(pymupdf, "PDF_REDACT_IMAGE_PIXELS",
                                   pymupdf.PDF_REDACT_IMAGE_NONE)}
            touched = getattr(pymupdf, "PDF_REDACT_LINE_ART_REMOVE_IF_TOUCHED", None)
            if touched is not None:
                kw["graphics"] = touched
            try:
                page.apply_redactions(**kw)
            except TypeError:
                try:
                    page.apply_redactions(images=kw["images"])
                except Exception:
                    page.apply_redactions(images=pymupdf.PDF_REDACT_IMAGE_NONE)
        except Exception:
            pass
        # 고해상도 렌더 (픽셀 상한으로 과도한 메모리 방지)
        dpi = 300
        area_in2 = (pw / 72.0) * (ph / 72.0)
        max_px = 18_000_000
        if area_in2 * dpi * dpi > max_px:
            dpi = max(220, int((max_px / max(1e-6, area_in2)) ** 0.5))
        pm = page.get_pixmap(dpi=dpi, alpha=False)
        try:
            zoom = dpi / 72.0
            wipes = []
            for mr in math_regions:
                wipes.append((mr["x0"] - 1.8, mr["y0"] - 1.8,
                              mr["x1"] + 1.8, mr["y1"] + 1.8))
            try:
                for rr in page_rules(page):
                    wipes.append((rr[0] - 0.8, rr[1] - 1.4,
                                  rr[2] + 0.8, rr[3] + 1.4))
            except Exception:
                pass
            for x0, y0, x1, y1 in wipes:
                pr = pymupdf.Rect(x0 * zoom, y0 * zoom, x1 * zoom, y1 * zoom)
                try:
                    pm.set_rect(pr, (255, 255, 255))
                except Exception:
                    pass
        except Exception:
            pass
        raw = pm.tobytes("png")
        pm = None
        d_url = _img_data_url(raw, keep_big=True)
        if not d_url:
            return None
        return _save_import_img(d_url)
    except Exception as e:
        print(f"[hibg] {pno}쪽 렌더 실패: {e}")
        return None
    finally:
        doc.close()


@app.route("/api/import/bg/<ref>/<int:pno>", methods=["GET"])
def import_hibg(ref, pno):
    """특정 쪽 배경을 고해상도로 렌더해 URL 을 돌려준다 (읽는 중 점점 또렷해짐)."""
    ref = re.sub(r"[^0-9a-zA-Z_\-]", "", ref or "")[:40]
    src = os.path.join(DOCS_DIR, f"{ref}.src")
    if not ref or not os.path.exists(src):
        return jsonify({"ok": False, "error": "원본 없음"}), 404
    key = (ref, pno)
    with _HIBG_LOCK:
        if key in _HIBG_CACHE:
            return jsonify({"ok": True, "url": _HIBG_CACHE[key]})
    try:
        url = _render_hi_bg(src, pno)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    if not url:
        return jsonify({"ok": False, "error": "렌더 실패"}), 500
    with _HIBG_LOCK:
        _HIBG_CACHE[key] = url
        if len(_HIBG_CACHE) > 400:
            _HIBG_CACHE.pop(next(iter(_HIBG_CACHE)), None)
    return jsonify({"ok": True, "url": url})
